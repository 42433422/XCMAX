from __future__ import annotations

import hashlib
import json
import os
import re
import threading
import uuid
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import UTC, datetime
from difflib import unified_diff
from pathlib import Path
from typing import Any

from app.infrastructure.rag import (
    HybridRetriever,
    RetrievedChunk,
    SemanticChunker,
    get_default_embedder,
)
from app.infrastructure.rag.citation_tracker import Citation, CitationTracker
from app.infrastructure.rag.dataset_vector_index import (
    DatasetVectorIndexBackend,
    DatasetVectorPgIndex,
    DatasetVectorSQLiteIndex,
    default_dataset_vector_index_path,
)
from app.utils.operational_errors import RECOVERABLE_ERRORS
from app.utils.path_utils import get_app_data_dir, get_upload_dir
from app.utils.safe_download_path import UnsafeDownloadPathError, resolve_under_allowed_dirs


@dataclass
class DatasetDocument:
    document_id: str
    source: str
    parser: str
    text_length: int
    chunk_count: int
    tenant_id: str = "default"
    version: int = 1
    version_label: str = "v1"
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "document_id": self.document_id,
            "source": self.source,
            "parser": self.parser,
            "text_length": self.text_length,
            "chunk_count": self.chunk_count,
            "tenant_id": self.tenant_id,
            "version": self.version,
            "version_label": self.version_label,
            "metadata": self.metadata,
        }


@dataclass
class DatasetRebuildJob:
    job_id: str
    dataset_id: str
    status: str = "queued"
    tenant_id: str = ""
    metadata_filter: dict[str, Any] = field(default_factory=dict)
    document_count: int = 0
    chunk_count: int = 0
    error: str = ""
    attempt_count: int = 0
    max_attempts: int = 1
    worker_id: str = ""
    created_at: str = field(default_factory=lambda: _utc_now_iso())
    queued_at: str = field(default_factory=lambda: _utc_now_iso())
    started_at: str = ""
    completed_at: str = ""
    cancelled_at: str = ""
    updated_at: str = field(default_factory=lambda: _utc_now_iso())

    def to_dict(self) -> dict[str, Any]:
        return {
            "job_id": self.job_id,
            "dataset_id": self.dataset_id,
            "status": self.status,
            "tenant_id": self.tenant_id,
            "metadata_filter": self.metadata_filter,
            "document_count": self.document_count,
            "chunk_count": self.chunk_count,
            "error": self.error,
            "attempt_count": self.attempt_count,
            "max_attempts": self.max_attempts,
            "worker_id": self.worker_id,
            "created_at": self.created_at,
            "queued_at": self.queued_at,
            "started_at": self.started_at,
            "completed_at": self.completed_at,
            "cancelled_at": self.cancelled_at,
            "updated_at": self.updated_at,
        }


DATASET_READ_PERMISSION = "dataset.read"
DATASET_WRITE_PERMISSION = "dataset.write"
DATASET_ADMIN_PERMISSION = "dataset.admin"
PUBLIC_KNOWLEDGE_TENANT_ID = "public"
REBUILD_TERMINAL_STATUSES = {"completed", "failed", "cancelled"}


def _semantic_embedding_available(embedding_count: int = 0) -> bool:
    """UI contract: true when a live embedder exists or chunks already carry vectors."""
    if int(embedding_count or 0) > 0:
        return True
    try:
        from app.infrastructure.rag import get_default_embedder

        return get_default_embedder() is not None
    except RECOVERABLE_ERRORS:
        return False


@dataclass(frozen=True)
class DatasetAccessContext:
    actor_id: str = ""
    tenant_id: str = ""
    permissions: frozenset[str] = field(default_factory=frozenset)
    is_admin: bool = False

    def to_dict(self) -> dict[str, Any]:
        return {
            "actor_id": self.actor_id,
            "tenant_id": self.tenant_id,
            "permissions": sorted(self.permissions),
            "is_admin": self.is_admin,
        }


@dataclass
class _DatasetState:
    dataset_id: str
    documents: dict[str, DatasetDocument] = field(default_factory=dict)
    chunks: list[RetrievedChunk] = field(default_factory=list)
    index: dict[str, Any] = field(default_factory=dict)
    rebuild_jobs: dict[str, DatasetRebuildJob] = field(default_factory=dict)


class DatasetRagApplicationService:
    """Minimal in-process dataset lifecycle for document RAG and citation QA."""

    def __init__(
        self,
        *,
        embedder: Callable[[str], list[float]] | None = None,
        allowed_roots: list[Path] | None = None,
        storage_path: str | Path | None = None,
        max_concurrent_rebuild_jobs: int | None = None,
        rebuild_workers_enabled: bool = True,
        vector_index_backend: DatasetVectorIndexBackend | None = None,
        vector_index_backend_name: str | None = None,
        vector_index_path: str | Path | None = None,
    ) -> None:
        self._embedder = embedder if embedder is not None else get_default_embedder()
        self._chunker = SemanticChunker(embedder=self._embedder)
        self._allowed_roots = allowed_roots
        self._storage_path = (
            Path(storage_path).resolve() if storage_path else _default_storage_path()
        )
        self._vector_index_backend = (
            vector_index_backend
            if vector_index_backend is not None
            else _build_dataset_vector_index_backend(
                backend_name=vector_index_backend_name,
                storage_path=self._storage_path,
                vector_index_path=vector_index_path,
            )
        )
        self._lock = threading.Lock()
        self._max_concurrent_rebuild_jobs = _resolve_max_concurrent_rebuild_jobs(
            max_concurrent_rebuild_jobs
        )
        self._rebuild_workers_enabled = bool(rebuild_workers_enabled)
        self._datasets: dict[str, _DatasetState] = {}
        self._storage_signature: tuple[int, int, int] | None = None
        self._load_persisted_state()
        if self._rebuild_workers_enabled:
            self._schedule_rebuild_jobs()

    def ingest_document(
        self,
        *,
        dataset_id: str,
        source: str = "",
        text: str = "",
        file_path: str = "",
        document_id: str = "",
        chunk_strategy: str = "semantic",
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        metadata: dict[str, Any] | None = None,
        tenant_id: str = "",
        version: int | str | None = None,
        version_label: str = "",
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        source_label = source.strip() or file_path.strip() or "inline"
        base_metadata = dict(metadata or {})
        requested_tenant = str(
            tenant_id or base_metadata.get("tenant_id") or base_metadata.get("user_id") or ""
        )
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            requested_tenant,
            required_permission=DATASET_WRITE_PERMISSION,
            default_without_context="default",
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied

        try:
            if text.strip():
                extracted_text = text.strip()
                parser = "inline_text"
                extract_metadata: dict[str, Any] = {}
            elif file_path.strip():
                path = self._resolve_file_path(file_path)
                extracted_text, parser, extract_metadata = self._extract_file_text(path)
                source_label = source.strip() or path.name
                base_metadata.setdefault("file_path", str(path))
            else:
                raise ValueError("text or file_path is required")

            if not extracted_text.strip():
                raise ValueError("document text is empty")

            chunks = self._split_text(
                extracted_text,
                strategy=chunk_strategy,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            if not chunks:
                raise ValueError("document produced no chunks")

            base_metadata.update(extract_metadata)
            base_metadata["tenant_id"] = tenant_key
            with self._lock:
                self._reload_external_state_locked()
                state = self._datasets.setdefault(dataset_key, _DatasetState(dataset_key))
                document_version = self._resolve_document_version(
                    state,
                    source=source_label,
                    tenant_id=tenant_key,
                    requested=version,
                )
            version_text = version_label.strip() or f"v{document_version}"
            base_metadata["document_version"] = document_version
            base_metadata["version_label"] = version_text
            doc_id = document_id.strip() or _stable_document_id(
                dataset_key,
                tenant_key,
                source_label,
                document_version,
                extracted_text,
            )
            document = DatasetDocument(
                document_id=doc_id,
                source=source_label,
                parser=parser,
                text_length=len(extracted_text),
                chunk_count=len(chunks),
                tenant_id=tenant_key,
                version=document_version,
                version_label=version_text,
                metadata=base_metadata,
            )
            retrieved_chunks = [
                RetrievedChunk(
                    text=chunk.text,
                    score=0.0,
                    source=source_label,
                    chunk_index=chunk.chunk_index,
                    char_start=chunk.char_start,
                    char_end=chunk.char_end,
                    metadata={
                        "dataset_id": dataset_key,
                        "document_id": doc_id,
                        "source": source_label,
                        "parser": parser,
                        "strategy": chunk.strategy,
                        "tenant_id": tenant_key,
                        "document_version": document_version,
                        "version_label": version_text,
                        **_embedding_metadata(self._embedder, chunk.text),
                        **base_metadata,
                    },
                    source_url=source_label,
                )
                for chunk in chunks
            ]

            with self._lock:
                self._reload_external_state_locked()
                state = self._datasets.setdefault(dataset_key, _DatasetState(dataset_key))
                state.documents[doc_id] = document
                state.chunks = [
                    c
                    for c in state.chunks
                    if not isinstance(c.metadata, dict) or c.metadata.get("document_id") != doc_id
                ]
                state.chunks.extend(retrieved_chunks)
                self._renumber_chunks(state)
                self._sync_vector_index_locked(state)
                self._refresh_index_metadata(state)
                self._persist_locked()

            return {
                "success": True,
                "dataset_id": dataset_key,
                "document": document.to_dict(),
                "chunk_count": len(chunks),
            }
        except (*RECOVERABLE_ERRORS, UnsafeDownloadPathError) as exc:
            return {
                "success": False,
                "dataset_id": dataset_key,
                "message": str(exc),
                "error_code": "dataset_ingest_failed",
            }

    def delete_document(
        self,
        dataset_id: str,
        document_id: str,
        *,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        doc_key = document_id.strip()
        context = _coerce_access_context(access_context)
        denied = _ensure_dataset_permission(
            context,
            DATASET_WRITE_PERMISSION,
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            if state is None or doc_key not in state.documents:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "document_id": doc_key,
                    "message": "document not found",
                }
            document = state.documents[doc_key]
            denied = _ensure_tenant_allowed(
                context,
                document.tenant_id,
                dataset_id=dataset_key,
                operation="delete_document",
            )
            if denied is not None:
                denied["document_id"] = doc_key
                return denied
            state.documents.pop(doc_key, None)
            before = len(state.chunks)
            state.chunks = [
                c
                for c in state.chunks
                if not isinstance(c.metadata, dict) or c.metadata.get("document_id") != doc_key
            ]
            self._renumber_chunks(state)
            self._sync_vector_index_locked(state)
            self._refresh_index_metadata(state)
            self._persist_locked()
        return {
            "success": True,
            "dataset_id": dataset_key,
            "document_id": doc_key,
            "deleted_chunks": before - len(state.chunks),
        }

    def set_document_publication_status(
        self,
        dataset_id: str,
        document_id: str,
        publication_status: str,
        *,
        reason: str,
        expected_status: str | None = None,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        from app.application.dataset_rag_publication import (
            set_document_publication_status as _impl,
        )

        return _impl(
            self,
            dataset_id,
            document_id,
            publication_status,
            reason=reason,
            expected_status=expected_status,
            access_context=access_context,
        )


    def diff_versions(
        self,
        *,
        dataset_id: str,
        source: str,
        from_version: str | int,
        to_version: str | int = "latest",
        tenant_id: str = "",
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        source_label = source.strip()
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_READ_PERMISSION,
            default_without_context="default",
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            if state is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "message": "dataset not found",
                }
            from_doc = self._resolve_document_for_version(
                state,
                source=source_label,
                tenant_id=tenant_key,
                version=from_version,
            )
            to_doc = self._resolve_document_for_version(
                state,
                source=source_label,
                tenant_id=tenant_key,
                version=to_version,
            )
            if from_doc is None or to_doc is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "source": source_label,
                    "tenant_id": tenant_key,
                    "message": "document version not found",
                    "from_version": str(from_version),
                    "to_version": str(to_version),
                }
            from_text = self._document_text_locked(state, from_doc.document_id)
            to_text = self._document_text_locked(state, to_doc.document_id)

        from_lines = from_text.splitlines() or ([from_text] if from_text else [])
        to_lines = to_text.splitlines() or ([to_text] if to_text else [])
        diff_lines = list(
            unified_diff(
                from_lines,
                to_lines,
                fromfile=f"{source_label}@{from_doc.version_label}",
                tofile=f"{source_label}@{to_doc.version_label}",
                lineterm="",
            )
        )
        added = [
            line[1:] for line in diff_lines if line.startswith("+") and not line.startswith("+++")
        ]
        removed = [
            line[1:] for line in diff_lines if line.startswith("-") and not line.startswith("---")
        ]
        return {
            "success": True,
            "dataset_id": dataset_key,
            "source": source_label,
            "tenant_id": tenant_key,
            "from_document": from_doc.to_dict(),
            "to_document": to_doc.to_dict(),
            "from_version": from_doc.version,
            "to_version": to_doc.version,
            "changed": from_text != to_text,
            "added_lines": added,
            "removed_lines": removed,
            "diff": diff_lines,
        }

    def rollback_document_version(
        self,
        *,
        dataset_id: str,
        source: str,
        target_version: str | int,
        tenant_id: str = "",
        metadata: dict[str, Any] | None = None,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        source_label = source.strip()
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_WRITE_PERMISSION,
            default_without_context="default",
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            if state is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "message": "dataset not found",
                }
            target_doc = self._resolve_document_for_version(
                state,
                source=source_label,
                tenant_id=tenant_key,
                version=target_version,
            )
            if target_doc is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "source": source_label,
                    "tenant_id": tenant_key,
                    "message": "document version not found",
                    "target_version": str(target_version),
                }
            rollback_text = self._document_text_locked(state, target_doc.document_id)
            rollback_metadata = dict(target_doc.metadata or {})
            rollback_metadata.update(dict(metadata or {}))
            rollback_metadata.update(
                {
                    "rollback": True,
                    "rollback_from_version": target_doc.version,
                    "rollback_from_document_id": target_doc.document_id,
                    "rollback_at": _utc_now_iso(),
                }
            )

        result = self.ingest_document(
            dataset_id=dataset_key,
            source=source_label,
            text=rollback_text,
            tenant_id=tenant_key,
            metadata=rollback_metadata,
            chunk_strategy="fixed",
            access_context=access_context,
        )
        result["rolled_back_from"] = target_doc.to_dict()
        return result

    def start_rebuild_index(
        self,
        *,
        dataset_id: str,
        tenant_id: str = "",
        metadata_filter: dict[str, Any] | None = None,
        background: bool = True,
        max_attempts: int = 1,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_WRITE_PERMISSION,
            default_without_context="",
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        job = DatasetRebuildJob(
            job_id=f"rag_rebuild_{uuid.uuid4().hex[:12]}",
            dataset_id=dataset_key,
            tenant_id=tenant_key,
            metadata_filter=dict(metadata_filter or {}),
            max_attempts=max(1, min(int(max_attempts or 1), 5)),
        )
        starts: list[tuple[str, str]] = []
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            if state is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "message": "dataset not found",
                }
            state.rebuild_jobs[job.job_id] = job
            if background and self._rebuild_workers_enabled:
                starts = self._claim_next_rebuild_jobs_locked()
            self._persist_locked()
            job_payload = self._rebuild_job_to_dict_locked(state, job)

        if background:
            self._start_rebuild_threads(starts)
            return {"success": True, "job": job_payload, "background": True}

        self._run_rebuild_job(dataset_key, job.job_id)
        return {
            "success": True,
            "job": self.get_rebuild_job(dataset_key, job.job_id).get("job", job.to_dict()),
            "background": False,
        }

    def cancel_rebuild_job(
        self,
        dataset_id: str,
        job_id: str,
        *,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        context = _coerce_access_context(access_context)
        denied = _ensure_dataset_permission(
            context,
            DATASET_WRITE_PERMISSION,
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            job = state.rebuild_jobs.get(job_id) if state is not None else None
            if job is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "job_id": job_id,
                    "message": "rebuild job not found",
                }
            denied = _ensure_tenant_allowed(
                context,
                job.tenant_id,
                dataset_id=dataset_key,
                operation="cancel_rebuild_job",
            )
            if denied is not None:
                denied["job_id"] = job_id
                return denied
            if job.status in REBUILD_TERMINAL_STATUSES:
                return {
                    "success": True,
                    "dataset_id": dataset_key,
                    "job": self._rebuild_job_to_dict_locked(state, job),
                    "message": f"rebuild job already {job.status}",
                }
            if job.status != "queued":
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "job_id": job_id,
                    "error_code": "dataset_rebuild_cancel_failed",
                    "message": "only queued rebuild jobs can be cancelled",
                    "job": self._rebuild_job_to_dict_locked(state, job),
                }
            now = _utc_now_iso()
            job.status = "cancelled"
            job.cancelled_at = now
            job.completed_at = now
            job.updated_at = now
            self._persist_locked()
            return {
                "success": True,
                "dataset_id": dataset_key,
                "job": self._rebuild_job_to_dict_locked(state, job),
            }

    def drain_rebuild_queue(self, *, max_jobs: int | None = None) -> dict[str, Any]:
        drained: list[dict[str, Any]] = []
        limit = max_jobs if max_jobs is not None else 1000
        while len(drained) < max(0, int(limit)):
            with self._lock:
                self._reload_external_state_locked()
                next_job = self._claim_next_rebuild_jobs_locked(limit=1)
                self._persist_locked()
            if not next_job:
                break
            dataset_id, job_id = next_job[0]
            self._run_rebuild_job(dataset_id, job_id, claimed=True, schedule_next=False)
            job_status = self.get_rebuild_job(dataset_id, job_id)
            if isinstance(job_status.get("job"), dict):
                drained.append(job_status["job"])
        return {
            "success": True,
            "drained_count": len(drained),
            "jobs": drained,
        }

    def get_rebuild_job(
        self,
        dataset_id: str,
        job_id: str,
        *,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        context = _coerce_access_context(access_context)
        denied = _ensure_dataset_permission(
            context,
            DATASET_READ_PERMISSION,
            dataset_id=dataset_key,
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            job = state.rebuild_jobs.get(job_id) if state is not None else None
            if job is None:
                return {
                    "success": False,
                    "dataset_id": dataset_key,
                    "job_id": job_id,
                    "message": "rebuild job not found",
                }
            denied = _ensure_tenant_allowed(
                context,
                job.tenant_id,
                dataset_id=dataset_key,
                operation="get_rebuild_job",
            )
            if denied is not None:
                denied["job_id"] = job_id
                return denied
            return {
                "success": True,
                "dataset_id": dataset_key,
                "job": self._rebuild_job_to_dict_locked(state, job),
            }

    def status(
        self,
        dataset_id: str = "",
        *,
        tenant_id: str = "",
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
        include_documents: bool = True,
    ) -> dict[str, Any]:
        tenant_filter, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_READ_PERMISSION,
            default_without_context="",
            dataset_id=_clean_key(dataset_id, default="default") if dataset_id.strip() else "",
        )
        if denied is not None:
            return denied
        with self._lock:
            self._reload_external_state_locked()
            if dataset_id.strip():
                dataset_key = _clean_key(dataset_id, default="default")
                state = self._datasets.get(dataset_key)
                return self._status_for_state(
                    dataset_key,
                    state,
                    tenant_id_filter=tenant_filter,
                    include_documents=include_documents,
                )
            datasets = {
                key: self._status_for_state(
                    key,
                    state,
                    tenant_id_filter=tenant_filter,
                    include_documents=include_documents,
                )
                for key, state in sorted(self._datasets.items())
            }
        return {
            "success": True,
            "datasets": datasets,
            "dataset_count": len(datasets),
            "document_count": sum(item["document_count"] for item in datasets.values()),
            "chunk_count": sum(item["chunk_count"] for item in datasets.values()),
            "storage_path": str(self._storage_path),
            "persistent": True,
        }

    def query(
        self,
        *,
        dataset_id: str,
        query: str,
        top_k: int = 5,
        tenant_id: str = "",
        version: str | int = "",
        metadata_filter: dict[str, Any] | None = None,
        rerank: bool = False,
        include_public: bool = False,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        dataset_key = _clean_key(dataset_id, default="default")
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_READ_PERMISSION,
            default_without_context="",
            dataset_id=dataset_key,
        )
        if denied is not None:
            denied.update({"chunks": [], "citations": [], "answer": ""})
            return denied
        query_text = query.strip()
        if not query_text:
            return {
                "success": False,
                "dataset_id": dataset_key,
                "message": "query is required",
                "chunks": [],
            }
        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            chunks = list(state.chunks) if state is not None else []
            index_snapshot = dict(state.index or {}) if state is not None else {}
        include_public_scope = bool(
            include_public
            and tenant_key
            and tenant_key != PUBLIC_KNOWLEDGE_TENANT_ID
        )
        candidate_limit = max(50, max(1, min(int(top_k), 50)) * 4)
        vector_candidates = self._query_vector_index_candidates(
            dataset_id=dataset_key,
            query=query_text,
            top_k=candidate_limit,
            tenant_id=tenant_key,
            version=version,
            metadata_filter=metadata_filter or {},
        )
        if vector_candidates is not None and include_public_scope:
            public_candidates = self._query_vector_index_candidates(
                dataset_id=dataset_key,
                query=query_text,
                top_k=candidate_limit,
                tenant_id=PUBLIC_KNOWLEDGE_TENANT_ID,
                version=version,
                metadata_filter={
                    **(metadata_filter or {}),
                    "publication_status": "published",
                },
            )
            if public_candidates is None:
                vector_candidates = None
            else:
                vector_candidates = _deduplicate_chunks(
                    [*vector_candidates, *public_candidates]
                )
        used_vector_backend = vector_candidates is not None
        if vector_candidates is not None:
            chunks = vector_candidates
            index_snapshot["query_backend"] = self._vector_index_backend_name()
        else:
            private_chunks = _filter_chunks(
                chunks,
                tenant_id=tenant_key,
                version=version,
                metadata_filter=metadata_filter or {},
            )
            public_chunks = (
                _filter_chunks(
                    chunks,
                    tenant_id=PUBLIC_KNOWLEDGE_TENANT_ID,
                    version=version,
                    metadata_filter={
                        **(metadata_filter or {}),
                        "publication_status": "published",
                    },
                )
                if include_public_scope
                else []
            )
            chunks = _deduplicate_chunks([*private_chunks, *public_chunks])
            index_snapshot["query_backend"] = "in_memory_hybrid"
        if not chunks:
            return {
                "success": True,
                "dataset_id": dataset_key,
                "query": query_text,
                "chunks": [],
                "citations": [],
                "answer": "",
                "tenant_id": tenant_key,
                "include_public": include_public_scope,
                "visible_tenant_ids": _visible_tenant_ids(
                    tenant_key, include_public_scope
                ),
                "version": str(version or ""),
                "vector_backend_used": used_vector_backend,
                "index": index_snapshot,
            }
        retriever = HybridRetriever(embedder=self._embedder, top_k=max(1, min(int(top_k), 50)))
        retriever.index(chunks)
        top = retriever.retrieve(query_text)
        if rerank:
            top = _rerank_chunks(query_text, top, top_k=max(1, min(int(top_k), 50)))
        return {
            "success": True,
            "dataset_id": dataset_key,
            "query": query_text,
            "chunks": [_chunk_to_dict(c, public=True) for c in top],
            "tenant_id": tenant_key,
            "include_public": include_public_scope,
            "visible_tenant_ids": _visible_tenant_ids(
                tenant_key, include_public_scope
            ),
            "version": str(version or ""),
            "metadata_filter": metadata_filter or {},
            "rerank": bool(rerank),
            "vector_backend_used": used_vector_backend,
            "index": index_snapshot,
        }

    def knowledge_graph(
        self,
        dataset_id: str,
        *,
        tenant_id: str = "",
        limit: int = 120,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """Build a bounded, explainable graph from persisted documents and chunks."""

        dataset_key = _clean_key(dataset_id, default="default")
        tenant_key, denied = _resolve_tenant_for_access(
            access_context,
            tenant_id,
            required_permission=DATASET_READ_PERMISSION,
            default_without_context="",
            dataset_id=dataset_key,
        )
        if denied is not None:
            denied.update({"nodes": [], "edges": [], "stats": {}})
            return denied

        with self._lock:
            self._reload_external_state_locked()
            state = self._datasets.get(dataset_key)
            if state is None:
                documents: list[DatasetDocument] = []
                chunks: list[RetrievedChunk] = []
            else:
                documents = [
                    doc
                    for doc in state.documents.values()
                    if not tenant_key or doc.tenant_id == tenant_key
                ]
                chunks = _filter_chunks(
                    list(state.chunks),
                    tenant_id=tenant_key,
                    version="",
                    metadata_filter={},
                )

        return _build_knowledge_graph_payload(
            dataset_id=dataset_key,
            tenant_id=tenant_key,
            documents=documents,
            chunks=chunks,
            limit=max(20, min(int(limit), 240)),
        )

    def answer(
        self,
        *,
        dataset_id: str,
        query: str,
        top_k: int = 5,
        llm_call: Callable[[str, str], str] | None = None,
        tenant_id: str = "",
        version: str | int = "",
        metadata_filter: dict[str, Any] | None = None,
        rerank: bool = False,
        include_public: bool = False,
        access_context: DatasetAccessContext | dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        result = self.query(
            dataset_id=dataset_id,
            query=query,
            top_k=top_k,
            tenant_id=tenant_id,
            version=version,
            metadata_filter=metadata_filter,
            rerank=rerank,
            include_public=include_public,
            access_context=access_context,
        )
        if not result.get("success") or not result.get("chunks"):
            result.setdefault("answer", "")
            result.setdefault("citations", [])
            result.setdefault("raw", "")
            return result

        retrieved = [_dict_to_retrieved_chunk(row) for row in result["chunks"]]
        tracker = CitationTracker(retrieved_chunks=retrieved)
        prompt = tracker.format_for_prompt()
        raw_answer = (
            llm_call(str(result.get("query") or ""), prompt)
            if llm_call is not None
            else _deterministic_answer(str(result.get("query") or ""), retrieved)
        )
        clean_answer, citations = tracker.attach_citations(raw_answer)
        result["answer"] = clean_answer
        result["raw"] = raw_answer
        result["citations"] = [_citation_to_dict(c) for c in citations]
        return result

    def _resolve_file_path(self, file_path: str) -> Path:
        return resolve_under_allowed_dirs(file_path, self._allowed_file_roots())

    def _allowed_file_roots(self) -> list[Path]:
        if self._allowed_roots is not None:
            return [Path(root).resolve() for root in self._allowed_roots]
        return [
            Path(get_upload_dir()).resolve(),
            Path(get_app_data_dir()).resolve(),
            Path.cwd().resolve(),
        ]

    def _extract_file_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        if not path.exists() or not path.is_file():
            raise ValueError(f"file not found: {path}")
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf_text(path)
        if suffix == ".docx":
            return self._extract_docx_text(path)
        if suffix in {".xlsx", ".xls"}:
            return self._extract_excel_text(path)
        if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
            return (
                path.read_text(encoding="utf-8", errors="replace"),
                "text_file",
                {
                    "extension": suffix,
                },
            )
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_pdf_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError("pdfplumber is required to ingest PDF documents") from exc

        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages.append(f"[page {index}]\n{page_text}")
            page_count = len(pdf.pages)
        return "\n\n".join(pages), "pdfplumber", {"extension": ".pdf", "page_count": page_count}

    def _extract_docx_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to ingest DOCX documents") from exc

        doc = Document(str(path))
        parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts), "python-docx", {"extension": ".docx"}

    def _extract_excel_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            return self._extract_xlsx_text(path)
        if suffix == ".xls":
            return self._extract_xls_text(path)
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_xlsx_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to ingest XLSX documents") from exc

        max_rows_per_sheet = 5_000
        max_cols = 100
        parts: list[str] = []
        sheet_count = 0
        row_count = 0
        wb = load_workbook(str(path), read_only=True, data_only=True)
        try:
            for sheet in wb.worksheets:
                sheet_count += 1
                sheet_lines: list[str] = []
                for index, row in enumerate(sheet.iter_rows(values_only=True)):
                    if index >= max_rows_per_sheet:
                        sheet_lines.append(f"... truncated after {max_rows_per_sheet} rows")
                        break
                    cells: list[str] = []
                    for col_index, cell in enumerate(row or ()):
                        if col_index >= max_cols:
                            break
                        if cell is None:
                            continue
                        text = str(cell).strip()
                        if text:
                            cells.append(text)
                    if cells:
                        sheet_lines.append("\t".join(cells))
                        row_count += 1
                if sheet_lines:
                    parts.append(f"[sheet {sheet.title}]\n" + "\n".join(sheet_lines))
        finally:
            wb.close()
        return (
            "\n\n".join(parts),
            "openpyxl",
            {
                "extension": ".xlsx",
                "sheet_count": sheet_count,
                "row_count": row_count,
            },
        )

    def _extract_xls_text(self, path: Path) -> tuple[str, str, dict[str, Any]]:
        try:
            import xlrd
        except ImportError as exc:
            raise RuntimeError("xlrd is required to ingest XLS documents") from exc

        max_rows_per_sheet = 5_000
        max_cols = 100
        parts: list[str] = []
        sheet_count = 0
        row_count = 0
        book = xlrd.open_workbook(str(path))
        for sheet in book.sheets():
            sheet_count += 1
            sheet_lines: list[str] = []
            nrows = min(int(sheet.nrows), max_rows_per_sheet)
            for row_index in range(nrows):
                cells: list[str] = []
                ncols = min(int(sheet.ncols), max_cols)
                for col_index in range(ncols):
                    value = sheet.cell_value(row_index, col_index)
                    if value is None or value == "":
                        continue
                    text = str(value).strip()
                    if text:
                        cells.append(text)
                if cells:
                    sheet_lines.append("\t".join(cells))
                    row_count += 1
            if int(sheet.nrows) > max_rows_per_sheet:
                sheet_lines.append(f"... truncated after {max_rows_per_sheet} rows")
            if sheet_lines:
                parts.append(f"[sheet {sheet.name}]\n" + "\n".join(sheet_lines))
        return (
            "\n\n".join(parts),
            "xlrd",
            {
                "extension": ".xls",
                "sheet_count": sheet_count,
                "row_count": row_count,
            },
        )

    def _split_text(
        self,
        text: str,
        *,
        strategy: str,
        chunk_size: int,
        chunk_overlap: int,
    ):
        if strategy == "semantic":
            return self._chunker.split_by_semantic(text)
        return self._chunker.split_by_fixed(
            text,
            chunk_size=max(50, min(int(chunk_size), 5000)),
            chunk_overlap=max(0, min(int(chunk_overlap), 500)),
        )

    def _load_persisted_state(self) -> None:
        try:
            if not self._storage_path.exists():
                self._datasets = {}
                self._storage_signature = None
                return
            raw = json.loads(self._storage_path.read_text(encoding="utf-8"))
            datasets = raw.get("datasets") if isinstance(raw, dict) else {}
            if not isinstance(datasets, dict):
                raise ValueError("dataset storage must contain a datasets object")
            loaded: dict[str, _DatasetState] = {}
            for dataset_id, payload in datasets.items():
                if not isinstance(payload, dict):
                    continue
                documents_payload = payload.get("documents")
                chunks_payload = payload.get("chunks")
                jobs_payload = payload.get("rebuild_jobs")
                documents = (
                    {
                        str(doc_id): _document_from_dict(row)
                        for doc_id, row in (documents_payload or {}).items()
                        if isinstance(row, dict)
                    }
                    if isinstance(documents_payload, dict)
                    else {}
                )
                chunks = (
                    [
                        _dict_to_retrieved_chunk(row)
                        for row in (chunks_payload or [])
                        if isinstance(row, dict)
                    ]
                    if isinstance(chunks_payload, list)
                    else []
                )
                index_payload = payload.get("index")
                state = _DatasetState(
                    str(dataset_id),
                    documents=documents,
                    chunks=chunks,
                    index=dict(index_payload or {}) if isinstance(index_payload, dict) else {},
                    rebuild_jobs={
                        str(job_id): _rebuild_job_from_dict(row)
                        for job_id, row in (jobs_payload or {}).items()
                        if isinstance(row, dict)
                    }
                    if isinstance(jobs_payload, dict)
                    else {},
                )
                self._recover_rebuild_jobs_locked(state)
                self._renumber_chunks(state)
                self._sync_vector_index_locked(state)
                self._refresh_index_metadata(state)
                loaded[str(dataset_id)] = state
            self._datasets = loaded
            self._storage_signature = self._current_storage_signature()
        except RECOVERABLE_ERRORS:
            return

    def _current_storage_signature(self) -> tuple[int, int, int] | None:
        try:
            stat = self._storage_path.stat()
            return (int(stat.st_mtime_ns), int(stat.st_size), int(stat.st_ino))
        except FileNotFoundError:
            return None

    def _reload_external_state_locked(self) -> bool:
        signature = self._current_storage_signature()
        if signature == self._storage_signature:
            return False
        previous_signature = self._storage_signature
        self._load_persisted_state()
        return self._storage_signature != previous_signature

    def _persist_locked(self) -> None:
        payload = {
            "version": 1,
            "datasets": {
                dataset_id: {
                    "dataset_id": state.dataset_id,
                    "documents": {
                        doc_id: doc.to_dict() for doc_id, doc in sorted(state.documents.items())
                    },
                    "chunks": [_chunk_to_dict(chunk) for chunk in state.chunks],
                    "index": dict(state.index or {}),
                    "rebuild_jobs": {
                        job_id: job.to_dict() for job_id, job in sorted(state.rebuild_jobs.items())
                    },
                }
                for dataset_id, state in sorted(self._datasets.items())
            },
        }
        self._storage_path.parent.mkdir(parents=True, exist_ok=True)
        tmp_path = self._storage_path.with_suffix(self._storage_path.suffix + ".tmp")
        encoded = json.dumps(payload, ensure_ascii=False, indent=2)
        with open(tmp_path, "w", encoding="utf-8") as handle:
            handle.write(encoded)
            handle.flush()
            os.fsync(handle.fileno())
        tmp_path.replace(self._storage_path)
        try:
            directory_fd = os.open(self._storage_path.parent, os.O_RDONLY)
            try:
                os.fsync(directory_fd)
            finally:
                os.close(directory_fd)
        except OSError:
            pass
        self._storage_signature = self._current_storage_signature()

    @staticmethod
    def _renumber_chunks(state: _DatasetState) -> None:
        for index, chunk in enumerate(state.chunks):
            chunk.chunk_index = index

    @staticmethod
    def _resolve_document_version(
        state: _DatasetState,
        *,
        source: str,
        tenant_id: str,
        requested: int | str | None,
    ) -> int:
        text = str(requested or "").strip()
        if text:
            if text.lower().startswith("v") and text[1:].isdigit():
                return max(1, int(text[1:]))
            if text.isdigit():
                return max(1, int(text))
            raise ValueError("version must be an integer, vN, or empty")
        versions = [
            int(doc.version or 1)
            for doc in state.documents.values()
            if doc.source == source and doc.tenant_id == tenant_id
        ]
        return (max(versions) + 1) if versions else 1

    @staticmethod
    def _resolve_document_for_version(
        state: _DatasetState,
        *,
        source: str,
        tenant_id: str,
        version: str | int,
    ) -> DatasetDocument | None:
        candidates = [
            doc
            for doc in state.documents.values()
            if doc.source == source and doc.tenant_id == tenant_id
        ]
        if not candidates:
            return None
        version_text = str(version or "").strip()
        if not version_text or version_text.lower() == "latest":
            return max(candidates, key=lambda doc: int(doc.version or 1))
        normalized = version_text[1:] if version_text.lower().startswith("v") else version_text
        for doc in candidates:
            if str(doc.version) == normalized or doc.version_label == version_text:
                return doc
        return None

    @staticmethod
    def _document_text_locked(state: _DatasetState, document_id: str) -> str:
        chunks = [
            chunk
            for chunk in state.chunks
            if isinstance(chunk.metadata, dict)
            and str(chunk.metadata.get("document_id") or "") == document_id
        ]
        chunks.sort(key=lambda item: (item.char_start, item.chunk_index))
        return "\n".join(chunk.text for chunk in chunks if chunk.text)

    def _schedule_rebuild_jobs(self) -> None:
        if not self._rebuild_workers_enabled:
            return
        with self._lock:
            starts = self._claim_next_rebuild_jobs_locked()
            if starts:
                self._persist_locked()
        self._start_rebuild_threads(starts)

    def _claim_next_rebuild_jobs_locked(
        self,
        *,
        limit: int | None = None,
    ) -> list[tuple[str, str]]:
        running = sum(
            1
            for state in self._datasets.values()
            for job in state.rebuild_jobs.values()
            if job.status == "running"
        )
        capacity = max(0, self._max_concurrent_rebuild_jobs - running)
        if limit is not None:
            capacity = min(capacity, max(0, int(limit)))
        if capacity <= 0:
            return []

        queued: list[tuple[str, DatasetRebuildJob]] = []
        for dataset_id, state in sorted(self._datasets.items()):
            for job in state.rebuild_jobs.values():
                if job.status == "queued":
                    queued.append((dataset_id, job))
        queued.sort(key=lambda item: (item[1].queued_at or item[1].created_at, item[1].job_id))

        starts: list[tuple[str, str]] = []
        for dataset_id, job in queued[:capacity]:
            now = _utc_now_iso()
            job.status = "running"
            job.started_at = now if not job.started_at else job.started_at
            job.updated_at = now
            job.attempt_count += 1
            job.worker_id = f"rag_worker_{uuid.uuid4().hex[:8]}"
            starts.append((dataset_id, job.job_id))
        return starts

    def _start_rebuild_threads(self, starts: list[tuple[str, str]]) -> None:
        for dataset_id, job_id in starts:
            thread = threading.Thread(
                target=self._run_rebuild_job,
                kwargs={
                    "dataset_id": dataset_id,
                    "job_id": job_id,
                    "claimed": True,
                },
                daemon=True,
            )
            thread.start()

    def _run_rebuild_job(
        self,
        dataset_id: str,
        job_id: str,
        *,
        claimed: bool = False,
        schedule_next: bool = True,
    ) -> None:
        with self._lock:
            state = self._datasets.get(dataset_id)
            job = state.rebuild_jobs.get(job_id) if state is not None else None
            if state is None or job is None:
                return
            if job.status in REBUILD_TERMINAL_STATUSES:
                return
            if not claimed:
                now = _utc_now_iso()
                job.status = "running"
                job.started_at = now if not job.started_at else job.started_at
                job.updated_at = now
                job.attempt_count += 1
                job.worker_id = f"rag_worker_{uuid.uuid4().hex[:8]}"
                self._persist_locked()

        try:
            with self._lock:
                state = self._datasets.get(dataset_id)
                job = state.rebuild_jobs.get(job_id) if state is not None else None
                if state is None or job is None:
                    return
                matched_document_ids: set[str] = set()
                rebuilt_chunks = 0
                for chunk in state.chunks:
                    metadata = dict(chunk.metadata or {})
                    if job.tenant_id and str(metadata.get("tenant_id") or "") != job.tenant_id:
                        continue
                    if job.metadata_filter and not _metadata_matches(chunk, job.metadata_filter):
                        continue
                    metadata.pop("_embedding", None)
                    metadata.update(_embedding_metadata(self._embedder, chunk.text))
                    chunk.metadata = metadata
                    rebuilt_chunks += 1
                    document_id = str(metadata.get("document_id") or "")
                    if document_id:
                        matched_document_ids.add(document_id)
                self._renumber_chunks(state)
                self._sync_vector_index_locked(state)
                self._refresh_index_metadata(state)
                job.status = "completed"
                job.document_count = len(matched_document_ids)
                job.chunk_count = rebuilt_chunks
                now = _utc_now_iso()
                job.completed_at = now
                job.updated_at = now
                self._persist_locked()
        except RECOVERABLE_ERRORS as exc:
            with self._lock:
                state = self._datasets.get(dataset_id)
                job = state.rebuild_jobs.get(job_id) if state is not None else None
                if job is not None:
                    now = _utc_now_iso()
                    if job.attempt_count < max(1, int(job.max_attempts or 1)):
                        job.status = "queued"
                        job.queued_at = now
                    else:
                        job.status = "failed"
                        job.completed_at = now
                    job.error = str(exc)
                    job.updated_at = now
                    self._persist_locked()
        finally:
            if schedule_next:
                self._schedule_rebuild_jobs()

    def _query_vector_index_candidates(
        self,
        *,
        dataset_id: str,
        query: str,
        top_k: int,
        tenant_id: str,
        version: str | int,
        metadata_filter: dict[str, Any],
    ) -> list[RetrievedChunk] | None:
        if self._vector_index_backend is None or self._embedder is None:
            return None
        try:
            query_vector = self._embedder(query)
            if not isinstance(query_vector, list) or not query_vector:
                return None
            return self._vector_index_backend.query(
                dataset_id,
                [float(value) for value in query_vector],
                top_k=top_k,
                tenant_id=tenant_id,
                version=version,
                metadata_filter=metadata_filter,
            )
        except RECOVERABLE_ERRORS:
            return None

    def _sync_vector_index_locked(self, state: _DatasetState) -> None:
        if self._vector_index_backend is None:
            state.index["vector_backend_sync_status"] = "disabled"
            return
        try:
            count = self._vector_index_backend.replace_dataset(state.dataset_id, state.chunks)
            state.index["vector_backend_sync_status"] = "synced"
            state.index["vector_backend_synced_chunks"] = count
            state.index["vector_backend_synced_at"] = _utc_now_iso()
        except RECOVERABLE_ERRORS as exc:
            state.index["vector_backend_sync_status"] = "failed"
            state.index["vector_backend_error"] = str(exc)

    def _vector_index_backend_name(self) -> str:
        if self._vector_index_backend is None:
            return "none"
        return str(getattr(self._vector_index_backend, "backend_name", "vector_index"))

    def _vector_index_status(self, dataset_id: str) -> dict[str, Any]:
        if self._vector_index_backend is None:
            return {
                "backend": "none",
                "persistent": False,
                "dataset_id": dataset_id,
                "chunk_count": 0,
                "index_exists": False,
            }
        try:
            return dict(self._vector_index_backend.status(dataset_id))
        except RECOVERABLE_ERRORS as exc:
            return {
                "backend": self._vector_index_backend_name(),
                "persistent": True,
                "dataset_id": dataset_id,
                "chunk_count": 0,
                "index_exists": False,
                "status": "failed",
                "error": str(exc),
            }

    def _refresh_index_metadata(self, state: _DatasetState) -> None:
        embedding_count = sum(
            1
            for chunk in state.chunks
            if isinstance(chunk.metadata, dict)
            and isinstance(chunk.metadata.get("_embedding"), list)
        )
        vector_status = self._vector_index_status(state.dataset_id)
        previous = dict(state.index or {})
        state.index = {
            "schema_version": 2,
            "retriever": "hybrid_bm25_vector",
            "reranker": "lexical_overlap_v1",
            "document_count": len(state.documents),
            "chunk_count": len(state.chunks),
            "embedding_count": embedding_count,
            "embedding_persisted": embedding_count > 0,
            "vector_backend": vector_status,
            "vector_backend_name": vector_status.get("backend", self._vector_index_backend_name()),
            "vector_backend_chunk_count": int(vector_status.get("chunk_count") or 0),
            "vector_backend_persistent": bool(vector_status.get("persistent")),
            "vector_backend_sync_status": previous.get("vector_backend_sync_status", ""),
            "vector_backend_synced_chunks": int(previous.get("vector_backend_synced_chunks") or 0),
            "vector_backend_synced_at": str(previous.get("vector_backend_synced_at") or ""),
            "vector_backend_error": str(previous.get("vector_backend_error") or ""),
            "updated_at": _utc_now_iso(),
        }

    def _status_for_state(
        self,
        dataset_id: str,
        state: _DatasetState | None,
        *,
        tenant_id_filter: str = "",
        include_documents: bool = True,
    ) -> dict[str, Any]:
        if state is None:
            empty_index = {
                "semantic_embedding_available": _semantic_embedding_available(0),
                "embedding_count": 0,
            }
            return {
                "success": True,
                "dataset_id": dataset_id,
                "document_count": 0,
                "chunk_count": 0,
                "documents": [],
                "tenant_ids": [],
                "versions": [],
                "index": empty_index,
                "rebuild_jobs": [],
                "rebuild_job_count": 0,
                "rebuild_queue": _empty_rebuild_queue_summary(
                    self._max_concurrent_rebuild_jobs,
                    worker_enabled=self._rebuild_workers_enabled,
                ),
                "storage_path": str(self._storage_path),
                "persistent": True,
            }
        documents = [
            doc
            for doc in state.documents.values()
            if not tenant_id_filter or doc.tenant_id == tenant_id_filter
        ]
        chunks = [
            chunk
            for chunk in state.chunks
            if not tenant_id_filter
            or str((chunk.metadata or {}).get("tenant_id") or "") == tenant_id_filter
        ]
        rebuild_jobs = [
            job
            for job in state.rebuild_jobs.values()
            if not tenant_id_filter or job.tenant_id == tenant_id_filter
        ]
        embedding_count = sum(
            1
            for chunk in chunks
            if isinstance(chunk.metadata, dict)
            and isinstance(chunk.metadata.get("_embedding"), list)
        )
        index = dict(state.index or {})
        if tenant_id_filter:
            index.update(
                {
                    "document_count": len(documents),
                    "chunk_count": len(chunks),
                    "embedding_count": embedding_count,
                    "filtered_by_tenant": tenant_id_filter,
                }
            )
        index["embedding_count"] = int(index.get("embedding_count") or embedding_count or 0)
        index["semantic_embedding_available"] = _semantic_embedding_available(
            int(index.get("embedding_count") or 0)
        )
        rebuild_jobs = [
            self._rebuild_job_to_dict_locked(state, job)
            for job in sorted(
                rebuild_jobs,
                key=lambda item: item.created_at,
                reverse=True,
            )[:10]
        ]
        return {
            "success": True,
            "dataset_id": dataset_id,
            "document_count": len(documents),
            "chunk_count": len(chunks),
            # Omniscient/graph HUD only needs counts; full document rows are heavy (hundreds+).
            "documents": [doc.to_dict() for doc in documents] if include_documents else [],
            "tenant_ids": sorted({doc.tenant_id for doc in documents}),
            "versions": sorted({doc.version_label for doc in documents}),
            "index": index,
            "rebuild_jobs": rebuild_jobs,
            "rebuild_job_count": len(rebuild_jobs),
            "rebuild_queue": self._rebuild_queue_summary_locked(
                state,
                tenant_id_filter=tenant_id_filter,
            ),
            "storage_path": str(self._storage_path),
            "persistent": True,
        }

    def _rebuild_job_to_dict_locked(
        self,
        state: _DatasetState,
        job: DatasetRebuildJob,
    ) -> dict[str, Any]:
        payload = job.to_dict()
        payload["queue_position"] = self._rebuild_job_queue_position_locked(state, job)
        payload["max_concurrent_jobs"] = self._max_concurrent_rebuild_jobs
        return payload

    @staticmethod
    def _rebuild_job_queue_position_locked(
        state: _DatasetState,
        job: DatasetRebuildJob,
    ) -> int:
        if job.status != "queued":
            return 0
        queued = [item for item in state.rebuild_jobs.values() if item.status == "queued"]
        queued.sort(key=lambda item: (item.queued_at or item.created_at, item.job_id))
        for index, item in enumerate(queued, start=1):
            if item.job_id == job.job_id:
                return index
        return 0

    def _rebuild_queue_summary_locked(
        self,
        state: _DatasetState,
        *,
        tenant_id_filter: str = "",
    ) -> dict[str, Any]:
        jobs = [
            job
            for job in state.rebuild_jobs.values()
            if not tenant_id_filter or job.tenant_id == tenant_id_filter
        ]
        counts = {
            "queued": 0,
            "running": 0,
            "completed": 0,
            "failed": 0,
            "cancelled": 0,
        }
        for job in jobs:
            if job.status in counts:
                counts[job.status] += 1
        queued = sorted(
            (job for job in jobs if job.status == "queued"),
            key=lambda item: (item.queued_at or item.created_at, item.job_id),
        )
        running = sorted(
            (job for job in jobs if job.status == "running"),
            key=lambda item: (item.started_at or item.updated_at, item.job_id),
        )
        return {
            "max_concurrent_jobs": self._max_concurrent_rebuild_jobs,
            "worker_enabled": self._rebuild_workers_enabled,
            "queued": counts["queued"],
            "running": counts["running"],
            "completed": counts["completed"],
            "failed": counts["failed"],
            "cancelled": counts["cancelled"],
            "next_job_id": queued[0].job_id if queued else "",
            "running_job_ids": [job.job_id for job in running],
        }

    @staticmethod
    def _recover_rebuild_jobs_locked(state: _DatasetState) -> None:
        for job in state.rebuild_jobs.values():
            if job.status == "running":
                now = _utc_now_iso()
                job.status = "queued"
                job.error = "requeued after service restart"
                job.worker_id = ""
                job.queued_at = now
                job.updated_at = now


def _default_storage_path() -> Path:
    configured = (
        os.environ.get("DATASET_RAG_STORE_PATH")
        or os.environ.get("XCAGI_DATASET_RAG_STORE_PATH")
        or ""
    ).strip()
    if configured:
        return Path(configured).expanduser().resolve()
    return Path(get_app_data_dir()).resolve() / "dataset_rag" / "datasets.json"


def _build_dataset_vector_index_backend(
    *,
    backend_name: str | None,
    storage_path: Path,
    vector_index_path: str | Path | None,
) -> DatasetVectorIndexBackend | None:
    configured = (
        backend_name
        if backend_name is not None
        else os.environ.get("DATASET_RAG_VECTOR_BACKEND")
        or os.environ.get("XCAGI_DATASET_RAG_VECTOR_BACKEND")
        or "sqlite"
    )
    name = str(configured or "").strip().lower()
    if name in {"", "none", "disabled", "off", "json", "memory"}:
        return None
    if name in {"sqlite", "sqlite_vector"}:
        path = (
            Path(vector_index_path).expanduser().resolve()
            if vector_index_path is not None
            else default_dataset_vector_index_path(storage_path)
        )
        return DatasetVectorSQLiteIndex(path)
    if name in {"pgvector", "postgres", "postgresql"}:
        database_url = (
            os.environ.get("DATASET_RAG_PGVECTOR_DATABASE_URL")
            or os.environ.get("XCAGI_DATASET_RAG_PGVECTOR_DATABASE_URL")
            or os.environ.get("PGVECTOR_DATABASE_URL")
            or os.environ.get("DATABASE_URL")
            or ""
        ).strip()
        dimension_raw = (
            os.environ.get("DATASET_RAG_PGVECTOR_DIMENSION")
            or os.environ.get("XCAGI_DATASET_RAG_PGVECTOR_DIMENSION")
            or "256"
        )
        try:
            dimension = int(dimension_raw)
        except (TypeError, ValueError):
            dimension = 256
        return DatasetVectorPgIndex(database_url, dimension=dimension)
    raise ValueError(f"unsupported dataset vector backend: {configured}")


def _resolve_max_concurrent_rebuild_jobs(configured: int | None) -> int:
    if configured is not None:
        return max(1, min(int(configured), 8))
    raw = os.environ.get("DATASET_RAG_REBUILD_MAX_CONCURRENT", "").strip()
    if not raw:
        raw = os.environ.get("XCAGI_DATASET_RAG_REBUILD_MAX_CONCURRENT", "").strip()
    if raw.isdigit():
        return max(1, min(int(raw), 8))
    return 1


def _empty_rebuild_queue_summary(
    max_concurrent_jobs: int,
    *,
    worker_enabled: bool,
) -> dict[str, Any]:
    return {
        "max_concurrent_jobs": max_concurrent_jobs,
        "worker_enabled": worker_enabled,
        "queued": 0,
        "running": 0,
        "completed": 0,
        "failed": 0,
        "cancelled": 0,
        "next_job_id": "",
        "running_job_ids": [],
    }


def _clean_key(value: str, *, default: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in {"-", "_", "."} else "_" for ch in value.strip())
    return cleaned.strip("._-") or default


def _coerce_access_context(
    value: DatasetAccessContext | dict[str, Any] | None,
) -> DatasetAccessContext | None:
    if value is None:
        return None
    if isinstance(value, DatasetAccessContext):
        return value
    permissions_value = value.get("permissions") if isinstance(value, dict) else None
    if isinstance(permissions_value, str):
        permissions = frozenset(
            part.strip() for part in permissions_value.replace(";", ",").split(",") if part.strip()
        )
    elif isinstance(permissions_value, (list, tuple, set, frozenset)):
        permissions = frozenset(
            str(part).strip() for part in permissions_value if str(part).strip()
        )
    else:
        permissions = frozenset()
    return DatasetAccessContext(
        actor_id=str(value.get("actor_id") or value.get("user_id") or ""),
        tenant_id=_clean_key(str(value.get("tenant_id") or ""), default="")
        if value.get("tenant_id")
        else "",
        permissions=permissions,
        is_admin=bool(value.get("is_admin") or value.get("admin")),
    )


def _has_dataset_permission(context: DatasetAccessContext | None, permission: str) -> bool:
    if context is None:
        return True
    if context.is_admin or DATASET_ADMIN_PERMISSION in context.permissions:
        return True
    if permission in context.permissions:
        return True
    prefix = permission.split(".", 1)[0]
    return f"{prefix}.*" in context.permissions or "*" in context.permissions


def _dataset_permission_denied(
    *,
    dataset_id: str,
    permission: str,
    message: str,
    context: DatasetAccessContext | None,
) -> dict[str, Any]:
    return {
        "success": False,
        "dataset_id": dataset_id,
        "error_code": "dataset_permission_denied",
        "message": message,
        "required_permission": permission,
        "access": context.to_dict() if context is not None else {},
    }


def _ensure_dataset_permission(
    context: DatasetAccessContext | None,
    permission: str,
    *,
    dataset_id: str,
) -> dict[str, Any] | None:
    if _has_dataset_permission(context, permission):
        return None
    return _dataset_permission_denied(
        dataset_id=dataset_id,
        permission=permission,
        message=f"{permission} permission is required",
        context=context,
    )


def _ensure_tenant_allowed(
    context: DatasetAccessContext | None,
    tenant_id: str,
    *,
    dataset_id: str,
    operation: str,
) -> dict[str, Any] | None:
    if context is None or context.is_admin or DATASET_ADMIN_PERMISSION in context.permissions:
        return None
    actor_tenant = _clean_key(context.tenant_id, default="") if context.tenant_id else ""
    target_tenant = _clean_key(str(tenant_id or ""), default="") if tenant_id else ""
    if not actor_tenant:
        return _dataset_permission_denied(
            dataset_id=dataset_id,
            permission=DATASET_READ_PERMISSION,
            message=f"{operation} requires an actor tenant context",
            context=context,
        )
    if not target_tenant:
        return _dataset_permission_denied(
            dataset_id=dataset_id,
            permission=DATASET_ADMIN_PERMISSION,
            message=f"{operation} across all tenants requires dataset admin",
            context=context,
        )
    if actor_tenant != target_tenant:
        return _dataset_permission_denied(
            dataset_id=dataset_id,
            permission=DATASET_ADMIN_PERMISSION,
            message=f"{operation} cannot access tenant {target_tenant}",
            context=context,
        )
    return None


def _resolve_tenant_for_access(
    access_context: DatasetAccessContext | dict[str, Any] | None,
    requested_tenant_id: str,
    *,
    required_permission: str,
    default_without_context: str,
    dataset_id: str,
) -> tuple[str, dict[str, Any] | None]:
    context = _coerce_access_context(access_context)
    denied = _ensure_dataset_permission(context, required_permission, dataset_id=dataset_id)
    if denied is not None:
        return "", denied
    requested = (
        _clean_key(str(requested_tenant_id or ""), default="") if requested_tenant_id else ""
    )
    if context is None:
        if requested:
            return requested, None
        return (
            _clean_key(str(default_without_context), default=default_without_context)
            if default_without_context
            else ""
        ), None
    if context.is_admin or DATASET_ADMIN_PERMISSION in context.permissions:
        if requested:
            return requested, None
        return (
            _clean_key(str(default_without_context), default=default_without_context)
            if default_without_context
            else ""
        ), None
    actor_tenant = _clean_key(context.tenant_id, default="") if context.tenant_id else ""
    if not actor_tenant:
        return "", _dataset_permission_denied(
            dataset_id=dataset_id,
            permission=required_permission,
            message="dataset tenant context is required",
            context=context,
        )
    if requested and requested != actor_tenant:
        return "", _dataset_permission_denied(
            dataset_id=dataset_id,
            permission=DATASET_ADMIN_PERMISSION,
            message=f"tenant {requested} is outside requester scope",
            context=context,
        )
    return actor_tenant, None


def _stable_document_id(
    dataset_id: str,
    tenant_id: str,
    source: str,
    version: int,
    text: str,
) -> str:
    digest = hashlib.sha256(
        f"{dataset_id}\0{tenant_id}\0{source}\0{version}\0{text}".encode()
    ).hexdigest()
    return f"doc_{digest[:16]}"


def _document_from_dict(data: dict[str, Any]) -> DatasetDocument:
    metadata = dict(data.get("metadata") or {})
    version = int(data.get("version") or metadata.get("document_version") or 1)
    return DatasetDocument(
        document_id=str(data.get("document_id") or ""),
        source=str(data.get("source") or ""),
        parser=str(data.get("parser") or ""),
        text_length=int(data.get("text_length") or 0),
        chunk_count=int(data.get("chunk_count") or 0),
        tenant_id=_clean_key(
            str(data.get("tenant_id") or metadata.get("tenant_id") or "default"),
            default="default",
        ),
        version=version,
        version_label=str(
            data.get("version_label") or metadata.get("version_label") or f"v{version}"
        ),
        metadata=metadata,
    )


def _rebuild_job_from_dict(data: dict[str, Any]) -> DatasetRebuildJob:
    created_at = str(data.get("created_at") or _utc_now_iso())
    queued_at = str(data.get("queued_at") or created_at)
    return DatasetRebuildJob(
        job_id=str(data.get("job_id") or ""),
        dataset_id=str(data.get("dataset_id") or ""),
        status=str(data.get("status") or "queued"),
        tenant_id=str(data.get("tenant_id") or ""),
        metadata_filter=dict(data.get("metadata_filter") or {}),
        document_count=int(data.get("document_count") or 0),
        chunk_count=int(data.get("chunk_count") or 0),
        error=str(data.get("error") or ""),
        attempt_count=int(data.get("attempt_count") or 0),
        max_attempts=max(1, int(data.get("max_attempts") or 1)),
        worker_id=str(data.get("worker_id") or ""),
        created_at=created_at,
        queued_at=queued_at,
        started_at=str(data.get("started_at") or ""),
        completed_at=str(data.get("completed_at") or ""),
        cancelled_at=str(data.get("cancelled_at") or ""),
        updated_at=str(data.get("updated_at") or queued_at),
    )


def _chunk_to_dict(chunk: RetrievedChunk, *, public: bool = False) -> dict[str, Any]:
    metadata = dict(chunk.metadata or {})
    retrieval_method = str(chunk.source or "")
    source = str(metadata.get("source") or chunk.source or "")
    if retrieval_method in {"hybrid", "vector", "bm25", "hybrid+rerank", "vector+rerank", "bm25+rerank"}:
        metadata.setdefault("retrieval_method", retrieval_method)
    if public:
        metadata = {key: value for key, value in metadata.items() if not str(key).startswith("_")}
    return {
        "text": chunk.text,
        "score": chunk.score,
        "source": source,
        "chunk_index": chunk.chunk_index,
        "char_start": chunk.char_start,
        "char_end": chunk.char_end,
        "metadata": metadata,
        "source_url": chunk.source_url,
        "page": chunk.page,
    }


def _dict_to_retrieved_chunk(data: dict[str, Any]) -> RetrievedChunk:
    return RetrievedChunk(
        text=str(data.get("text") or ""),
        score=float(data.get("score") or 0.0),
        source=str(data.get("source") or ""),
        chunk_index=int(data.get("chunk_index") or 0),
        char_start=int(data.get("char_start") or 0),
        char_end=int(data.get("char_end") or 0),
        metadata=dict(data.get("metadata") or {}),
        source_url=str(data.get("source_url") or ""),
        page=data.get("page") if isinstance(data.get("page"), int) else None,
    )


def _citation_to_dict(citation: Citation) -> dict[str, Any]:
    return {
        "index": citation.index,
        "text": citation.text,
        "source": citation.source,
        "chunk_index": citation.chunk_index,
        "char_range": list(citation.char_range),
        "source_url": citation.source_url,
        "page": citation.page,
    }


def _deterministic_answer(query: str, chunks: list[RetrievedChunk]) -> str:
    if not chunks:
        return ""
    excerpt = chunks[0].text.strip().replace("\n", " ")
    if len(excerpt) > 320:
        excerpt = excerpt[:317].rstrip() + "..."
    prefix = f"Based on the retrieved dataset evidence for {query!r}: " if query else ""
    return f"{prefix}{excerpt} [1]"


_GRAPH_TOPIC_METADATA_KEYS = frozenset(
    {"topic", "topics", "tag", "tags", "entity", "entities", "keywords", "category", "doc_type"}
)
_GRAPH_TOPIC_STOPWORDS = frozenset(
    {
        "about",
        "after",
        "before",
        "default",
        "document",
        "from",
        "into",
        "persy",
        "that",
        "their",
        "this",
        "with",
        "以及",
        "他们",
        "内容",
        "可以",
        "如何",
        "我们",
        "文件",
        "是否",
        "知识",
        "系统",
        "资料",
        "这个",
        "这些",
        "进行",
        "需要",
    }
)


from app.application.dataset_rag_graph import (
    _build_knowledge_graph_payload,
)


def _embedding_metadata(
    embedder: Callable[[str], list[float]] | None,
    text: str,
) -> dict[str, Any]:
    if embedder is None:
        return {}
    try:
        embedding = embedder(text)
    except RECOVERABLE_ERRORS:
        return {}
    if not isinstance(embedding, list) or not embedding:
        return {}
    try:
        return {"_embedding": [float(value) for value in embedding]}
    except (TypeError, ValueError):
        return {}


def _filter_chunks(
    chunks: list[RetrievedChunk],
    *,
    tenant_id: str,
    version: str | int,
    metadata_filter: dict[str, Any],
) -> list[RetrievedChunk]:
    selected = list(chunks)
    tenant_key = _clean_key(str(tenant_id or ""), default="") if tenant_id else ""
    if tenant_key:
        selected = [
            chunk
            for chunk in selected
            if str((chunk.metadata or {}).get("tenant_id") or "") == tenant_key
        ]
    if metadata_filter:
        selected = [chunk for chunk in selected if _metadata_matches(chunk, metadata_filter)]

    version_text = str(version or "").strip()
    if not version_text:
        return selected
    if version_text.lower() == "latest":
        latest_by_scope: dict[tuple[str, str], int] = {}
        for chunk in selected:
            metadata = chunk.metadata or {}
            scope = (
                str(metadata.get("tenant_id") or ""),
                str(metadata.get("source") or chunk.source or ""),
            )
            latest_by_scope[scope] = max(
                latest_by_scope.get(scope, 0),
                int(metadata.get("document_version") or 1),
            )
        return [
            chunk
            for chunk in selected
            if int((chunk.metadata or {}).get("document_version") or 1)
            == latest_by_scope.get(
                (
                    str((chunk.metadata or {}).get("tenant_id") or ""),
                    str((chunk.metadata or {}).get("source") or chunk.source or ""),
                ),
                1,
            )
        ]
    normalized = version_text[1:] if version_text.lower().startswith("v") else version_text
    return [
        chunk
        for chunk in selected
        if str((chunk.metadata or {}).get("document_version") or "") == normalized
        or str((chunk.metadata or {}).get("version_label") or "") == version_text
    ]


def _deduplicate_chunks(chunks: list[RetrievedChunk]) -> list[RetrievedChunk]:
    seen: set[tuple[str, int, str]] = set()
    result: list[RetrievedChunk] = []
    for chunk in chunks:
        metadata = chunk.metadata or {}
        key = (
            str(metadata.get("document_id") or chunk.source or ""),
            int(chunk.chunk_index or 0),
            str(metadata.get("tenant_id") or ""),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(chunk)
    return result


def _visible_tenant_ids(tenant_id: str, include_public: bool) -> list[str]:
    visible = [tenant_id] if tenant_id else []
    if include_public and PUBLIC_KNOWLEDGE_TENANT_ID not in visible:
        visible.append(PUBLIC_KNOWLEDGE_TENANT_ID)
    return visible


def _metadata_matches(chunk: RetrievedChunk, metadata_filter: dict[str, Any]) -> bool:
    metadata = dict(chunk.metadata or {})
    metadata.setdefault("source", chunk.source)
    for key, expected in metadata_filter.items():
        actual = metadata.get(str(key))
        if isinstance(expected, list):
            expected_values = {str(item) for item in expected}
            if str(actual) not in expected_values:
                return False
        elif isinstance(expected, dict):
            if not isinstance(actual, dict):
                return False
            for nested_key, nested_expected in expected.items():
                if str(actual.get(str(nested_key))) != str(nested_expected):
                    return False
        elif str(actual) != str(expected):
            return False
    return True


def _rerank_chunks(
    query: str,
    chunks: list[RetrievedChunk],
    *,
    top_k: int,
) -> list[RetrievedChunk]:
    query_terms = set(_tokenize_for_rerank(query))
    if not query_terms:
        return chunks[:top_k]
    reranked: list[RetrievedChunk] = []
    for chunk in chunks:
        chunk_terms = set(_tokenize_for_rerank(chunk.text))
        overlap = len(query_terms & chunk_terms)
        exact_bonus = (
            1.0 if query.strip().lower() and query.strip().lower() in chunk.text.lower() else 0.0
        )
        boost = overlap / max(1, len(query_terms)) + exact_bonus
        reranked.append(
            RetrievedChunk(
                text=chunk.text,
                score=float(chunk.score) + boost,
                source=f"{chunk.source}+rerank" if "rerank" not in chunk.source else chunk.source,
                chunk_index=chunk.chunk_index,
                char_start=chunk.char_start,
                char_end=chunk.char_end,
                metadata=chunk.metadata,
                source_url=chunk.source_url,
                page=chunk.page,
            )
        )
    return sorted(reranked, key=lambda item: item.score, reverse=True)[:top_k]


def _tokenize_for_rerank(text: str) -> list[str]:
    tokens: list[str] = []
    for part in re.split(r"([\u4e00-\u9fff])", text):
        if not part:
            continue
        if re.fullmatch(r"[\u4e00-\u9fff]", part):
            tokens.append(part)
        else:
            tokens.extend(re.findall(r"[A-Za-z0-9]+", part.lower()))
    return tokens


def _utc_now_iso() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")


_dataset_rag_app_service: DatasetRagApplicationService | None = None
_dataset_rag_lock = threading.Lock()


def get_dataset_rag_app_service() -> DatasetRagApplicationService:
    global _dataset_rag_app_service
    if _dataset_rag_app_service is None:
        with _dataset_rag_lock:
            if _dataset_rag_app_service is None:
                _dataset_rag_app_service = DatasetRagApplicationService()
    return _dataset_rag_app_service


def reset_dataset_rag_app_service_for_tests() -> None:
    global _dataset_rag_app_service
    with _dataset_rag_lock:
        _dataset_rag_app_service = None
