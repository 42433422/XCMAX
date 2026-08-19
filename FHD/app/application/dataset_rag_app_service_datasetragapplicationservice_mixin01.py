# ruff: noqa
# mypy: ignore-errors
"""Behavior mixin extracted from the public facade class."""
from __future__ import annotations
import importlib

def _facade():
    return importlib.import_module('app.application.dataset_rag_app_service')

class _DatasetRagApplicationServicePart01Mixin:

    def __init__(self, *, embedder: _facade().Callable[[str], list[float]] | None=None, allowed_roots: list[_facade().Path] | None=None, storage_path: str | _facade().Path | None=None, max_concurrent_rebuild_jobs: int | None=None, rebuild_workers_enabled: bool=True, vector_index_backend: _facade().DatasetVectorIndexBackend | None=None, vector_index_backend_name: str | None=None, vector_index_path: str | _facade().Path | None=None) -> None:
        self._embedder = embedder if embedder is not None else _facade().get_default_embedder()
        self._chunker = _facade().SemanticChunker(embedder=self._embedder)
        self._allowed_roots = allowed_roots
        self._storage_path = _facade().Path(storage_path).resolve() if storage_path else _facade()._default_storage_path()
        self._vector_index_backend = vector_index_backend if vector_index_backend is not None else _facade()._build_dataset_vector_index_backend(backend_name=vector_index_backend_name, storage_path=self._storage_path, vector_index_path=vector_index_path)
        self._lock = _facade().threading.Lock()
        self._max_concurrent_rebuild_jobs = _facade()._resolve_max_concurrent_rebuild_jobs(max_concurrent_rebuild_jobs)
        self._rebuild_workers_enabled = bool(rebuild_workers_enabled)
        self._datasets: dict[str, _facade()._DatasetState] = {}
        self._load_persisted_state()
        if self._rebuild_workers_enabled:
            self._schedule_rebuild_jobs()

    def ingest_document(self, *, dataset_id: str, source: str='', text: str='', file_path: str='', document_id: str='', chunk_strategy: str='semantic', chunk_size: int=500, chunk_overlap: int=50, metadata: dict[str, _facade().Any] | None=None, tenant_id: str='', version: int | str | None=None, version_label: str='', access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        source_label = source.strip() or file_path.strip() or 'inline'
        base_metadata = dict(metadata or {})
        requested_tenant = str(tenant_id or base_metadata.get('tenant_id') or base_metadata.get('user_id') or '')
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, requested_tenant, required_permission=_facade().DATASET_WRITE_PERMISSION, default_without_context='default', dataset_id=dataset_key)
        if denied is not None:
            return denied
        try:
            if text.strip():
                extracted_text = text.strip()
                parser = 'inline_text'
                extract_metadata: dict[str, _facade().Any] = {}
            elif file_path.strip():
                path = self._resolve_file_path(file_path)
                (extracted_text, parser, extract_metadata) = self._extract_file_text(path)
                source_label = source.strip() or path.name
                base_metadata.setdefault('file_path', str(path))
            else:
                raise ValueError('text or file_path is required')
            if not extracted_text.strip():
                raise ValueError('document text is empty')
            chunks = self._split_text(extracted_text, strategy=chunk_strategy, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            if not chunks:
                raise ValueError('document produced no chunks')
            base_metadata.update(extract_metadata)
            base_metadata['tenant_id'] = tenant_key
            with self._lock:
                state = self._datasets.setdefault(dataset_key, _facade()._DatasetState(dataset_key))
                document_version = self._resolve_document_version(state, source=source_label, tenant_id=tenant_key, requested=version)
            version_text = version_label.strip() or f'v{document_version}'
            base_metadata['document_version'] = document_version
            base_metadata['version_label'] = version_text
            doc_id = document_id.strip() or _facade()._stable_document_id(dataset_key, tenant_key, source_label, document_version, extracted_text)
            document = _facade().DatasetDocument(document_id=doc_id, source=source_label, parser=parser, text_length=len(extracted_text), chunk_count=len(chunks), tenant_id=tenant_key, version=document_version, version_label=version_text, metadata=base_metadata)
            retrieved_chunks = [_facade().RetrievedChunk(text=chunk.text, score=0.0, source=source_label, chunk_index=chunk.chunk_index, char_start=chunk.char_start, char_end=chunk.char_end, metadata={'dataset_id': dataset_key, 'document_id': doc_id, 'source': source_label, 'parser': parser, 'strategy': chunk.strategy, 'tenant_id': tenant_key, 'document_version': document_version, 'version_label': version_text, **_facade()._embedding_metadata(self._embedder, chunk.text), **base_metadata}, source_url=source_label) for chunk in chunks]
            with self._lock:
                state.documents[doc_id] = document
                state.chunks = [c for c in state.chunks if not isinstance(c.metadata, dict) or c.metadata.get('document_id') != doc_id]
                state.chunks.extend(retrieved_chunks)
                self._renumber_chunks(state)
                self._sync_vector_index_locked(state)
                self._refresh_index_metadata(state)
                self._persist_locked()
            return {'success': True, 'dataset_id': dataset_key, 'document': document.to_dict(), 'chunk_count': len(chunks)}
        except _facade()._DATASET_DOWNLOAD_ERRORS as exc:
            return {'success': False, 'dataset_id': dataset_key, 'message': str(exc), 'error_code': 'dataset_ingest_failed'}

    def delete_document(self, dataset_id: str, document_id: str, *, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        doc_key = document_id.strip()
        context = _facade()._coerce_access_context(access_context)
        denied = _facade()._ensure_dataset_permission(context, _facade().DATASET_WRITE_PERMISSION, dataset_id=dataset_key)
        if denied is not None:
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            if state is None or doc_key not in state.documents:
                return {'success': False, 'dataset_id': dataset_key, 'document_id': doc_key, 'message': 'document not found'}
            document = state.documents[doc_key]
            denied = _facade()._ensure_tenant_allowed(context, document.tenant_id, dataset_id=dataset_key, operation='delete_document')
            if denied is not None:
                denied['document_id'] = doc_key
                return denied
            state.documents.pop(doc_key, None)
            before = len(state.chunks)
            state.chunks = [c for c in state.chunks if not isinstance(c.metadata, dict) or c.metadata.get('document_id') != doc_key]
            self._renumber_chunks(state)
            self._sync_vector_index_locked(state)
            self._refresh_index_metadata(state)
            self._persist_locked()
        return {'success': True, 'dataset_id': dataset_key, 'document_id': doc_key, 'deleted_chunks': before - len(state.chunks)}

    def diff_versions(self, *, dataset_id: str, source: str, from_version: str | int, to_version: str | int='latest', tenant_id: str='', access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        source_label = source.strip()
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_READ_PERMISSION, default_without_context='default', dataset_id=dataset_key)
        if denied is not None:
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            if state is None:
                return {'success': False, 'dataset_id': dataset_key, 'message': 'dataset not found'}
            from_doc = self._resolve_document_for_version(state, source=source_label, tenant_id=tenant_key, version=from_version)
            to_doc = self._resolve_document_for_version(state, source=source_label, tenant_id=tenant_key, version=to_version)
            if from_doc is None or to_doc is None:
                return {'success': False, 'dataset_id': dataset_key, 'source': source_label, 'tenant_id': tenant_key, 'message': 'document version not found', 'from_version': str(from_version), 'to_version': str(to_version)}
            from_text = self._document_text_locked(state, from_doc.document_id)
            to_text = self._document_text_locked(state, to_doc.document_id)
        from_lines = from_text.splitlines() or ([from_text] if from_text else [])
        to_lines = to_text.splitlines() or ([to_text] if to_text else [])
        diff_lines = list(_facade().unified_diff(from_lines, to_lines, fromfile=f'{source_label}@{from_doc.version_label}', tofile=f'{source_label}@{to_doc.version_label}', lineterm=''))
        added = [line[1:] for line in diff_lines if line.startswith('+') and (not line.startswith('+++'))]
        removed = [line[1:] for line in diff_lines if line.startswith('-') and (not line.startswith('---'))]
        return {'success': True, 'dataset_id': dataset_key, 'source': source_label, 'tenant_id': tenant_key, 'from_document': from_doc.to_dict(), 'to_document': to_doc.to_dict(), 'from_version': from_doc.version, 'to_version': to_doc.version, 'changed': from_text != to_text, 'added_lines': added, 'removed_lines': removed, 'diff': diff_lines}

    def rollback_document_version(self, *, dataset_id: str, source: str, target_version: str | int, tenant_id: str='', metadata: dict[str, _facade().Any] | None=None, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        source_label = source.strip()
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_WRITE_PERMISSION, default_without_context='default', dataset_id=dataset_key)
        if denied is not None:
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            if state is None:
                return {'success': False, 'dataset_id': dataset_key, 'message': 'dataset not found'}
            target_doc = self._resolve_document_for_version(state, source=source_label, tenant_id=tenant_key, version=target_version)
            if target_doc is None:
                return {'success': False, 'dataset_id': dataset_key, 'source': source_label, 'tenant_id': tenant_key, 'message': 'document version not found', 'target_version': str(target_version)}
            rollback_text = self._document_text_locked(state, target_doc.document_id)
            rollback_metadata = dict(target_doc.metadata or {})
            rollback_metadata.update(dict(metadata or {}))
            rollback_metadata.update({'rollback': True, 'rollback_from_version': target_doc.version, 'rollback_from_document_id': target_doc.document_id, 'rollback_at': _facade()._utc_now_iso()})
        result = self.ingest_document(dataset_id=dataset_key, source=source_label, text=rollback_text, tenant_id=tenant_key, metadata=rollback_metadata, chunk_strategy='fixed', access_context=access_context)
        result['rolled_back_from'] = target_doc.to_dict()
        return result

    def start_rebuild_index(self, *, dataset_id: str, tenant_id: str='', metadata_filter: dict[str, _facade().Any] | None=None, background: bool=True, max_attempts: int=1, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_WRITE_PERMISSION, default_without_context='', dataset_id=dataset_key)
        if denied is not None:
            return denied
        job = _facade().DatasetRebuildJob(job_id=f'rag_rebuild_{_facade().uuid.uuid4().hex[:12]}', dataset_id=dataset_key, tenant_id=tenant_key, metadata_filter=dict(metadata_filter or {}), max_attempts=max(1, min(int(max_attempts or 1), 5)))
        starts: list[tuple[str, str]] = []
        with self._lock:
            state = self._datasets.get(dataset_key)
            if state is None:
                return {'success': False, 'dataset_id': dataset_key, 'message': 'dataset not found'}
            state.rebuild_jobs[job.job_id] = job
            if background and self._rebuild_workers_enabled:
                starts = self._claim_next_rebuild_jobs_locked()
            self._persist_locked()
            job_payload = self._rebuild_job_to_dict_locked(state, job)
        if background:
            self._start_rebuild_threads(starts)
            return {'success': True, 'job': job_payload, 'background': True}
        self._run_rebuild_job(dataset_key, job.job_id)
        return {'success': True, 'job': self.get_rebuild_job(dataset_key, job.job_id).get('job', job.to_dict()), 'background': False}

    def cancel_rebuild_job(self, dataset_id: str, job_id: str, *, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        context = _facade()._coerce_access_context(access_context)
        denied = _facade()._ensure_dataset_permission(context, _facade().DATASET_WRITE_PERMISSION, dataset_id=dataset_key)
        if denied is not None:
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            job = state.rebuild_jobs.get(job_id) if state is not None else None
            if job is None:
                return {'success': False, 'dataset_id': dataset_key, 'job_id': job_id, 'message': 'rebuild job not found'}
            if state is None:
                return {'success': False, 'dataset_id': dataset_key, 'message': 'dataset not found'}
            denied = _facade()._ensure_tenant_allowed(context, job.tenant_id, dataset_id=dataset_key, operation='cancel_rebuild_job')
            if denied is not None:
                denied['job_id'] = job_id
                return denied
            if job.status in _facade().REBUILD_TERMINAL_STATUSES:
                return {'success': True, 'dataset_id': dataset_key, 'job': self._rebuild_job_to_dict_locked(state, job), 'message': f'rebuild job already {job.status}'}
            if job.status != 'queued':
                return {'success': False, 'dataset_id': dataset_key, 'job_id': job_id, 'error_code': 'dataset_rebuild_cancel_failed', 'message': 'only queued rebuild jobs can be cancelled', 'job': self._rebuild_job_to_dict_locked(state, job)}
            now = _facade()._utc_now_iso()
            job.status = 'cancelled'
            job.cancelled_at = now
            job.completed_at = now
            job.updated_at = now
            self._persist_locked()
            return {'success': True, 'dataset_id': dataset_key, 'job': self._rebuild_job_to_dict_locked(state, job)}

    def drain_rebuild_queue(self, *, max_jobs: int | None=None) -> dict[str, _facade().Any]:
        drained: list[dict[str, _facade().Any]] = []
        limit = max_jobs if max_jobs is not None else 1000
        while len(drained) < max(0, int(limit)):
            with self._lock:
                next_job = self._claim_next_rebuild_jobs_locked(limit=1)
                self._persist_locked()
            if not next_job:
                break
            (dataset_id, job_id) = next_job[0]
            self._run_rebuild_job(dataset_id, job_id, claimed=True, schedule_next=False)
            job_status = self.get_rebuild_job(dataset_id, job_id)
            if isinstance(job_status.get('job'), dict):
                drained.append(job_status['job'])
        return {'success': True, 'drained_count': len(drained), 'jobs': drained}

    def get_rebuild_job(self, dataset_id: str, job_id: str, *, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        context = _facade()._coerce_access_context(access_context)
        denied = _facade()._ensure_dataset_permission(context, _facade().DATASET_READ_PERMISSION, dataset_id=dataset_key)
        if denied is not None:
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            job = state.rebuild_jobs.get(job_id) if state is not None else None
            if job is None:
                return {'success': False, 'dataset_id': dataset_key, 'job_id': job_id, 'message': 'rebuild job not found'}
            if state is None:
                return {'success': False, 'dataset_id': dataset_key, 'message': 'dataset not found'}
            denied = _facade()._ensure_tenant_allowed(context, job.tenant_id, dataset_id=dataset_key, operation='get_rebuild_job')
            if denied is not None:
                denied['job_id'] = job_id
                return denied
            return {'success': True, 'dataset_id': dataset_key, 'job': self._rebuild_job_to_dict_locked(state, job)}

    def status(self, dataset_id: str='', *, tenant_id: str='', access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None, include_documents: bool=True) -> dict[str, _facade().Any]:
        (tenant_filter, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_READ_PERMISSION, default_without_context='', dataset_id=_facade()._clean_key(dataset_id, default='default') if dataset_id.strip() else '')
        if denied is not None:
            return denied
        with self._lock:
            if dataset_id.strip():
                dataset_key = _facade()._clean_key(dataset_id, default='default')
                state = self._datasets.get(dataset_key)
                return self._status_for_state(dataset_key, state, tenant_id_filter=tenant_filter, include_documents=include_documents)
            datasets = {key: self._status_for_state(key, state, tenant_id_filter=tenant_filter, include_documents=include_documents) for (key, state) in sorted(self._datasets.items())}
        return {'success': True, 'datasets': datasets, 'dataset_count': len(datasets), 'document_count': sum((item['document_count'] for item in datasets.values())), 'chunk_count': sum((item['chunk_count'] for item in datasets.values())), 'storage_path': str(self._storage_path), 'persistent': True}

    def query(self, *, dataset_id: str, query: str, top_k: int=5, tenant_id: str='', version: str | int='', metadata_filter: dict[str, _facade().Any] | None=None, rerank: bool=False, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_READ_PERMISSION, default_without_context='', dataset_id=dataset_key)
        if denied is not None:
            denied.update({'chunks': [], 'citations': [], 'answer': ''})
            return denied
        query_text = query.strip()
        if not query_text:
            return {'success': False, 'dataset_id': dataset_key, 'message': 'query is required', 'chunks': []}
        with self._lock:
            state = self._datasets.get(dataset_key)
            chunks = list(state.chunks) if state is not None else []
            index_snapshot = dict(state.index or {}) if state is not None else {}
        vector_candidates = self._query_vector_index_candidates(dataset_id=dataset_key, query=query_text, top_k=max(50, max(1, min(int(top_k), 50)) * 4), tenant_id=tenant_key, version=version, metadata_filter=metadata_filter or {})
        used_vector_backend = vector_candidates is not None
        if vector_candidates is not None:
            chunks = vector_candidates
            index_snapshot['query_backend'] = self._vector_index_backend_name()
        else:
            chunks = _facade()._filter_chunks(chunks, tenant_id=tenant_key, version=version, metadata_filter=metadata_filter or {})
            index_snapshot['query_backend'] = 'in_memory_hybrid'
        if not chunks:
            return {'success': True, 'dataset_id': dataset_key, 'query': query_text, 'chunks': [], 'citations': [], 'answer': '', 'tenant_id': tenant_key, 'version': str(version or ''), 'vector_backend_used': used_vector_backend, 'index': index_snapshot}
        retriever = _facade().HybridRetriever(embedder=self._embedder, top_k=max(1, min(int(top_k), 50)))
        retriever.index(chunks)
        top = retriever.retrieve(query_text)
        if rerank:
            top = _facade()._rerank_chunks(query_text, top, top_k=max(1, min(int(top_k), 50)))
        return {'success': True, 'dataset_id': dataset_key, 'query': query_text, 'chunks': [_facade()._chunk_to_dict(c, public=True) for c in top], 'tenant_id': tenant_key, 'version': str(version or ''), 'metadata_filter': metadata_filter or {}, 'rerank': bool(rerank), 'vector_backend_used': used_vector_backend, 'index': index_snapshot}

    def knowledge_graph(self, dataset_id: str, *, tenant_id: str='', limit: int=120, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        """Build a bounded, explainable graph from persisted documents and chunks."""
        dataset_key = _facade()._clean_key(dataset_id, default='default')
        (tenant_key, denied) = _facade()._resolve_tenant_for_access(access_context, tenant_id, required_permission=_facade().DATASET_READ_PERMISSION, default_without_context='', dataset_id=dataset_key)
        if denied is not None:
            denied.update({'nodes': [], 'edges': [], 'stats': {}})
            return denied
        with self._lock:
            state = self._datasets.get(dataset_key)
            if state is None:
                documents: list[_facade().DatasetDocument] = []
                chunks: list[_facade().RetrievedChunk] = []
            else:
                documents = [doc for doc in state.documents.values() if not tenant_key or doc.tenant_id == tenant_key]
                chunks = _facade()._filter_chunks(list(state.chunks), tenant_id=tenant_key, version='', metadata_filter={})
        return _facade()._build_knowledge_graph_payload(dataset_id=dataset_key, tenant_id=tenant_key, documents=documents, chunks=chunks, limit=max(20, min(int(limit), 240)))

    def answer(self, *, dataset_id: str, query: str, top_k: int=5, llm_call: _facade().Callable[[str, str], str] | None=None, tenant_id: str='', version: str | int='', metadata_filter: dict[str, _facade().Any] | None=None, rerank: bool=False, access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None=None) -> dict[str, _facade().Any]:
        result = self.query(dataset_id=dataset_id, query=query, top_k=top_k, tenant_id=tenant_id, version=version, metadata_filter=metadata_filter, rerank=rerank, access_context=access_context)
        if not result.get('success') or not result.get('chunks'):
            result.setdefault('answer', '')
            result.setdefault('citations', [])
            result.setdefault('raw', '')
            return result
        retrieved = [_facade()._dict_to_retrieved_chunk(row) for row in result['chunks']]
        tracker = _facade().CitationTracker(retrieved_chunks=retrieved)
        prompt = tracker.format_for_prompt()
        raw_answer = llm_call(str(result.get('query') or ''), prompt) if llm_call is not None else _facade()._deterministic_answer(str(result.get('query') or ''), retrieved)
        (clean_answer, citations) = tracker.attach_citations(raw_answer)
        result['answer'] = clean_answer
        result['raw'] = raw_answer
        result['citations'] = [_facade()._citation_to_dict(c) for c in citations]
        return result

    def _resolve_file_path(self, file_path: str) -> _facade().Path:
        return _facade().resolve_under_allowed_dirs(file_path, self._allowed_file_roots())

    def _allowed_file_roots(self) -> list[_facade().Path]:
        if self._allowed_roots is not None:
            return [_facade().Path(root).resolve() for root in self._allowed_roots]
        return [_facade().Path(_facade().get_upload_dir()).resolve(), _facade().Path(_facade().get_app_data_dir()).resolve(), _facade().Path.cwd().resolve()]

    def _extract_file_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        if not path.exists() or not path.is_file():
            raise ValueError(f'file not found: {path}')
        suffix = path.suffix.lower()
        if suffix == '.pdf':
            return self._extract_pdf_text(path)
        if suffix == '.docx':
            return self._extract_docx_text(path)
        if suffix in {'.xlsx', '.xls'}:
            return self._extract_excel_text(path)
        if suffix in {'.txt', '.md', '.csv', '.json', '.log'}:
            return (path.read_text(encoding='utf-8', errors='replace'), 'text_file', {'extension': suffix})
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_pdf_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError('pdfplumber is required to ingest PDF documents') from exc
        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for (index, page) in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or '').strip()
                if page_text:
                    pages.append(f'[page {index}]\n{page_text}')
            page_count = len(pdf.pages)
        return ('\n\n'.join(pages), 'pdfplumber', {'extension': '.pdf', 'page_count': page_count})

    def _extract_docx_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError('python-docx is required to ingest DOCX documents') from exc
        doc = Document(str(path))
        parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))
        return ('\n'.join(parts), 'python-docx', {'extension': '.docx'})

    def _extract_excel_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        suffix = path.suffix.lower()
        if suffix == '.xlsx':
            return self._extract_xlsx_text(path)
        if suffix == '.xls':
            return self._extract_xls_text(path)
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_xlsx_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError('openpyxl is required to ingest XLSX documents') from exc
        max_rows_per_sheet = 5000
        max_cols = 100
        parts: list[str] = []
        sheet_count = 0
        row_count = 0
        wb = load_workbook(str(path), read_only=True, data_only=True)
        try:
            for sheet in wb.worksheets:
                sheet_count += 1
                sheet_lines: list[str] = []
                for (index, row) in enumerate(sheet.iter_rows(values_only=True)):
                    if index >= max_rows_per_sheet:
                        sheet_lines.append(f'... truncated after {max_rows_per_sheet} rows')
                        break
                    cells: list[str] = []
                    for (col_index, cell) in enumerate(row or ()):
                        if col_index >= max_cols:
                            break
                        if cell is None:
                            continue
                        text = str(cell).strip()
                        if text:
                            cells.append(text)
                    if cells:
                        sheet_lines.append('\t'.join(cells))
                        row_count += 1
                if sheet_lines:
                    parts.append(f'[sheet {sheet.title}]\n' + '\n'.join(sheet_lines))
        finally:
            wb.close()
        return ('\n\n'.join(parts), 'openpyxl', {'extension': '.xlsx', 'sheet_count': sheet_count, 'row_count': row_count})
