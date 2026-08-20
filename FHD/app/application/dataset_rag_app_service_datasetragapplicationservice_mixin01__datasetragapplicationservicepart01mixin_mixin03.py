# mypy: disable-error-code="attr-defined, no-any-return, valid-type"
"""Behavior mixin extracted from the public facade class."""

from __future__ import annotations

import importlib


def _facade():
    return importlib.import_module("app.application.dataset_rag_app_service")


class __DatasetRagApplicationServicePart01MixinPart03Mixin:
    def answer(
        self,
        *,
        dataset_id: str,
        query: str,
        top_k: int = 5,
        llm_call: _facade().Callable[[str, str], str] | None = None,
        tenant_id: str = "",
        version: str | int = "",
        metadata_filter: dict[str, _facade().Any] | None = None,
        rerank: bool = False,
        access_context: _facade().DatasetAccessContext | dict[str, _facade().Any] | None = None,
    ) -> dict[str, _facade().Any]:
        result = self.query(
            dataset_id=dataset_id,
            query=query,
            top_k=top_k,
            tenant_id=tenant_id,
            version=version,
            metadata_filter=metadata_filter,
            rerank=rerank,
            access_context=access_context,
        )
        if not result.get("success") or not result.get("chunks"):
            result.setdefault("answer", "")
            result.setdefault("citations", [])
            result.setdefault("raw", "")
            return result
        retrieved = [_facade()._dict_to_retrieved_chunk(row) for row in result["chunks"]]
        tracker = _facade().CitationTracker(retrieved_chunks=retrieved)
        prompt = tracker.format_for_prompt()
        raw_answer = (
            llm_call(str(result.get("query") or ""), prompt)
            if llm_call is not None
            else _facade()._deterministic_answer(str(result.get("query") or ""), retrieved)
        )
        (clean_answer, citations) = tracker.attach_citations(raw_answer)
        result["answer"] = clean_answer
        result["raw"] = raw_answer
        result["citations"] = [_facade()._citation_to_dict(c) for c in citations]
        return result

    def _resolve_file_path(self, file_path: str) -> _facade().Path:
        return _facade().resolve_under_allowed_dirs(file_path, self._allowed_file_roots())

    def _allowed_file_roots(self) -> list[_facade().Path]:
        if self._allowed_roots is not None:
            return [_facade().Path(root).resolve() for root in self._allowed_roots]
        return [
            _facade().Path(_facade().get_upload_dir()).resolve(),
            _facade().Path(_facade().get_app_data_dir()).resolve(),
            _facade().Path.cwd().resolve(),
        ]

    def _extract_file_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        if not path.exists() or not path.is_file():
            raise ValueError(f"file not found: {path}")
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf_text(path)
        if suffix == ".docx":
            return self._extract_docx_text(path)
        if suffix in {".xlsx", ".xls"}:
            return self._extract_excel_text(path)
        if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
            return (
                path.read_text(encoding="utf-8", errors="replace"),
                "text_file",
                {"extension": suffix},
            )
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_pdf_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError("pdfplumber is required to ingest PDF documents") from exc
        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages.append(f"[page {index}]\n{page_text}")
            page_count = len(pdf.pages)
        return ("\n\n".join(pages), "pdfplumber", {"extension": ".pdf", "page_count": page_count})

    def _extract_docx_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to ingest DOCX documents") from exc
        doc = Document(str(path))
        parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return ("\n".join(parts), "python-docx", {"extension": ".docx"})

    def _extract_excel_text(
        self, path: _facade().Path
    ) -> tuple[str, str, dict[str, _facade().Any]]:
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            return self._extract_xlsx_text(path)
        if suffix == ".xls":
            return self._extract_xls_text(path)
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")

    def _extract_xlsx_text(self, path: _facade().Path) -> tuple[str, str, dict[str, _facade().Any]]:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to ingest XLSX documents") from exc
        max_rows_per_sheet = 5000
        max_cols = 100
        parts: list[str] = []
        sheet_count = 0
        row_count = 0
        wb = load_workbook(str(path), read_only=True, data_only=True)
        try:
            for sheet in wb.worksheets:
                sheet_count += 1
                sheet_lines: list[str] = []
                for index, row in enumerate(sheet.iter_rows(values_only=True)):
                    if index >= max_rows_per_sheet:
                        sheet_lines.append(f"... truncated after {max_rows_per_sheet} rows")
                        break
                    cells: list[str] = []
                    for col_index, cell in enumerate(row or ()):
                        if col_index >= max_cols:
                            break
                        if cell is None:
                            continue
                        text = str(cell).strip()
                        if text:
                            cells.append(text)
                    if cells:
                        sheet_lines.append("\t".join(cells))
                        row_count += 1
                if sheet_lines:
                    parts.append(f"[sheet {sheet.title}]\n" + "\n".join(sheet_lines))
        finally:
            wb.close()
        return (
            "\n\n".join(parts),
            "openpyxl",
            {"extension": ".xlsx", "sheet_count": sheet_count, "row_count": row_count},
        )
