"""
价格表生成器

生成客户价格表 PDF 文档并支持打印
"""

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, cast

from app.utils.operational_errors import RECOVERABLE_ERRORS
from app.utils.path_utils import resolve_fhd_repo_root

logger = logging.getLogger(__name__)


class PriceListGenerator:
    """价格表生成器"""

    def __init__(self, output_dir: str | None = None):
        if output_dir is None:
            output_dir = str(Path(__file__).parent.parent.parent / "generated_price_lists")

        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info("[PriceListGenerator] 输出目录：%s", self.output_dir)

    def generate(
        self, customer_name: str, products: list[dict[str, Any]], printer_name: str | None = None
    ) -> dict[str, Any]:
        """
        生成价格表

        Args:
            customer_name: 客户名称
            products: 产品列表，每个产品包含 model_number, name, spec, unit, unit_price
            printer_name: 打印机名称（可选）

        Returns:
            生成结果，包含 success, filename, filepath, error 等字段
        """
        try:
            logger.info("[PriceListGenerator] 开始生成价格表 - 客户：%s", customer_name)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = customer_name.replace("/", "_").replace("\\", "_")
            filename = f"{safe_name}_价格表_{timestamp}.docx"
            filepath = self.output_dir / filename

            # 生成 PDF 内容
            self._create_price_list_pdf(filepath, customer_name, products)

            logger.info("[PriceListGenerator] PDF 生成成功：%s", filepath)

            # 如果指定了打印机，执行打印
            if printer_name:
                logger.info("[PriceListGenerator] 发送到打印机：%s", printer_name)
                self._print_file(str(filepath), printer_name)
            else:
                # 自动选择打印机
                auto_printer = self._get_default_printer()
                if auto_printer:
                    logger.info("[PriceListGenerator] 使用默认打印机：%s", auto_printer)
                    self._print_file(str(filepath), auto_printer)
                else:
                    logger.warning("[PriceListGenerator] 未找到可用打印机，仅生成文件")

            return {
                "success": True,
                "filename": filename,
                "filepath": str(filepath),
                "message": "价格表已生成",
            }

        except RECOVERABLE_ERRORS as e:
            logger.error("[PriceListGenerator] 生成失败：%s", e, exc_info=True)
            return {"success": False, "message": str(e)}

    def _create_price_list_pdf(
        self, filepath: Path, customer_name: str, products: list[dict[str, Any]]
    ):
        """
        创建价格表 Word 文件

        使用模板生成 Word 文档
        """
        try:
            # 尝试使用模板生成 Word
            try:
                fhd_root = resolve_fhd_repo_root(anchor=Path(__file__).resolve())
                if not fhd_root:
                    raise RuntimeError("未解析到 FHD 仓库根目录（需存在 app 包）")

                from app.infrastructure.documents.price_list_export import (
                    build_price_list_docx_bytes,
                )

                template_path = fhd_root / "424" / "模板.docx"
                if not template_path.exists():
                    raise FileNotFoundError(f"模板文件不存在: {template_path}")

                # 使用模板生成文档
                docx_bytes = build_price_list_docx_bytes(
                    template_path=template_path,
                    customer_name=customer_name,
                    quote_date=datetime.now().strftime("%Y-%m-%d"),
                    products=products,
                )

                # 保存到文件
                with open(filepath, "wb") as f:
                    f.write(docx_bytes)

                logger.info("[PriceListGenerator] 使用模板生成 Word 成功：%s", filepath)
                return

            except ImportError as e:
                logger.warning(
                    "[PriceListGenerator] 导入模板生成失败：%s，使用 python-docx 直接生成", e
                )
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.shared import Pt

                # 创建 Word 文档
                doc = Document()

                # 设置中文字体
                doc.styles["Normal"].font.name = "SimSun"
                doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

                # 标题
                title = doc.add_heading(f"{customer_name}价格表", level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0]
                title_run.font.name = "SimSun"
                title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                title_run.font.size = Pt(16)
                title_run.bold = True

                # 日期
                date_str = datetime.now().strftime("%Y年%m月%d日")
                date_para = doc.add_paragraph(f"生成日期：{date_str}")
                date_run = date_para.runs[0]
                date_run.font.name = "SimSun"
                date_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                date_run.font.size = Pt(10)

                doc.add_paragraph()  # 空行

                # 产品表格
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"

                # 表头
                header_cells = table.rows[0].cells
                headers = ["序号", "型号", "产品名称", "规格", "单位", "单价 (元)"]
                for i, header_text in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header_text
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(10)
                    run.bold = True

                # 添加产品数据
                for idx, product in enumerate(products, 1):
                    # 支持字典和 Pydantic 模型
                    if isinstance(product, dict):
                        model_number = product.get("model_number", "")
                        name = product.get("name", "")
                        spec = product.get("spec", "")
                        unit = product.get("unit", "")
                        unit_price = product.get("unit_price", "")
                    else:
                        model_number = getattr(product, "model_number", "")
                        name = getattr(product, "name", "")
                        spec = getattr(product, "spec", "")
                        unit = getattr(product, "unit", "")
                        unit_price = getattr(product, "unit_price", "")

                    row_cells = table.add_row().cells
                    row_data = [str(idx), model_number, name, spec, unit, unit_price]

                    for i, text in enumerate(row_data):
                        cell = row_cells[i]
                        cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cell.paragraphs[0].runs[0]
                        run.font.name = "SimSun"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                        run.font.size = Pt(10)

                # 保存 Word 文件
                doc.save(str(filepath))
                logger.info("[PriceListGenerator] 使用 python-docx 生成 Word 成功")

            except ImportError:
                # python-docx 不可用，生成文本文件
                logger.warning("[PriceListGenerator] python-docx 不可用，生成文本文件")
                txt_filepath = filepath.with_suffix(".txt")
                with open(txt_filepath, "w", encoding="utf-8") as f:
                    f.write(f"{customer_name}价格表\n")
                    f.write("=" * 80 + "\n")
                    f.write(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}\n\n")

                    f.write(
                        f"{'序号':<6}{'型号':<15}{'产品名称':<30}{'规格':<15}{'单位':<8}{'单价 (元)':<10}\n"
                    )
                    f.write("-" * 84 + "\n")

                    for idx, product in enumerate(products, 1):
                        # 支持字典和 Pydantic 模型
                        if isinstance(product, dict):
                            model_number = product.get("model_number", "")
                            name = product.get("name", "")
                            spec = product.get("spec", "")
                            unit = product.get("unit", "")
                            unit_price = product.get("unit_price", "")
                        else:
                            model_number = getattr(product, "model_number", "")
                            name = getattr(product, "name", "")
                            spec = getattr(product, "spec", "")
                            unit = getattr(product, "unit", "")
                            unit_price = getattr(product, "unit_price", "")

                        f.write(
                            f"{idx:<6}"
                            f"{model_number:<15}"
                            f"{name:<30}"
                            f"{spec:<15}"
                            f"{unit:<8}"
                            f"{unit_price:<10}\n"
                        )

                    f.write("\n" + "=" * 80 + "\n")

                logger.info("[PriceListGenerator] 文本文件生成成功：%s", txt_filepath)

        except RECOVERABLE_ERRORS as e:
            logger.error("[PriceListGenerator] 创建 Word 失败：%s", e, exc_info=True)
            raise

    def _get_default_printer(self) -> str | None:
        """获取默认打印机名称"""
        try:
            if os.name == "nt":  # Windows
                import win32print

                default_printer = win32print.GetDefaultPrinter()
                return cast("str | None", default_printer)
        except RECOVERABLE_ERRORS as e:
            logger.warning("[PriceListGenerator] 获取默认打印机失败：%s", e)
        return None

    def _print_file(self, filepath: str, printer_name: str):
        """
        打印文件（支持 Word 和文本文件）

        Args:
            filepath: 文件路径
            printer_name: 打印机名称
        """
        try:
            if os.name == "nt":  # Windows
                import win32api
                import win32print

                # 验证打印机是否存在
                printers = [
                    p[2]
                    for p in win32print.EnumPrinters(
                        win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS
                    )
                ]

                if printer_name not in printers:
                    logger.warning(
                        "[PriceListGenerator] 打印机 %s 不存在，使用默认打印机", printer_name
                    )
                    printer_name = win32print.GetDefaultPrinter()

                # 使用 Word 打印（如果安装了 Word）
                file_ext = Path(filepath).suffix.lower()
                if file_ext == ".docx" or file_ext == ".doc":
                    try:
                        # 尝试使用 Word COM 对象打印
                        import win32com.client

                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(filepath)
                        doc.PrintOut()
                        doc.Close()
                        word.Quit()
                        logger.info("[PriceListGenerator] 使用 Word COM 打印成功：%s", printer_name)
                        return
                    except RECOVERABLE_ERRORS as e:
                        logger.warning(
                            "[PriceListGenerator] Word COM 打印失败：%s，使用 Shell 打印", e
                        )

                # 使用 Shell 打印（Windows 默认方式）
                win32api.ShellExecute(0, "print", filepath, f'/d:"{printer_name}"', ".", 0)

                logger.info("[PriceListGenerator] 文件已发送到打印机：%s", printer_name)
            else:
                logger.warning("[PriceListGenerator] 非 Windows 系统，跳过打印")

        except RECOVERABLE_ERRORS as e:
            logger.error("[PriceListGenerator] 打印失败：%s", e, exc_info=True)
            raise
