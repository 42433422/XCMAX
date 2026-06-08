import os
import uuid
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class SimpleSalesContractGenerator:
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "generated_contracts"
        )
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(
        self,
        customer_name: str,
        contract_date: str = None,
        products: list[dict] = None,
        return_buckets_expected: int = None,
        return_buckets_actual: int = None,
    ) -> dict:
        if contract_date is None:
            now = datetime.now()
            contract_date = f"{now.year}年{now.month:02d}月{now.day:02d}日"

        if products is None:
            products = [
                {
                    "model_number": "306B",
                    "name": "PU亮光硬化剂",
                    "spec": "10KG×1",
                    "unit": "桶",
                    "quantity": "10",
                    "unit_price": "39.2",
                    "amount": "392",
                }
            ]

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("销售合同")
        run.font.size = Pt(18)
        run.font.bold = True

        doc.add_paragraph()

        info_para = doc.add_paragraph()
        info_para.add_run(f"客户：{customer_name}").font.size = Pt(12)
        info_para.add_run(f"\n日期：{contract_date}").font.size = Pt(12)

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        headers = ["编号", "品名", "规格", "单位", "数量", "单价", "金额"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_amount = 0
        total_quantity = 0

        for product in products:
            row_cells = table.add_row().cells
            model_no = product.get("model_number", "")
            name = product.get("name", "")
            product_name = f"{model_no}\n{name}" if model_no else name

            row_cells[0].text = model_no
            row_cells[1].text = product_name
            row_cells[2].text = product.get("spec", "")
            row_cells[3].text = product.get("unit", "")
            row_cells[4].text = product.get("quantity", "")
            row_cells[5].text = f"{product.get('unit_price', '')}元/KG"
            row_cells[6].text = f"{product.get('amount', '')}元"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            qty = float(product.get("quantity", 0))
            amt = float(product.get("amount", 0))
            total_quantity += qty
            total_amount += amt

        doc.add_paragraph()

        total_para = doc.add_paragraph()
        total_para.add_run(f"总重量：{total_quantity} KG").font.size = Pt(11)
        total_para.add_run(f"        合计金额：{total_amount} 元").font.size = Pt(11)

        doc.add_paragraph()

        note_para = doc.add_paragraph()
        note_para.add_run("注：以上价格均为实价，无折扣！").font.size = Pt(10)

        note2_para = doc.add_paragraph()
        return_expected_str = (
            f"({return_buckets_expected})" if return_buckets_expected is not None else "(  )"
        )
        return_actual_str = (
            f"({return_buckets_actual})" if return_buckets_actual is not None else "(  )"
        )
        note2_para.add_run(
            f"注：160KG桶需退回，如未退回一个按80元计算 应退桶{return_expected_str}个，实退桶{return_actual_str}个."
        ).font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_para.add_run(
            "核準：黄种霜    會計：胡小玲    經辨：姚胜华    倉庫：廖振卷"
        ).font.size = Pt(10)

        contract_id = str(uuid.uuid4())[:8]
        filename = f"{customer_name}_{contract_id}_销售合同.docx"
        filename = filename.replace(" ", "_").replace("/", "_")
        file_path = os.path.join(self.output_dir, filename)

        doc.save(file_path)

        return {
            "success": True,
            "contract_id": contract_id,
            "filename": filename,
            "file_path": file_path,
            "customer_name": customer_name,
            "contract_date": contract_date,
            "products": products,
            "total_quantity": total_quantity,
            "total_amount": total_amount,
        }


if __name__ == "__main__":
    generator = SimpleSalesContractGenerator()

    result = generator.generate(
        customer_name="深圳市百木鼎家具有限公司",
        contract_date="2026年04月11日",
        products=[
            {
                "model_number": "306B",
                "name": "PU亮光硬化剂",
                "spec": "10KG×1",
                "unit": "桶",
                "quantity": "10",
                "unit_price": "39.2",
                "amount": "392",
            }
        ],
        return_buckets_expected=1,
        return_buckets_actual=0,
    )

    print("=== 简洁版销售合同生成 ===")
    print(f"状态: {'成功' if result['success'] else '失败'}")
    print(f"文件名: {result['filename']}")
    print(f"文件路径: {result['file_path']}")
    print(f"客户: {result['customer_name']}")
    print(f"日期: {result['contract_date']}")
    print(f"总金额: {result['total_amount']}元")
