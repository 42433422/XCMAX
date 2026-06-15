"""Draft document structure via LLM, render to DOCX / XLSX."""

from __future__ import annotations

import json
import logging
import os
import re
from io import BytesIO
from typing import Any, Literal, cast

from app.infrastructure.llm.client import (
    get_openai_compatible_client,
    resolve_chat_model,
    resolve_mode,
)

logger = logging.getLogger(__name__)


def _document_spec_timeout_sec() -> float:
    """起草合同/表格 JSON 的 LLM 调用超时（秒）；过长会阻塞 Planner 流式 SSE。"""
    raw = (os.environ.get("FHD_DOCUMENT_SPEC_TIMEOUT_SEC") or "180").strip()
    try:
        v = float(raw)
    except ValueError:
        v = 180.0
    return max(15.0, min(v, 600.0))


_SYSTEM_DOCX = """你是中文商务与法务文书助手。用户会描述要生成的 Word 文档（如合同、协议、服务方案、告知书等）。
请输出**唯一一个** JSON 对象（不要 markdown 代码块，不要任何解释文字），格式严格如下:
{
  "title": "文档标题",
  "sections": [
    { "heading": "一、合同双方", "paragraphs": ["甲方（服务方）：[公司名称]，地址：[地址]，联系人：[姓名]，电话：[电话]。", "乙方（委托方）：[公司名称]，地址：[地址]，联系人：[姓名]，电话：[电话]。"] },
    { "heading": "二、服务内容与范围", "paragraphs": ["1. 乙方向甲方购买以下AI服务：[具体服务内容]。", "2. 服务范围包括：[详细范围]。"] },
    { "heading": "三、合同金额与支付方式", "paragraphs": ["1. 合同总金额：人民币 [金额] 元（大写：[大写金额]）。", "2. 支付方式：[一次性支付/分期支付]。", "3. 付款时间：[具体时间节点]。"] },
    { "heading": "四、交付与验收", "paragraphs": ["1. 交付时间：自本合同签署之日起 [X] 个工作日内完成。", "2. 验收标准：乙方应在收到服务成果后 [X] 个工作日内完成验收。"] },
    { "heading": "五、知识产权与保密", "paragraphs": ["1. 双方确认，本合同项下产生的知识产权归 [甲方/乙方/双方共有]。", "2. 双方应对在工作过程中知悉的对方商业秘密负有保密义务。"] },
    { "heading": "六、违约责任", "paragraphs": ["1. 任何一方违反本合同约定，应向守约方支付合同总金额 [X]% 的违约金。", "2. 因一方违约造成对方损失的，违约方应赔偿守约方因此遭受的全部损失。"] },
    { "heading": "七、争议解决", "paragraphs": ["本合同的履行过程中如发生争议，双方应友好协商解决；协商不成的，提交 [仲裁机构/法院] 解决。"] },
    { "heading": "八、其他约定", "paragraphs": ["1. 本合同自双方签字（盖章）之日起生效。", "2. 本合同一式 [X] 份，甲乙双方各执 [X] 份，具有同等法律效力。", "3. 本合同未尽事宜，由双方另行签订补充协议，补充协议与本合同具有同等法律效力。"] }
  ],
  "tables": [],
  "signatures": ["甲方（盖章）：________________    乙方（盖章）：________________", "签署日期：____年____月____日        签署日期：____年____月____日"]
}
要求：内容须贴合用户描述的主题；若用户要「AI 服务合同 / 技术服务合同」，应包含服务内容、数据与知识产权、保密、费用与支付、交付与验收、期限、变更与终止、违约、不可抗力、争议解决等常见可协商条款（表述为中文法律商务常用语，非正式法律意见）；段落用 string 数组；无表格时 tables 为 []。"""

_SYSTEM_XLSX = """你是中文办公表格助手。用户会描述要的 Excel。请输出**唯一一个** JSON 对象（不要 markdown）:
{
  "title": "用于文件名的简短标题",
  "sheets": [
    {
      "name": "工作表名",
      "headers": ["列1","列2"],
      "rows": [["示例1","示例2"]],
      "column_widths": [18, 12]
    }
  ]
}
column_widths 可选。根据需求生成合理业务数据（报价、清单、排期、对账等）。至少一个 sheet。"""


def _strip_json_fence(text: str) -> str:
    t = (text or "").strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z0-9_-]*\n?", "", t)
        t = re.sub(r"\n?```\s*$", "", t)
    return t.strip()


def _extract_first_json_object(text: str) -> str | None:
    """从模型输出中抠出第一个顶层 JSON 对象（忽略前后说明、尾随文字）。"""
    t = _strip_json_fence((text or "").strip())
    start = t.find("{")
    if start < 0:
        return None
    depth = 0
    in_str = False
    esc = False
    for i in range(start, len(t)):
        ch = t[i]
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return t[start : i + 1]
    return None


def draft_document_spec(user_prompt: str, output: Literal["docx", "xlsx"]) -> dict[str, Any]:
    if resolve_mode() == "offline":
        raise RuntimeError(
            "当前为 offline 模式：生成文档需要 online 模式并配置 DEEPSEEK_API_KEY 或 OPENAI_API_KEY"
        )
    cli = get_openai_compatible_client()
    model = resolve_chat_model()
    sys = _SYSTEM_DOCX if output == "docx" else _SYSTEM_XLSX
    resp = cli.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": sys},
            {"role": "user", "content": (user_prompt or "").strip()},
        ],
        temperature=0.35,
        max_tokens=8192,
        timeout=_document_spec_timeout_sec(),
        response_format={"type": "json_object"},
    )
    raw = (resp.choices[0].message.content or "").strip()
    raw = _strip_json_fence(raw)
    try:
        return cast("dict[str, Any]", json.loads(raw))
    except json.JSONDecodeError:
        blob = _extract_first_json_object(raw)
        if blob:
            try:
                return cast("dict[str, Any]", json.loads(blob))
            except json.JSONDecodeError:
                logger.warning("document spec JSON parse failed (extracted), head=%s", blob[:400])
        else:
            logger.warning("document spec JSON parse failed, head=%s", raw[:400])
        raise RuntimeError("模型未返回合法 JSON，请缩短需求或改为分条说明后重试") from None


def build_docx_bytes(spec: dict[str, Any]) -> tuple[bytes, str]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
        from docx.shared import Inches, Pt, RGBColor  # noqa: F401
    except ImportError as e:
        raise RuntimeError("未安装 python-docx") from e

    title = str(spec.get("title") or "生成文档").strip() or "生成文档"
    doc = Document()
    doc.add_heading(title, level=0)
    for section in spec.get("sections") or []:
        if not isinstance(section, dict):
            continue
        head = str(section.get("heading") or "").strip()
        if head:
            h = doc.add_heading(head, level=1)
            h.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        for p in section.get("paragraphs") or []:
            s = str(p).strip()
            if s:
                para = doc.add_paragraph(s)
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_after = Pt(6)
    for tbl in spec.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        tt = str(tbl.get("title") or "").strip()
        if tt:
            doc.add_paragraph(tt)
        headers = tbl.get("headers") if isinstance(tbl.get("headers"), list) else []
        rows = tbl.get("rows") if isinstance(tbl.get("rows"), list) else []
        col_counts = [len(headers)]
        for r in rows:
            if isinstance(r, list):
                col_counts.append(len(r))
        n_col = max(col_counts + [1])
        n_row = 1 + len([r for r in rows if isinstance(r, list)])
        table = doc.add_table(rows=n_row, cols=n_col)
        table.style = "Table Grid"
        for j in range(n_col):
            h = headers[j] if j < len(headers) else ""
            table.rows[0].cells[j].text = str(h)
        for i, r in enumerate(rows, start=1):
            if not isinstance(r, list):
                continue
            for j in range(n_col):
                val = r[j] if j < len(r) else ""
                table.rows[i].cells[j].text = str(val)
    for line in spec.get("signatures") or []:
        s = str(line).strip()
        if s:
            para = doc.add_paragraph(s)
            para.paragraph_format.line_spacing = 2.0
    buf = BytesIO()
    doc.save(buf)
    fn = f"{title[:48]}.docx".replace("/", "-").replace("\\", "-")
    return buf.getvalue(), fn


def build_xlsx_bytes(spec: dict[str, Any]) -> tuple[bytes, str]:
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    title = str(spec.get("title") or "生成表格").strip() or "生成表格"
    wb = Workbook()
    sheets = spec.get("sheets") if isinstance(spec.get("sheets"), list) else []
    if not sheets:
        sheets = [{"name": "Sheet1", "headers": ["说明"], "rows": [["（无内容）"]]}]
    first = True
    for sh in sheets:
        if not isinstance(sh, dict):
            continue
        name = str(sh.get("name") or "Sheet1").strip()[:31] or "Sheet1"
        if first:
            ws = wb.active
            ws.title = name
            first = False
        else:
            ws = wb.create_sheet(title=name)
        headers = sh.get("headers") if isinstance(sh.get("headers"), list) else []
        rows = sh.get("rows") if isinstance(sh.get("rows"), list) else []
        r_i = 1
        if headers:
            for j, h in enumerate(headers, start=1):
                ws.cell(row=r_i, column=j, value=str(h))
            r_i += 1
        for row in rows:
            if not isinstance(row, list):
                continue
            for j, v in enumerate(row, start=1):
                ws.cell(row=r_i, column=j, value=v)
            r_i += 1
        widths = sh.get("column_widths")
        if isinstance(widths, list):
            for i, w in enumerate(widths[:40], start=1):
                try:
                    ws.column_dimensions[get_column_letter(i)].width = float(w)
                except (TypeError, ValueError):
                    pass
    buf = BytesIO()
    wb.save(buf)
    fn = f"{title[:48]}.xlsx".replace("/", "-").replace("\\", "-")
    return buf.getvalue(), fn


def generate_office_file(user_prompt: str, output: Literal["docx", "xlsx"]) -> tuple[bytes, str]:
    spec = draft_document_spec(user_prompt, output)
    if output == "docx":
        return build_docx_bytes(spec)
    return build_xlsx_bytes(spec)
