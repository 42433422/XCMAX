"""小猫分析 Word 报告导出（与 XLSX 导出共享插件结论）。"""

from __future__ import annotations

from app.utils.operational_errors import OPERATIONAL_ERRORS
import json
from datetime import datetime
from io import BytesIO
from typing import Any

from .service import KittenReportExportService


def _html_to_plain(content: str) -> str:
    return KittenReportExportService._html_to_text(str(content or ""))


def build_kitten_docx(payload: dict[str, Any]) -> dict[str, Any]:
    svc = KittenReportExportService()
    plugin_results = svc.collect_plugin_results(payload or {})

    dataset = payload.get("dataset") or {}
    messages = payload.get("messages") or []
    result = payload.get("result") or {}
    phase = str(payload.get("phase") or "")
    industry = str(payload.get("industry") or "通用")
    web_hits = payload.get("web_search_results") or []

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("python-docx 未安装，无法导出 Word") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    title = doc.add_heading("小猫分析报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"分析阶段：{phase or '-'}　行业：{industry}")

    doc.add_heading("结论摘要", level=1)
    doc.add_paragraph(str(result.get("title") or "AI 分析"), style="Heading 2")
    doc.add_paragraph(_html_to_plain(str(result.get("summary") or "")))

    doc.add_heading("数据与上传文件", level=1)
    if dataset:
        rows = int(dataset.get("rows") or 0)
        cols = int(dataset.get("columns") or 0)
        fn = str(dataset.get("name") or "")
        fields = dataset.get("fieldNames") or dataset.get("fields") or []
        field_txt = "、".join(str(x) for x in fields[:40])
        if len(fields) > 40:
            field_txt += "…"
        doc.add_paragraph(f"文件：{fn}")
        doc.add_paragraph(f"规模：{rows} 行 / {cols} 列")
        if field_txt:
            doc.add_paragraph(f"字段：{field_txt}")
        preview = str(dataset.get("previewText") or "").strip()
        if preview:
            doc.add_paragraph("样本预览：")
            doc.add_paragraph(preview[:8000] + ("…" if len(preview) > 8000 else ""))
    else:
        doc.add_paragraph("本次未附带上传表格。")

    if isinstance(web_hits, list) and web_hits:
        doc.add_heading("联网检索摘要", level=1)
        for idx, hit in enumerate(web_hits[:8], start=1):
            if not isinstance(hit, dict):
                continue
            t = str(hit.get("title") or "").strip()
            u = str(hit.get("url") or "").strip()
            s = str(hit.get("snippet") or "").strip()
            doc.add_paragraph(f"{idx}. {t or u}", style="Heading 3")
            if u:
                doc.add_paragraph(u)
            if s:
                doc.add_paragraph(s[:1200] + ("…" if len(s) > 1200 else ""))

    doc.add_heading("算法与插件洞察", level=1)
    for p in plugin_results:
        doc.add_paragraph(str(p.get("title") or p.get("key") or "插件"), style="Heading 3")
        doc.add_paragraph(f"级别：{p.get('level', '')}　摘要：{p.get('summary', '')}")
        details = p.get("details")
        if details:
            try:
                detail_txt = json.dumps(details, ensure_ascii=False, default=str)
            except OPERATIONAL_ERRORS:
                detail_txt = str(details)
            if len(detail_txt) > 6000:
                detail_txt = detail_txt[:6000] + "…"
            doc.add_paragraph(detail_txt)

    doc.add_heading("对话摘录", level=1)
    if messages:
        for msg in messages[-40:]:
            role = str(msg.get("role") or "")
            label = "助手" if role == "ai" else "用户"
            line = (
                f"[{label} {msg.get('time') or ''}] {_html_to_plain(str(msg.get('content') or ''))}"
            )
            doc.add_paragraph(line[:2000] + ("…" if len(line) > 2000 else ""))
    else:
        doc.add_paragraph("暂无对话记录。")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    file_name = f"小猫分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return {"file_name": file_name, "content": buf.read()}
