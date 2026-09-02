"""XC企业版 V10.0.0 文档鉴别材料（用户手册）生成器

输入：补正材料目录下的 screenshots/（16 张功能截图 + 10 张模块操作截图）
输出：V10-申请材料-<日期>/文档鉴别材料-用户手册.docx
规范：封面含软件名称+版本号+著作权人；每页页眉含"软件名称+版本号"（V9 补正教训）。
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

SOFTWARE_NAME = "XC企业版"
VERSION = "V10.0.0"
OWNER = "成都修茈科技有限公司"

BASE = os.path.dirname(os.path.abspath(__file__))
SHOT_DIR = os.path.join(BASE, "补正材料-XC企业版-20260704", "screenshots")
OUT_DIR = os.path.join(BASE, f"V10-申请材料-{date.today().strftime('%Y%m%d')}")

CHAPTERS: list[tuple[str, list[tuple[str, str]]]] = [
    ("一、登录与主界面", [
        ("01-login.png", "启动客户端后进入登录界面，输入账号密码完成身份验证；支持会话保持与安全退出。"),
        ("02-home-chat.png", "登录后进入主界面，左侧为功能导航，中央为智能对话工作区，可发起自然语言指令查询数据、生成任务。"),
    ]),
    ("二、智能对话", [
        ("02-home-chat.png", "在对话工作区输入自然语言，系统调用 AI 员工完成意图识别、数据查询与结果播报；支持多轮会话与员工副窗切换。"),
        ("13-ai-groups.png", "员工分组管理界面：按业务域维护 AI 员工分组、成员与调度关系，支持启停与能力查看。"),
    ]),
    ("三、业务对象管理", [
        ("03-products.png", "业务对象（产品资料）列表：维护产品、规格、价格、单位等基础资料，支持 Excel 导入与价格表导出。"),
        ("click-01-业务对象.png", "点击左侧菜单【业务对象】进入资料维护页，可新增、编辑、删除与批量导入基础资料。"),
    ]),
    ("四、业务单据与业务记录", [
        ("04-orders.png", "业务单据列表：按客户与状态筛选订单，查看出货单记录并导出。"),
        ("05-create-order.png", "新建订单：选择客户、录入商品明细与数量，保存后生成业务记录并进入履约流程。"),
        ("06-shipment-records.png", "出货记录页：汇总发货流水，支持按时间与客户维度查询、导出对账。"),
        ("click-03-业务单据.png", "点击【业务单据】进入单据管理，完成开单、审核、发货状态跟踪。"),
        ("click-04-业务记录.png", "点击【业务记录】查看全部业务流水，支持按类型与时间范围检索。"),
    ]),
    ("五、组织管理", [
        ("07-customers.png", "组织管理（客户/购买单位资料）：维护客户档案、联系方式与结算信息，与订单链路共享数据。"),
        ("click-02-组织管理.png", "点击【组织管理】进入客户资料维护，支持新增客户、分级与归属调整。"),
    ]),
    ("六、资源库", [
        ("click-05-资源库.png", "资源库管理：维护物料、库存、仓库库位、供应商及出入库信息，支撑单据履约。"),
    ]),
    ("七、数据来源授权", [
        ("11-data-sources.png", "数据来源授权：管理 AI 员工可读取的数据源，支持本地消息数据库等适配器的授权接入与撤销。"),
        ("click-06-数据来源.png", "点击【数据来源】进入授权列表，可查看连接状态、读取范围与最近同步时间。"),
    ]),
    ("八、模板与打印", [
        ("08-template-preview.png", "模板预览：对 Word、Excel、CSV、PDF、PPT 等办公模板进行输出预览。"),
        ("09-print.png", "打印输出：选择打印机与份数，完成单据与标签的实物打印。"),
        ("10-label-editor.png", "标签编辑器：可视化设计标签版式，绑定数据字段后保存至模板库。"),
        ("click-07-模板与打印.png", "点击【模板与打印】进入模板管理，完成模板新建、编辑与启停。"),
        ("click-08-打印机列表.png", "打印机列表：维护本机打印机配置与默认打印机。"),
        ("click-09-模板库.png", "模板库：集中管理历史模板，支持复制、版本回退与共享。"),
    ]),
    ("九、智能生态与应用市场", [
        ("12-mod-store.png", "应用市场（Mod 商店）：浏览、购买与安装业务扩展模块，安装后即时生效。"),
        ("13-ai-groups.png", "员工工作台：管理 AI 员工、流程、模块与能力库，配置服务器运行能力。"),
    ]),
    ("十、系统设置与工具", [
        ("14-settings.png", "系统设置：维护账号资料、模型服务、钱包套餐与安全选项。"),
        ("16-tools.png", "工具箱：提供数据备份、导入导出等辅助工具入口。"),
        ("click-10-系统设置.png", "点击【系统设置】进入设置页，按类目完成参数配置。"),
    ]),
    ("十一、桌面运行时", [
        ("15-desktop-runtime.png", "桌面运行时：查看本地后端服务状态、SQLite 数据目录与端口，支持服务重启与备份恢复。"),
    ]),
]


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SOFTWARE_NAME)
    r.font.size = Pt(36)
    r.bold = True
    r.font.name = "Calibri"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{VERSION}  用户手册（文档鉴别材料）")
    r2.font.size = Pt(20)

    for _ in range(6):
        doc.add_paragraph()
    for line in (f"著作权人：{OWNER}", f"编写日期：{date.today().strftime('%Y年%m月%d日')}"):
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(line)
        r3.font.size = Pt(14)
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    doc.add_heading("软件概述", level=1)
    doc.add_paragraph(
        f"{SOFTWARE_NAME}（{VERSION}）是面向中小企业经营场景的企业 AI 员工平台，"
        "由前端 SPA、桌面壳（Electron）、本地后端服务与 SQLite 数据库构成，"
        "并通过模块（Mod）扩展机制与移动端实现协同。"
    )
    funcs = [
        "智能对话：自然语言入口，查询数据、生成任务、调用员工副窗并接入语音播报；",
        "业务对象管理：维护产品、规格、价格、单位等基础资料，支持 Excel 导入与模板化输出；",
        "组织管理：维护客户/购买单位资料，与订单、发货记录等业务链路共享数据；",
        "业务单据与业务记录：新建订单、查看出货记录、按客户与状态筛选、导出；",
        "资源库：物料、库存、仓库库位、供应商及出入库管理；",
        "数据来源授权：管理 AI 员工可读取的数据来源，支持适配器授权接入；",
        "模板与打印：标签生成、打印机配置、模板库管理，覆盖 Word/Excel/CSV/PDF/PPT；",
        "智能生态：AI 员工分组调度、Mod 应用市场、能力库与服务器运行能力；",
        "系统设置与桌面运行时：账号资料、模型服务、钱包套餐、本地后端与数据目录管理；",
        "移动协同：移动端会话、名录、探索与个人设置，与桌面端绑定协作。",
    ]
    for f in funcs:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("运行环境", level=1)
    env = [
        "操作系统：Windows 10 及以上 / macOS 12 及以上；",
        "运行时：客户端内置本地后端服务（Python）与 SQLite 数据库，无需额外安装数据库；",
        "硬件：内存 8GB 及以上，磁盘可用空间 2GB 及以上；",
        "网络：首次激活与模型服务、应用市场需要互联网连接，本地业务可离线运行。",
    ]
    for e in env:
        doc.add_paragraph(e, style="List Bullet")
    doc.add_page_break()


def main() -> None:
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    set_base_style(doc)
    add_cover(doc)

    doc.add_heading("目录", level=1)
    toc = ["软件概述", "运行环境"] + [c[0] for c in CHAPTERS]
    for i, t in enumerate(toc, 1):
        doc.add_paragraph(f"{i}. {t}")
    doc.add_page_break()

    add_overview(doc)

    for title, shots in CHAPTERS:
        doc.add_heading(title, level=1)
        for img_name, caption in shots:
            img_path = os.path.join(SHOT_DIR, img_name)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(img_path):
                p.add_run().add_picture(img_path, width=Cm(14))
            else:
                p.add_run(f"【缺图：{img_name}】")
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(caption)
            cr.font.size = Pt(10)
            cr.italic = True

    doc.add_heading("附：技术特点", level=1)
    for t in (
        "前后端分离架构：Vue 3 + TypeScript 前端与 Python 后端服务通过 HTTP/IPC 通信；",
        "桌面壳：Electron 封装，集成自动更新、签名校验与异常回滚；",
        "数据安全：SQLite WAL 模式、在线热备份与启动自检恢复；",
        "扩展机制：Mod 模块化扩展，支持市场分发与本地加载；",
        "AI 能力：意图识别、员工调度与多模型服务接入。",
    ):
        doc.add_paragraph(t, style="List Bullet")

    # 页眉：软件名称 + 版本号（V9 补正教训：全文档统一）
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.text = f"{SOFTWARE_NAME} {VERSION}  用户手册（文档鉴别材料）"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)

    out = os.path.join(OUT_DIR, "文档鉴别材料-用户手册.docx")
    doc.save(out)
    print(f"[完成] {out}")
    print(f"[完成] 章节 {len(CHAPTERS)} 个，截图引用 {sum(len(s) for _, s in CHAPTERS)} 处")


if __name__ == "__main__":
    main()
