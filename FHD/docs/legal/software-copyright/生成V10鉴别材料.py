"""XC企业版 V10.0.0 软著登记材料生成器（CPCC 申请专用）

产出（写入本目录 V10-申请材料-20260901/）：
1. 程序鉴别材料-完整源代码.txt   全量去注释源代码
2. 程序鉴别材料-前后30页.txt    前 1500 行 + 后 1500 行（每页 50 行）
3. 程序鉴别材料.docx            排版版（页眉含软件名称+版本号，9pt 密排）
4. 源代码统计报告.md            V10 口径统计

来源目录（相对 FHD 仓库根）：
- app/            Python 后端（FastAPI）
- frontend/src/   Vue 3 + TypeScript 前端
- desktop/        Electron 桌面端主进程

排除：测试、构建产物、依赖、数据文件。
"""

from __future__ import annotations

import os
import re
from datetime import date

SOFTWARE_NAME = "XC企业版"
VERSION = "V10.0.0"
LINES_PER_PAGE = 50
FRONT_PAGES = 30
BACK_PAGES = 30

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
# BASE_DIR = FHD/（脚本位于 FHD/docs/legal/software-copyright/）
OUT_DIR = os.path.dirname(os.path.abspath(__file__))

SOURCES = [
    ("后端(app)", os.path.join(BASE_DIR, "app")),
    ("前端(frontend/src)", os.path.join(BASE_DIR, "frontend", "src")),
    ("桌面端(desktop)", os.path.join(BASE_DIR, "desktop")),
]

EXCLUDE_DIRS = {
    "tests", "test", "__pycache__", "node_modules", "dist", "dist-electron",
    "release", "build", "coverage", "e2e", "test-results", "playwright-report",
    ".venv", "venv", "mock", "fixtures", "resources", "branding", "scripts",
    "templates", "vue-dist", ".vite", "stories",
}

EXTENSIONS = {".py", ".vue", ".ts", ".js", ".tsx", ".jsx"}

EXCLUDE_FILE_PATTERNS = re.compile(r"(\.test\.|\.spec\.|\.d\.ts$|\.config\.(ts|js)$|\.backup$)")


def remove_comments_python(content: str) -> str:
    lines = content.split("\n")
    result: list[str] = []
    in_multiline = False
    opener = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if in_multiline:
            if opener in stripped:
                in_multiline = False
            continue
        if stripped.startswith('"""') or stripped.startswith("'''"):
            opener = stripped[:3]
            if stripped.count(opener) >= 2 and len(stripped) > 3:
                continue
            in_multiline = True
            continue
        if "#" in line:
            code = line.split("#")[0]
            if code.strip():
                result.append(code.rstrip())
        else:
            result.append(line.rstrip())
    return "\n".join(result)


def remove_comments_c_like(content: str) -> str:
    content = re.sub(r"<!--.*?-->", "", content, flags=re.DOTALL)
    content = re.sub(r"/\*.*?\*/", "", content, flags=re.DOTALL)
    content = re.sub(r"^\s*//.*$", "", content, flags=re.MULTILINE)
    return content


def clean_lines(content: str) -> list[str]:
    content = content.replace("\t", "    ")
    lines = [ln.rstrip() for ln in content.split("\n")]
    return [ln for ln in lines if ln.strip()]


def collect(source_label: str, source_dir: str) -> tuple[list[str], int]:
    """返回 (带文件头注释的代码行列表, 文件数)。目录内按路径稳定排序。"""
    out: list[str] = []
    file_count = 0
    if not os.path.isdir(source_dir):
        print(f"[跳过] 目录不存在：{source_dir}")
        return out, 0
    for root, dirs, files in os.walk(source_dir):
        dirs[:] = sorted(d for d in dirs if d not in EXCLUDE_DIRS and not d.startswith("."))
        for name in sorted(files):
            ext = os.path.splitext(name)[1].lower()
            if ext not in EXTENSIONS or EXCLUDE_FILE_PATTERNS.search(name):
                continue
            path = os.path.join(root, name)
            try:
                with open(path, encoding="utf-8", errors="ignore") as f:
                    raw = f.read()
            except OSError:
                continue
            if ext == ".py":
                body = clean_lines(remove_comments_python(raw))
            else:
                body = clean_lines(remove_comments_c_like(raw))
            if not body:
                continue
            rel = os.path.relpath(path, source_dir)
            out.append(f"// ===== [{source_label}] {rel} =====")
            out.extend(body)
            file_count += 1
    return out, file_count


def write_txt(path: str, lines: list[str]) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        f.write("\n")


def build_front_back_30(all_lines: list[str], out_path: str) -> tuple[int, int]:
    """前 30 页 + 后 30 页，中间省略处有明确标注。返回 (前取行数, 后取行数)。"""
    n = LINES_PER_PAGE * FRONT_PAGES
    m = LINES_PER_PAGE * BACK_PAGES
    head = all_lines[:n]
    tail = all_lines[-m:] if len(all_lines) > n + m else []
    omitted = len(all_lines) - len(head) - len(tail)
    out = list(head)
    if tail:
        out.append(f"// ……（中间源代码省略 {omitted} 行，全文见《程序鉴别材料-完整源代码》）……")
        out.extend(tail)
    write_txt(out_path, out)
    return len(head), len(tail)


def build_docx(all_lines: list[str], out_path: str, total_pages: int) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.6)
        sec.bottom_margin = Cm(1.6)
        sec.left_margin = Cm(1.9)
        sec.right_margin = Cm(1.9)
        header_p = sec.header.paragraphs[0]
        header_p.text = f"软件名称：{SOFTWARE_NAME} {VERSION}    程序鉴别材料    第 页"
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_p.runs:
            run.font.size = Pt(9)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(11.5)

    batch = 400
    for i in range(0, len(all_lines), batch):
        chunk = all_lines[i : i + batch]
        p = doc.add_paragraph("\n".join(chunk))
        p.paragraph_format.line_spacing = Pt(11.5)

    # 页码域：替换页眉中"第 页"为"第 X 页"
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            for run in p.runs:
                if "第 页" in run.text:
                    run.text = run.text.replace("第 页", "第 ")
                    fld = run._r.makeelement(qn("w:fldSimple"), {qn("w:instr"): "PAGE"})
                    run._r.addnext(fld)
                    tail = run._r.makeelement(qn("w:r"), {})
                    t = tail.makeelement(qn("w:t"), {})
                    t.text = " 页"
                    tail.append(t)
                    fld.addnext(tail)
    doc.save(out_path)
    print(f"已生成 docx（约 {total_pages} 页）：{os.path.basename(out_path)}")


def main() -> None:
    stamp = date.today().strftime("%Y%m%d")
    out = os.path.join(OUT_DIR, f"V10-申请材料-{stamp}")
    os.makedirs(out, exist_ok=True)

    all_lines: list[str] = []
    report: list[tuple[str, int, int]] = []
    for label, src in SOURCES:
        lines, cnt = collect(label, src)
        report.append((label, cnt, len(lines)))
        all_lines.extend(lines)
        print(f"[完成] {label}: {cnt} 个文件, {len(lines)} 行")

    full = os.path.join(out, "程序鉴别材料-完整源代码.txt")
    fb = os.path.join(out, "程序鉴别材料-前后30页.txt")
    write_txt(full, all_lines)
    head_n, tail_n = build_front_back_30(all_lines, fb)

    total_pages = (len(all_lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    docx_path = os.path.join(out, "程序鉴别材料.docx")
    try:
        build_docx(all_lines, docx_path, total_pages)
    except Exception as exc:  # noqa: BLE001 — docx 失败不阻塞 txt 产出
        print(f"[警告] docx 生成失败（可用 txt 自行排版）：{exc}")

    report_path = os.path.join(out, "源代码统计报告.md")
    total = len(all_lines)
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(f"# {SOFTWARE_NAME} - 源代码统计报告\n\n")
        f.write("## 一、基本信息\n\n")
        f.write(f"- **软件名称**：{SOFTWARE_NAME}\n- **版本号**：{VERSION}\n")
        f.write(f"- **著作权人**：成都修茈科技有限公司\n- **材料用途**：计算机软件著作权登记（程序鉴别材料）\n")
        f.write(f"- **生成日期**：{date.today().strftime('%Y年%m月%d日')}\n\n")
        f.write("## 二、源程序构成\n\n| 模块 | 文件数 | 去注释后行数 |\n|---|---:|---:|\n")
        for label, cnt, ln in report:
            f.write(f"| {label} | {cnt} | {ln:,} |\n")
        f.write(f"| **合计** | **{sum(c for _, c, _ in report)}** | **{total:,}** |\n\n")
        f.write("## 三、鉴别材料口径\n\n")
        f.write(f"- 每页 {LINES_PER_PAGE} 行；源程序共约 {total_pages:,} 页\n")
        f.write(f"- 程序鉴别材料提交前后各 {FRONT_PAGES} 页（前 {head_n:,} 行 + 后 {tail_n:,} 行），全文备查于完整源代码文档\n")
        f.write("- 已剔除：注释、空行、测试代码、构建产物、第三方依赖\n")
    print(f"\n[完成] 统计报告：{report_path}")
    print(f"[完成] 程序总量：{total:,} 行 ≈ {total_pages:,} 页（每页 {LINES_PER_PAGE} 行）")
    print(f"[完成] 输出目录：{out}")


if __name__ == "__main__":
    main()
