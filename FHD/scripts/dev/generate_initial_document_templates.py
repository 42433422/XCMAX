#!/usr/bin/env python3
"""生成模板库通用 Excel（按 templateScopeRules 必填列）+ 发货单/价格表。"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _style_header(ws, cols: int) -> None:
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )
    for c in range(1, cols + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin
    ws.row_dimensions[1].height = 22


def _write_sheet(
    path: Path,
    *,
    sheet_title: str,
    headers: list[str],
    rows: list[list[Any]],
    title: str | None = None,
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_title[:31]
    start = 1
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        cell = ws.cell(row=1, column=1, value=title)
        cell.font = Font(name="微软雅黑", size=14, bold=True, color="1F4E79")
        cell.alignment = Alignment(horizontal="center")
        start = 2
    for i, h in enumerate(headers, 1):
        ws.cell(row=start, column=i, value=h)
    # header style on start row
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=start, column=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for r_i, row in enumerate(rows, start=start + 1):
        for c_i, val in enumerate(row, 1):
            ws.cell(row=r_i, column=c_i, value=val)
    for i, h in enumerate(headers, 1):
        ws.column_dimensions[chr(64 + i) if i <= 26 else "A"].width = max(12, min(22, len(h) * 2 + 4))
    # simpler width
    from openpyxl.utils import get_column_letter

    for i, h in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(i)].width = max(12, min(24, len(h) * 2 + 4))
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


# 与 frontend/src/shared/templateScopeRules.json / backend _DEFAULT_TEMPLATE_SCOPE_RULES 对齐
GENERIC_EXCEL_SPECS: list[dict[str, Any]] = [
    {
        "filename": "通用_出货明细.xlsx",
        "template_key": "SEED_ORDERS_DEFAULT",
        "template_name": "通用出货明细表",
        "template_type": "出货明细",
        "business_scope": "orders",
        "sheet": "出货明细",
        "title": "XCMAX 通用出货明细表",
        "headers": ["产品型号", "产品名称", "数量", "单价", "金额", "备注"],
        "rows": [
            ["A001", "演示产品A", 3, 12.5, 37.5, "开箱样例"],
            ["9803", "演示哑光清面漆", 2, 18.8, 37.6, ""],
        ],
    },
    {
        "filename": "通用_出货记录.xlsx",
        "template_key": "SEED_SHIPMENT_RECORDS_DEFAULT",
        "template_name": "通用出货记录表",
        "template_type": "出货记录",
        "business_scope": "shipmentRecords",
        "sheet": "出货记录",
        "title": "XCMAX 通用出货记录",
        "headers": ["购买单位", "产品名称", "型号", "数量", "单价", "金额", "出货日期"],
        "rows": [
            ["演示客户有限公司", "演示产品A", "A001", 3, 12.5, 37.5, "2026-07-28"],
            ["演示客户有限公司", "演示哑光清面漆", "9803", 2, 18.8, 37.6, "2026-07-28"],
        ],
    },
    {
        "filename": "通用_产品目录.xlsx",
        "template_key": "SEED_PRODUCTS_DEFAULT",
        "template_name": "通用产品目录表",
        "template_type": "产品目录",
        "business_scope": "products",
        "sheet": "产品目录",
        "title": "XCMAX 通用产品目录",
        "headers": ["产品型号", "产品名称", "规格", "单价", "单位", "库存"],
        "rows": [
            ["A001", "演示产品A", "28", 12.5, "桶", 100],
            ["9803", "演示哑光清面漆", "28", 18.8, "桶", 50],
            ["DEMO-002", "演示产品B", "20", 9.8, "桶", 80],
        ],
    },
    {
        "filename": "通用_原材料.xlsx",
        "template_key": "SEED_MATERIALS_DEFAULT",
        "template_name": "通用原材料表",
        "template_type": "原材料",
        "business_scope": "materials",
        "sheet": "原材料",
        "title": "XCMAX 通用原材料仓库",
        "headers": ["原材料编码", "名称", "分类", "规格", "单位", "库存数量", "单价", "供应商"],
        "rows": [
            ["M-001", "树脂基料", "主料", "工业级", "kg", 500, 8.5, "演示供应商A"],
            ["M-002", "稀释剂", "辅料", "标准", "L", 200, 6.2, "演示供应商B"],
        ],
    },
    {
        "filename": "通用_客户.xlsx",
        "template_key": "SEED_CUSTOMERS_DEFAULT",
        "template_name": "通用客户表",
        "template_type": "客户",
        "business_scope": "customers",
        "sheet": "客户",
        "title": "XCMAX 通用客户名单",
        "headers": ["客户名称", "联系人", "电话", "地址", "备注"],
        "rows": [
            ["演示客户有限公司", "张三", "13800000000", "成都市高新区演示路 1 号", "开箱样例"],
            ["七彩乐园", "李四", "13900000000", "成都市武侯区示例街 8 号", ""],
        ],
    },
    {
        "filename": "通用_汇总统计.xlsx",
        "template_key": "SEED_SHIPMENT_SUMMARY_DEFAULT",
        "template_name": "通用汇总统计表",
        "template_type": "汇总统计",
        "business_scope": "shipmentSummary",
        "sheet": "汇总统计",
        "title": "XCMAX 通用汇总统计",
        "headers": ["统计周期", "单据数", "金额", "金额合计", "金额总计", "备注"],
        "rows": [
            ["2026-07", 12, 12800.5, 12800.5, 12800.5, "月汇总样例"],
            ["2026-06", 9, 9600.0, 9600.0, 9600.0, ""],
        ],
    },
    {
        "filename": "通用_销售报表.xlsx",
        "template_key": "SEED_SALES_REPORT_DEFAULT",
        "template_name": "通用销售报表",
        "template_type": "销售报表",
        "business_scope": "salesReport",
        "sheet": "销售报表",
        "title": "XCMAX 通用销售报表",
        "headers": ["客户名称", "销售金额", "实收款", "下欠款金额", "业务员", "月份"],
        "rows": [
            ["演示客户有限公司", 50000, 42000, 8000, "王五", "2026-07"],
            ["七彩乐园", 32000, 32000, 0, "赵六", "2026-07"],
        ],
    },
    {
        "filename": "通用_考勤记录.xlsx",
        "template_key": "SEED_ATTENDANCE_DEFAULT",
        "template_name": "通用考勤记录表",
        "template_type": "考勤记录",
        "business_scope": "shipmentRecords",
        "sheet": "考勤记录",
        "title": "XCMAX 通用考勤记录",
        "headers": ["购买单位", "产品名称", "型号", "数量", "单价", "金额", "日期", "备注"],
        "rows": [
            # 复用出货记录必填列别名，便于考勤行业映射「部门/人员」
            ["XC 演示部门", "XC 演示人员", "1001", 1, 0, 0, "2026-07-28", "出勤样例"],
            ["XC 演示部门", "演示员工B", "1002", 1, 0, 0, "2026-07-28", "请假半天"],
        ],
    },
]


def build_shipment_xlsx(path: Path) -> None:
    from openpyxl.styles import Alignment, Border, Font, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "送货单"
    thin = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    ws["A1"] = "XCMAX 演示送货单"
    ws["A1"].font = Font(name="宋体", size=16, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells("A1:J1")
    ws["A2"] = "购货单位：{{购买单位}}       联系人：              {{日期}}      订单编号："
    ws.merge_cells("A2:J2")
    headers = [
        (1, "产品型号"),
        (4, "产品名称"),
        (5, "数量/件"),
        (6, "规格/KG"),
        (7, "数量/KG"),
        (8, "单价/元"),
        (9, "金额/元"),
        (10, "备注"),
    ]
    for col, text in headers:
        cell = ws.cell(row=3, column=col, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    for row in range(4, 15):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        for col in range(1, 11):
            ws.cell(row=row, column=col).border = thin
    ws["A15"] = "合计"
    ws.merge_cells("A15:D15")
    ws["E15"] = "=SUM(E4:E14)"
    ws["G15"] = "=SUM(G4:G14)"
    ws["I15"] = "=SUM(I4:I14)"
    ws["A16"] = "大写人民币："
    ws.merge_cells("A16:C16")
    ws["A18"] = "收货人签收："
    ws.merge_cells("A18:E18")
    ws["F18"] = "送货人："
    ws.merge_cells("F18:J18")
    from openpyxl.utils import get_column_letter

    widths = {1: 12, 2: 8, 3: 8, 4: 22, 5: 10, 6: 10, 7: 10, 8: 10, 9: 12, 10: 12}
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def build_price_list_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    title = doc.add_paragraph("产品价格表")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "宋体"
    doc.add_paragraph("客户：{{客户}}　　报价日期：{{报价日期}}")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["型号", "名称", "规格", "单价"]):
        table.rows[0].cells[i].text = h
    for r, sample in enumerate(
        (("DEMO-001", "演示产品A", "28", "12.5"), ("DEMO-002", "演示产品B", "20", "9.8")),
        start=1,
    ):
        for c, val in enumerate(sample):
            table.rows[r].cells[c].text = val
    foot = doc.add_paragraph(
        "说明：本表为 XCMAX 开箱演示模板，可在模板库中替换为企业自有版式。"
    )
    foot.runs[0].font.size = Pt(9)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    out = _repo_root() / "resources" / "templates"
    written: list[str] = []

    ship = out / "发货单模板.xlsx"
    build_shipment_xlsx(ship)
    written.append(str(ship))
    # 仅作老生成器 uploads 别名源；运行时不会拷进可扫描 templates/
    alias = out / "尹玉华1.xlsx"
    alias.write_bytes(ship.read_bytes())
    written.append(str(alias))

    price = out / "price_list_default.docx"
    build_price_list_docx(price)
    written.append(str(price))

    for spec in GENERIC_EXCEL_SPECS:
        path = out / str(spec["filename"])
        _write_sheet(
            path,
            sheet_title=str(spec["sheet"]),
            headers=list(spec["headers"]),
            rows=list(spec["rows"]),
            title=str(spec.get("title") or ""),
        )
        written.append(str(path))

    print(f"wrote {len(written)} files under {out}")
    for p in written:
        print(" -", Path(p).name, Path(p).stat().st_size)


if __name__ == "__main__":
    main()
