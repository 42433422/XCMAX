from __future__ import annotations

import json
from unittest.mock import patch

import pytest

from app.application.agent_orchestrator.chat_trace import attach_chat_trace_run
from app.application.agent_orchestrator.run_repository import InMemoryAgentRunRepository


@pytest.fixture(autouse=True)
def _isolated_agent_usage_ledger(tmp_path, monkeypatch):
    monkeypatch.setenv("MODEL_USAGE_LEDGER_PATH", str(tmp_path / "model_usage_ledger.json"))
    monkeypatch.delenv("MODEL_USAGE_WALLET_BACKEND", raising=False)
    monkeypatch.delenv("MODEL_USAGE_WALLET_REQUIRED", raising=False)


def test_attach_chat_trace_run_persists_completed_run() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {"success": True, "response": "done", "data": {"text": "done"}}

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="查库存",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run_id = result["run_id"]
    assert result["data"]["run_id"] == run_id

    run = repo.get(run_id)
    assert run is not None
    assert run.user_id == "u1"
    assert run.status == "completed"
    assert run.intent == "legacy_chat_adapter"
    assert run.metadata["source"] == "desktop"
    assert run.metadata["channel"] == "compat_chat"
    assert run.final_output["chat_payload"]["response"] == "done"
    assert [event.event_type for event in run.events] == ["run.created", "run.completed"]


def test_attach_chat_trace_run_accepts_explicit_intent() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {"success": True, "response": "done", "data": {"text": "done"}}

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="导入数据库",
            runtime_context={"workflow_trace_mode": "deterministic_shortcut"},
            user_id="u1",
            source="pro",
            channel="deterministic_workflow",
            intent="excel_import_to_db",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.intent == "excel_import_to_db"
    assert run.metadata["channel"] == "deterministic_workflow"
    assert run.metadata["runtime_context"]["workflow_trace_mode"] == "deterministic_shortcut"


def test_attach_chat_trace_run_records_llm_calls(tmp_path, monkeypatch) -> None:
    repo = InMemoryAgentRunRepository()
    ledger_path = tmp_path / "model_usage_ledger.json"
    monkeypatch.setenv("MODEL_USAGE_LEDGER_PATH", str(ledger_path))
    trace = {
        "provider_id": "openai_compatible",
        "provider": "xcauto",
        "model": "xcauto-account",
        "prompt_tokens": 2,
        "completion_tokens": 3,
        "total_tokens": 5,
        "latency_ms": 12.5,
        "cost_units": 1,
        "billing_status": "metered",
        "billing_source": "estimated_token_units",
    }
    payload = {
        "success": True,
        "response": "done",
        "data": {"text": "done", "llm_trace": trace},
        "_xcagi_trace": trace,
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="普通问答",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert len(run.llm_calls) == 1
    assert run.llm_calls[0].provider == "xcauto"
    assert run.llm_calls[0].model == "xcauto-account"
    assert run.llm_calls[0].total_tokens == 5
    assert run.llm_calls[0].cost_units == 1
    assert run.llm_calls[0].billing_status == "metered"
    assert run.llm_calls[0].metadata["usage_ledger"]["status"] == "recorded"
    assert run.metadata["llm_call_count"] == 1
    assert run.metadata["llm_token_total"] == 5
    assert run.metadata["llm_cost_units_total"] == 1
    assert run.metadata["ai_cost_units_total"] == 1
    assert run.metadata["model_usage_entry_count"] == 1
    assert run.metadata["model_usage_cost_units_total"] == 1
    assert run.metadata["model_usage_ledger_status"] == "recorded"
    assert run.final_output["llm_calls"][0]["provider"] == "xcauto"
    assert run.final_output["llm_calls"][0]["metadata"]["usage_ledger"]["status"] == "recorded"
    assert run.final_output["llm_cost_units_total"] == 1
    assert run.final_output["ai_cost_units_total"] == 1
    assert [event.event_type for event in run.events] == [
        "run.created",
        "llm.completed",
        "billing.recorded",
        "run.completed",
    ]
    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    assert ledger["summary"]["entry_count"] == 1
    assert ledger["summary"]["cost_units_total"] == 1
    assert ledger["entries"][0]["run_id"] == run.run_id
    assert ledger["entries"][0]["provider"] == "xcauto"
    assert ledger["entries"][0]["model"] == "xcauto-account"


def test_attach_chat_trace_run_debits_local_model_wallet(tmp_path, monkeypatch) -> None:
    repo = InMemoryAgentRunRepository()
    ledger_path = tmp_path / "model_usage_ledger.json"
    monkeypatch.setenv("MODEL_USAGE_LEDGER_PATH", str(ledger_path))
    from app.infrastructure.billing.model_usage import get_model_wallet, set_model_wallet_balance

    set_model_wallet_balance("u1", 2, reason="test")
    trace = {
        "provider_id": "openai_compatible",
        "provider": "xcauto",
        "model": "xcauto-account",
        "total_tokens": 5,
        "cost_units": 1,
        "billing_status": "metered",
        "billing_source": "estimated_token_units",
    }
    payload = {"success": True, "response": "done", "data": {"llm_trace": trace}}

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="普通问答",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert run.llm_calls[0].billing_status == "debited"
    assert run.llm_calls[0].billing_source == "local_model_wallet"
    assert run.llm_calls[0].metadata["wallet_debit"]["balance_after_units"] == 1
    assert run.metadata["model_wallet_balance_units"] == 1
    assert "billing.debited" in [event.event_type for event in run.events]
    assert get_model_wallet("u1")["balance_units"] == 1


def test_attach_chat_trace_run_fails_when_model_wallet_insufficient(tmp_path, monkeypatch) -> None:
    repo = InMemoryAgentRunRepository()
    ledger_path = tmp_path / "model_usage_ledger.json"
    monkeypatch.setenv("MODEL_USAGE_LEDGER_PATH", str(ledger_path))
    monkeypatch.setenv("MODEL_USAGE_WALLET_REQUIRED", "1")
    trace = {
        "provider_id": "openai_compatible",
        "provider": "xcauto",
        "model": "xcauto-account",
        "total_tokens": 5,
        "cost_units": 1,
        "billing_status": "metered",
        "billing_source": "estimated_token_units",
    }
    payload = {"success": True, "response": "done", "data": {"llm_trace": trace}}

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="普通问答",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "failed"
    assert run.error == "AI wallet balance insufficient"
    assert run.llm_calls[0].billing_status == "insufficient_balance"
    assert run.llm_calls[0].metadata["wallet_debit"]["shortfall_units"] == 1
    assert "billing.insufficient_balance" in [event.event_type for event in run.events]


def test_attach_chat_trace_run_records_rag_retrieval() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "付款条款是 30 天",
        "data": {
            "text": "付款条款是 30 天",
            "query": "合同付款条款",
            "rag_enabled": True,
            "dataset_id": "contracts",
            "chunks": [
                {
                    "chunk_index": 0,
                    "text": "付款条款：验收后 30 天内付款",
                    "score": 0.91,
                    "source": "contract.pdf",
                }
            ],
            "citations": [
                {
                    "index": 1,
                    "text": "付款条款：验收后 30 天内付款",
                    "source": "contract.pdf",
                    "chunk_index": 0,
                }
            ],
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="合同付款条款是什么",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert len(run.retrieval_calls) == 1
    assert run.retrieval_calls[0].query == "合同付款条款"
    assert run.retrieval_calls[0].source == "contracts"
    assert run.retrieval_calls[0].chunks[0]["source"] == "contract.pdf"
    assert run.retrieval_calls[0].citations[0]["source"] == "contract.pdf"
    assert run.metadata["retrieval_call_count"] == 1
    assert run.metadata["retrieval_chunk_count"] == 1
    assert run.metadata["citation_count"] == 1
    assert run.final_output["retrieval_calls"][0]["citations"][0]["source"] == "contract.pdf"
    assert [event.event_type for event in run.events] == [
        "run.created",
        "rag.retrieved",
        "run.completed",
    ]


def test_attach_chat_trace_run_records_user_memory_reference() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "按上次客户偏好优先看涂料类",
        "data": {
            "text": "按上次客户偏好优先看涂料类",
            "query": "上次客户偏好是什么",
            "user_memory_rag_summary": "【UserMemoryRAG】召回 1 条：客户偏好先看涂料类报价",
            "user_memory_hits": [
                {
                    "chunk_id": "mem1",
                    "content": "客户偏好先看涂料类报价",
                    "score": 0.88,
                    "metadata": {"intent": "business_db_read"},
                }
            ],
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="按上次客户偏好查询",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert len(run.memory_references) == 1
    assert run.memory_references[0].query == "上次客户偏好是什么"
    assert run.memory_references[0].source == "user_memory_rag"
    assert run.memory_references[0].hits[0]["chunk_id"] == "mem1"
    assert run.memory_references[0].summary.startswith("【UserMemoryRAG】")
    assert run.metadata["memory_reference_count"] == 1
    assert run.metadata["memory_hit_count"] == 1
    assert run.metadata["memory_sources"] == ["user_memory_rag"]
    assert run.final_output["memory_references"][0]["hits"][0]["chunk_id"] == "mem1"
    assert run.final_output["memory_hit_count"] == 1
    assert [event.event_type for event in run.events] == [
        "run.created",
        "memory.recalled",
        "run.completed",
    ]


def test_attach_chat_trace_run_records_ocr_artifact(tmp_path, monkeypatch) -> None:
    from app.application.dataset_rag_app_service import (
        get_dataset_rag_app_service,
        reset_dataset_rag_app_service_for_tests,
    )

    monkeypatch.setenv("DATASET_RAG_STORE_PATH", str(tmp_path / "dataset_store.json"))
    reset_dataset_rag_app_service_for_tests()
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "识别到购货单位",
        "data": {
            "text": "识别到购货单位",
            "ocr_result": {
                "success": True,
                "text": "购货单位：星光贸易\n联系人：张三",
                "confidence": 0.93,
                "file_path": "/tmp/label.png",
                "structured_data": {
                    "purchase_unit": "星光贸易",
                    "contact_person": "张三",
                },
            },
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="识别这张产品图",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert len(run.artifacts) == 1
    assert run.artifacts[0].artifact_type == "ocr_text"
    assert run.artifacts[0].uri == "/tmp/label.png"
    assert run.artifacts[0].fields[0]["name"] == "purchase_unit"
    assert run.artifacts[0].preview["confidence"] == 0.93
    assert run.metadata["artifact_count"] == 1
    assert run.metadata["artifact_types"] == ["ocr_text"]
    assert run.metadata["dataset_ingest_count"] == 1
    assert run.metadata["dataset_ids"] == ["user_u1"]
    assert run.final_output["artifacts"][0]["artifact_type"] == "ocr_text"
    assert run.final_output["dataset_ingest_count"] == 1
    answer = get_dataset_rag_app_service().answer(
        dataset_id="user_u1",
        query="购货单位",
        top_k=1,
    )
    assert answer["success"] is True
    assert "星光贸易" in answer["answer"]
    assert [event.event_type for event in run.events] == [
        "run.created",
        "artifact.attached",
        "dataset.ingested",
        "run.completed",
    ]
    reset_dataset_rag_app_service_for_tests()


def test_attach_chat_trace_run_ingests_pdf_file_analysis_artifact(tmp_path, monkeypatch) -> None:
    from app.application.dataset_rag_app_service import (
        get_dataset_rag_app_service,
        reset_dataset_rag_app_service_for_tests,
    )
    from evals.run_agent_eval import _minimal_pdf_bytes

    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("DATASET_RAG_STORE_PATH", str(tmp_path / "dataset_store.json"))
    reset_dataset_rag_app_service_for_tests()
    repo = InMemoryAgentRunRepository()
    pdf_path = tmp_path / "xcauto_policy.pdf"
    pdf_path.write_bytes(
        _minimal_pdf_bytes(
            "All AI routes should use the server platform XCauto model. "
            "Artifacts from PDF analysis must enter Dataset RAG with citations."
        )
    )
    payload = {
        "success": True,
        "response": "PDF 已解析",
        "data": {
            "text": "PDF 已解析",
            "file_analysis": {
                "success": True,
                "parser_used": "pdfplumber",
                "extension": ".pdf",
                "file_path": str(pdf_path),
                "saved_name": "xcauto_policy.pdf",
                "text_preview": "All AI routes should use the server platform XCauto model.",
                "ai_summary": "XCauto policy PDF",
            },
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="解析并加入知识库",
            runtime_context={"workspace": "demo", "dataset_id": "platform-docs"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert run.artifacts[0].artifact_type == "pdf_document"
    assert run.metadata["dataset_ingest_count"] == 1
    assert run.metadata["dataset_ids"] == ["platform-docs"]
    ingest = run.metadata["dataset_ingests"][0]
    assert ingest["success"] is True
    assert ingest["parser"] == "pdfplumber"
    assert ingest["source"] == "xcauto_policy.pdf"
    assert "dataset.ingested" in [event.event_type for event in run.events]

    answer = get_dataset_rag_app_service().answer(
        dataset_id="platform-docs",
        query="Which model should AI routes use?",
        top_k=2,
    )
    assert answer["success"] is True
    assert "XCauto" in answer["answer"]
    assert answer["citations"]

    reset_dataset_rag_app_service_for_tests()
    reload_answer = get_dataset_rag_app_service().answer(
        dataset_id="platform-docs",
        query="Which model should AI routes use?",
        top_k=2,
    )
    assert reload_answer["success"] is True
    assert "XCauto" in reload_answer["answer"]
    reset_dataset_rag_app_service_for_tests()


def test_attach_chat_trace_run_ingests_office_file_analysis_artifact(tmp_path, monkeypatch) -> None:
    from docx import Document

    from app.application.dataset_rag_app_service import (
        get_dataset_rag_app_service,
        reset_dataset_rag_app_service_for_tests,
    )

    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("DATASET_RAG_STORE_PATH", str(tmp_path / "dataset_store.json"))
    reset_dataset_rag_app_service_for_tests()
    repo = InMemoryAgentRunRepository()
    docx_path = tmp_path / "platform_notes.docx"
    doc = Document()
    doc.add_paragraph("Office artifacts should enter Dataset RAG automatically.")
    doc.add_paragraph("The server platform model is XCauto.")
    doc.save(str(docx_path))
    payload = {
        "success": True,
        "response": "DOCX 已解析",
        "data": {
            "text": "DOCX 已解析",
            "file_analysis": {
                "success": True,
                "parser_used": "python-docx",
                "extension": ".docx",
                "file_path": str(docx_path),
                "saved_name": "platform_notes.docx",
                "text_preview": "The server platform model is XCauto.",
                "ai_summary": "Platform notes DOCX",
            },
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="解析 Office 并加入知识库",
            runtime_context={"workspace": "demo", "dataset_id": "office-docs"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.artifacts[0].artifact_type == "office_document"
    assert run.metadata["dataset_ingest_count"] == 1
    assert run.metadata["dataset_ingests"][0]["parser"] == "python-docx"
    answer = get_dataset_rag_app_service().answer(
        dataset_id="office-docs",
        query="Which model should Office artifacts mention?",
        top_k=2,
    )
    assert answer["success"] is True
    assert "XCauto" in answer["answer"]
    reset_dataset_rag_app_service_for_tests()


def test_attach_chat_trace_run_records_generated_office_artifact() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "合同已生成",
        "data": {
            "text": "合同已生成",
            "file_name": "contract.docx",
            "download_url": "/api/ai/kitten/document/pickup/doc-token",
            "pickup_token": "doc-token",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(
            payload,
            message="生成合同",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert len(run.artifacts) == 1
    assert run.artifacts[0].artifact_type == "office_document"
    assert run.artifacts[0].name == "contract.docx"
    assert run.artifacts[0].uri == "/api/ai/kitten/document/pickup/doc-token"
    assert run.artifacts[0].preview["pickup_token"] == "doc-token"
    assert run.metadata["artifact_types"] == ["office_document"]
    assert run.final_output["artifacts"][0]["artifact_type"] == "office_document"
    assert [event.event_type for event in run.events] == [
        "run.created",
        "artifact.attached",
        "run.completed",
    ]


def test_attach_chat_trace_run_marks_token_waiting() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "requires_token": True,
        "token_name": "DB_READ_TOKEN",
        "message": "需要授权",
        "data": {"requires_token": True},
    }

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(payload, message="查看数据库", runtime_context={})

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "waiting_user"
    assert run.events[-1].event_type == "step.waiting_user"


def test_attach_chat_trace_run_skips_existing_run_id() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {"success": True, "run_id": "run_existing", "data": {"text": "ok"}}

    with patch(
        "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
        return_value=repo,
    ):
        result = attach_chat_trace_run(payload, message="hello", runtime_context={})

    assert result["run_id"] == "run_existing"
    assert repo.list_recent() == []


def test_attach_chat_trace_run_orchestrates_low_risk_tool_call() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "查询到 1 个产品",
        "toolCall": {
            "tool_id": "products",
            "action": "执行",
            "params": {"keyword": "5003"},
        },
        "data": {"text": "查询到 1 个产品"},
    }

    with (
        patch(
            "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
            return_value=repo,
        ),
        patch(
            "app.application.facades.tools_facade.execute_registered_workflow_tool",
            return_value={"success": True, "data": [{"model_number": "5003"}]},
        ) as mock_execute,
    ):
        result = attach_chat_trace_run(
            payload,
            message="查产品 5003",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert run.intent == "products_query"
    assert run.metadata["trace_mode"] == "orchestrated_tool_call"
    assert run.metadata["channel"] == "compat_chat"
    assert run.metadata["source"] == "desktop"
    assert run.metadata["tool_call_count"] == 1
    assert run.tool_calls[0].tool_id == "products"
    assert run.tool_calls[0].action == "query"
    assert run.tool_calls[0].cost_units == 1
    assert "tool.started" in [event.event_type for event in run.events]
    assert "tool.completed" in [event.event_type for event in run.events]
    assert result["data"]["agent_run_id"] == run.run_id

    mock_execute.assert_called_once()
    tool_id, action, params = mock_execute.call_args.args
    assert tool_id == "products"
    assert action == "query"
    assert params["keyword"] == "5003"
    assert params["_runtime_context"]["trace_mode"] == "orchestrated_tool_call"


def test_attach_chat_trace_run_does_not_orchestrate_medium_nested_action() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "准备创建产品",
        "toolCall": {
            "tool_id": "products",
            "action": "执行",
            "params": {"action": "create", "name_or_model": "A1", "unit_name": "客户A"},
        },
        "data": {"text": "准备创建产品"},
    }

    with (
        patch(
            "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
            return_value=repo,
        ),
        patch(
            "app.application.facades.tools_facade.execute_registered_workflow_tool"
        ) as mock_execute,
    ):
        result = attach_chat_trace_run(
            payload,
            message="新增产品 A1",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert run.intent == "legacy_chat_adapter"
    assert run.metadata["trace_mode"] == "post_execution"
    assert run.tool_calls == []
    mock_execute.assert_not_called()


def test_attach_chat_trace_run_observes_legacy_tool_records_without_reexecution() -> None:
    repo = InMemoryAgentRunRepository()
    payload = {
        "success": True,
        "response": "查询完成",
        "data": {
            "text": "查询完成",
            "legacy_tool_records": [
                {
                    "tool_id": "products",
                    "tool_name": "products",
                    "action": "query",
                    "params": {"keyword": "5003"},
                    "output": {"success": True, "data": [{"model_number": "5003"}]},
                    "tool_call_id": "tc1",
                }
            ],
        },
    }

    with (
        patch(
            "app.application.agent_orchestrator.chat_trace.get_agent_run_repository",
            return_value=repo,
        ),
        patch(
            "app.application.facades.tools_facade.execute_registered_workflow_tool"
        ) as mock_execute,
    ):
        result = attach_chat_trace_run(
            payload,
            message="查产品 5003",
            runtime_context={"workspace": "demo"},
            user_id="u1",
            source="desktop",
            channel="compat_chat",
        )

    run = repo.get(result["run_id"])
    assert run is not None
    assert run.status == "completed"
    assert run.intent == "legacy_tool_chain"
    assert run.metadata["trace_mode"] == "legacy_tool_records"
    assert run.metadata["tool_call_count"] == 1
    assert run.metadata["cost_units_total"] == 1
    assert run.tool_calls[0].metadata["observed"] is True
    assert run.tool_calls[0].metadata["legacy_tool_call_id"] == "tc1"
    assert run.tool_calls[0].tool_id == "products"
    assert run.tool_calls[0].action == "query"
    assert run.final_output["node_outputs"][run.steps[0].node_id]["success"] is True
    assert "tool.completed" in [event.event_type for event in run.events]
    mock_execute.assert_not_called()
