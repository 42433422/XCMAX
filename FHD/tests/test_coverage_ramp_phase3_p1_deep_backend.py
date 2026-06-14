"""COVERAGE_RAMP Phase 3 round 1: workflow/planner deep branches, price_list_export,
ai_chat_v2, service_bridge + excel_extract routes, compat_db writes."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pandas as pd
import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.application.ai_chat_app_service_v2 import (
    AiChatAppServiceV2,
    get_ai_chat_app_service_v2,
)
from app.application.tools.workflow import (
    execute_workflow_tool,
    handle_excel_analysis,
    run_natural_language_pandas,
)
from app.application.workflow.planner import (
    _execute_excel_analysis_tool,
    _execute_excel_decompose_tool,
    _execute_excel_schema_tool,
    _execute_import_excel_tool,
    _execute_print_label_tool,
    _execute_template_extract_tool,
    _execute_wechat_preview_tool,
    execute_tool,
)
from app.infrastructure.documents.price_list_export import (
    _format_price_cell,
    _header_text,
    _parse_header_serial_and_column_map,
    _product_row_cell_values,
    _row_keyword_score,
    build_price_list_docx_bytes,
    build_price_list_template_preview_json,
)
from app.infrastructure.persistence.compat_db.writes import (
    _customer_pg_insert,
    products_pg_delete_row,
    products_pg_update_row,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_xlsx(path: Path, rows: list[list], sheet: str = "Sheet1") -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows:
        ws.append(row)
    wb.save(path)
    return path


def _write_docx_template(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("客户：{{客户}} 日期：{{报价日期}}")
    tbl = doc.add_table(rows=2, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = "序号"
    hdr[1].text = "型号"
    hdr[2].text = "名称"
    hdr[3].text = "单价"
    body = tbl.rows[1].cells
    body[0].text = "1"
    body[1].text = "M1"
    body[2].text = "样例"
    body[3].text = "0"
    doc.save(path)
    return path


@pytest.fixture
def bridge_client() -> TestClient:
    from app.fastapi_routes import service_bridge

    app = FastAPI()
    app.include_router(service_bridge.router)
    return TestClient(app, raise_server_exceptions=False)


@pytest.fixture
def excel_client() -> TestClient:
    from app.fastapi_routes import excel_extract

    app = FastAPI()
    app.include_router(excel_extract.router)
    return TestClient(app, raise_server_exceptions=False)


def _mock_db_ctx(query_result=None, count_result=0):
    db = MagicMock()
    q = MagicMock()
    q.filter.return_value = q
    q.first.return_value = query_result
    q.count.return_value = count_result
    q.order_by.return_value = q
    q.offset.return_value = q
    q.limit.return_value = q
    q.all.return_value = query_result or []
    db.query.return_value = q
    ctx = MagicMock()
    ctx.__enter__ = MagicMock(return_value=db)
    ctx.__exit__ = MagicMock(return_value=False)
    return ctx, db


# ---------------------------------------------------------------------------
# workflow execute_workflow_tool branches
# ---------------------------------------------------------------------------


@pytest.fixture(autouse=True)
def _bypass_native_planner_tools(monkeypatch: pytest.MonkeyPatch) -> None:
    """跳过 mod/员工工具分发，命中 workflow 本地实现。"""
    monkeypatch.setattr(
        "app.mod_sdk.planner_native_tools.try_execute_native_planner_tool",
        lambda *a, **k: (None, None),
    )
    monkeypatch.setattr(
        "app.mod_sdk.employee_tool_registry.is_employee_tool",
        lambda _n: False,
    )
    monkeypatch.setattr(
        "app.application.employee_pack_runner.try_execute_employee_planner_tool",
        lambda *a, **k: None,
    )


def test_execute_workflow_tool_json_string_args() -> None:
    raw = execute_workflow_tool("excel_chart_recommend", "{}")
    out = json.loads(raw)
    assert "suggestions" in out


def test_execute_workflow_tool_invalid_json_args() -> None:
    raw = execute_workflow_tool("excel_chart_recommend", "not-json")
    out = json.loads(raw)
    assert "suggestions" in out


def test_execute_workflow_tool_excel_join_missing_files(tmp_path) -> None:
    raw = execute_workflow_tool(
        "excel_join_compare",
        {"action": "join", "file_paths": ["a.xlsx", "b.xlsx"]},
        workspace_root=str(tmp_path),
    )
    out = json.loads(raw)
    assert out["success"] is False


def test_execute_workflow_tool_excel_join_success(tmp_path) -> None:
    p1 = _write_xlsx(tmp_path / "a.xlsx", [["id", "val"], [1, 10]])
    p2 = _write_xlsx(tmp_path / "b.xlsx", [["id", "extra"], [1, "x"]])
    d1 = pd.DataFrame({"id": [1], "val": [10]})
    d2 = pd.DataFrame({"id": [1], "extra": ["x"]})
    with patch("app.application.tools.workflow.pd.read_excel", side_effect=[d1, d2]):
        raw = execute_workflow_tool(
            "excel_join_compare",
            {"action": "join", "file_paths": [str(p1), str(p2)], "join_keys": ["id"]},
            workspace_root=str(tmp_path),
        )
    out = json.loads(raw)
    assert out["action"] == "join"
    assert out["row_count"] == 1


def test_execute_workflow_tool_excel_diff_with_keys(tmp_path) -> None:
    p1 = _write_xlsx(tmp_path / "a.xlsx", [["id", "v"], [1, 1]])
    p2 = _write_xlsx(tmp_path / "b.xlsx", [["id", "v"], [1, 2], [2, 3]])
    d1 = pd.DataFrame({"id": [1], "v": [1]})
    d2 = pd.DataFrame({"id": [1, 2], "v": [2, 3]})
    with patch("app.application.tools.workflow.pd.read_excel", side_effect=[d1, d2]):
        raw = execute_workflow_tool(
            "excel_join_compare",
            {
                "action": "diff",
                "file_path_a": str(p1),
                "file_path_b": str(p2),
                "key_columns": ["id"],
            },
            workspace_root=str(tmp_path),
        )
    out = json.loads(raw)
    assert out["action"] == "diff"
    assert "only_in_right" in out


def test_execute_workflow_tool_excel_diff_no_keys(tmp_path) -> None:
    p1 = _write_xlsx(tmp_path / "a.xlsx", [["x"], [1]])
    p2 = _write_xlsx(tmp_path / "b.xlsx", [["x"], [2]])
    with patch(
        "app.application.tools.workflow.pd.read_excel",
        side_effect=[pd.DataFrame({"x": [1]}), pd.DataFrame({"x": [2]})],
    ):
        raw = execute_workflow_tool(
            "excel_join_compare",
            {"action": "diff", "file_path_a": str(p1), "file_path_b": str(p2)},
            workspace_root=str(tmp_path),
        )
    out = json.loads(raw)
    assert out["row_count"] == 1


def test_execute_workflow_tool_excel_unknown_action() -> None:
    raw = execute_workflow_tool("excel_join_compare", {"action": "unknown"})
    out = json.loads(raw)
    assert out["success"] is False


def test_execute_workflow_tool_excel_prophet_insufficient_data() -> None:
    raw = execute_workflow_tool("excel_prophet", {"periods": 3})
    out = json.loads(raw)
    assert out["action"] == "forecast"
    assert len(out["future_forecast"]) == 3


def test_execute_workflow_tool_excel_prophet_with_file(tmp_path) -> None:
    p = _write_xlsx(tmp_path / "sales.xlsx", [["month", "amt"], [1, 10], [2, 20], [3, 30]])
    with patch(
        "app.application.tools.workflow._read_excel_dataframe",
        return_value=pd.DataFrame({"amt": [10, 20, 30]}),
    ):
        raw = execute_workflow_tool(
            "excel_prophet",
            {"file_path": str(p), "value_column": "amt", "periods": 2},
            workspace_root=str(tmp_path),
        )
    out = json.loads(raw)
    assert out["model"] == "linear_regression"
    assert len(out["future_forecast"]) == 2


def test_execute_workflow_tool_excel_schema_understand(tmp_path) -> None:
    p = _write_xlsx(tmp_path / "t.xlsx", [["产品", "单价"], ["漆", 99]])
    with patch(
        "app.application.tools.workflow._read_excel_dataframe",
        return_value=pd.DataFrame({"产品": ["漆"], "单价": [99]}),
    ):
        raw = execute_workflow_tool(
            "excel_schema_understand",
            {"file_path": str(p)},
            workspace_root=str(tmp_path),
        )
    out = json.loads(raw)
    assert out.get("success") is True or "columns" in out or "snapshot" in out


def test_handle_excel_analysis_read_query_aggregate(tmp_path) -> None:
    df = pd.DataFrame({"cat": ["A", "A", "B"], "qty": [2, 3, 1]})
    p = tmp_path / "data.xlsx"
    _write_xlsx(p, [["cat", "qty"], ["A", 2], ["A", 3], ["B", 1]])
    with patch("app.application.tools.workflow._read_excel_dataframe", return_value=df):
        read = handle_excel_analysis(
            {"file_path": p.name, "action": "read"},
            workspace_root=str(tmp_path),
        )
    assert read["success"] is True
    assert read["row_count"] == 3

    with patch("app.application.tools.workflow._read_excel_dataframe", return_value=df):
        query = handle_excel_analysis(
            {"file_path": p.name, "action": "query", "query_expression": "qty > 1"},
            workspace_root=str(tmp_path),
        )
    assert query["success"] is True

    with patch("app.application.tools.workflow._read_excel_dataframe", return_value=df):
        agg = handle_excel_analysis(
            {
                "file_path": p.name,
                "action": "aggregate",
                "group_by": ["cat"],
                "metrics": [{"column": "qty", "op": "sum"}],
            },
            workspace_root=str(tmp_path),
        )
    assert agg["success"] is True

    with patch("app.application.tools.workflow._read_excel_dataframe", return_value=df):
        stats = handle_excel_analysis(
            {"file_path": p.name, "action": "statistics"},
            workspace_root=str(tmp_path),
        )
    assert stats["success"] is True


def test_handle_excel_analysis_missing_file(tmp_path) -> None:
    out = handle_excel_analysis(
        {"file_path": "missing.xlsx", "action": "read"},
        workspace_root=str(tmp_path),
    )
    assert out["success"] is False


def test_handle_excel_analysis_unsupported_action(tmp_path) -> None:
    p = _write_xlsx(tmp_path / "t.xlsx", [["a"], [1]])
    out = handle_excel_analysis(
        {"file_path": p.name, "action": "nope"},
        workspace_root=str(tmp_path),
    )
    assert out["success"] is False


def test_run_natural_language_pandas_fallback() -> None:
    df = pd.DataFrame({"qty": [1, 2, 3]})
    out = run_natural_language_pandas(df, "求和")
    assert out["row_count"] == 3
    assert "records" in out


# ---------------------------------------------------------------------------
# planner tool handlers
# ---------------------------------------------------------------------------


def test_execute_print_label_with_products() -> None:
    with patch(
        "app.infrastructure.documents.shipment_document_generator_impl.SimpleLabelGenerator"
    ) as mock_gen_cls:
        mock_gen_cls.return_value.generate_labels_for_order.return_value = [{"id": 1}]
        out = _execute_print_label_tool(
            {"products": [{"name": "漆", "model": "5003"}], "unit_name": "甲公司"}
        )
    assert out["success"] is True


def test_execute_excel_decompose_missing_path() -> None:
    out = _execute_excel_decompose_tool({})
    assert out["success"] is False
    assert out["error_code"] == "missing_file_path"


@patch("app.bootstrap.get_template_app_service")
def test_execute_excel_decompose_ok(mock_get: MagicMock) -> None:
    mock_get.return_value.decompose_template.return_value = {"success": True}
    out = _execute_excel_decompose_tool({"file_path": "/tmp/t.xlsx"})
    assert out["success"] is True


def test_execute_template_extract_delegates() -> None:
    with patch(
        "app.application.workflow.planner._execute_excel_decompose_tool",
        return_value={"success": True},
    ):
        out = _execute_template_extract_tool({"file_path": "a.xlsx"})
    assert out["success"] is True


@patch("app.bootstrap.get_wechat_contact_app_service")
def test_execute_wechat_preview_tool(mock_get: MagicMock) -> None:
    mock_get.return_value.get_contacts.return_value = [{"name": "张三"}]
    out = _execute_wechat_preview_tool({"keyword": "张"})
    assert out["success"] is True
    assert out["data"]


def test_execute_excel_schema_openpyxl_fallback(tmp_path) -> None:
    p = _write_xlsx(tmp_path / "schema.xlsx", [["型号", "单价"], ["M1", 10]])
    out = _execute_excel_schema_tool({"file_path": str(p)})
    assert out["success"] is True
    assert out["fields"]


def test_execute_excel_schema_missing_path() -> None:
    out = _execute_excel_schema_tool({})
    assert out["error_code"] == "missing_file_path"


def test_execute_excel_analysis_openpyxl_fallback(tmp_path) -> None:
    p = _write_xlsx(tmp_path / "data.xlsx", [["产品", "单价"], ["漆", 88]])
    out = _execute_excel_analysis_tool({"file_path": str(p)})
    assert out["success"] is True


@patch("app.bootstrap.get_products_service")
def test_execute_import_excel_tool(mock_products: MagicMock, tmp_path) -> None:
    p = _write_xlsx(
        tmp_path / "import.xlsx",
        [["产品名称", "型号", "单价", "单位"], ["清漆", "5003", 120, "甲公司"]],
    )
    svc = MagicMock()
    svc.get_products.return_value = {"success": True, "data": []}
    svc.create_product.return_value = {"success": True}
    mock_products.return_value = svc
    with patch("app.bootstrap.get_customer_app_service") as mock_cust:
        mock_cust.return_value.match_purchase_unit.return_value = {"unit_name": "甲公司"}
        out = _execute_import_excel_tool(
            {"file_path": str(p), "unit_name": "甲公司", "skip_duplicates": False}
        )
    assert out["success"] is True


def test_execute_tool_action_defaults() -> None:
    with patch("app.bootstrap.get_shipment_app_service") as mock_ship:
        mock_ship.return_value.get_shipment_records.return_value = []
        out = execute_tool("shipment_records", {})
    assert out["success"] is True


# ---------------------------------------------------------------------------
# price_list_export
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "val,expected",
    [
        (None, ""),
        ("", ""),
        (100.0, "100"),
        (99.5, "99.50"),
        ("文本", "文本"),
    ],
)
def test_format_price_cell(val, expected) -> None:
    assert _format_price_cell(val) == expected


def test_row_keyword_score_and_header_map() -> None:
    cells = [MagicMock(text=v) for v in ["序号", "型号", "名称", "规格", "单价"]]
    assert _row_keyword_score(cells) >= 3
    has_serial, col_map = _parse_header_serial_and_column_map(cells)
    assert has_serial is True
    assert "model" in col_map or "name" in col_map


def test_product_row_cell_values_cn_keys() -> None:
    vals = _product_row_cell_values(
        {"型号": "M9", "产品名称": "漆", "规格": "1L", "单价": 50}
    )
    assert any("M9" in v or "漆" in v for v in vals)


@patch("app.infrastructure.documents.price_list_export.resolve_price_list_docx_template")
def test_build_price_list_template_preview(mock_resolve: MagicMock) -> None:
    mock_resolve.return_value = (Path("/tmp/t.docx"), "templates/t.docx")
    out = build_price_list_template_preview_json()
    assert out["success"] is True


def test_build_price_list_docx_bytes(tmp_path) -> None:
    tpl = _write_docx_template(tmp_path / "tpl.docx")
    data = build_price_list_docx_bytes(
        template_path=tpl,
        customer_name="甲公司",
        quote_date="2026-06-14",
        products=[{"model": "5003", "name": "清漆", "spec": "25kg", "price": 120}],
    )
    assert isinstance(data, bytes)
    assert len(data) > 100


def test_build_price_list_docx_bytes_no_template_raises() -> None:
    with pytest.raises(ValueError):
        build_price_list_docx_bytes(template_path=None)


def test_header_text_from_cell() -> None:
    cell = MagicMock()
    cell.text = "  单价  "
    assert _header_text(cell) == "单价"


# ---------------------------------------------------------------------------
# ai_chat_app_service_v2
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_ai_chat_v2_execute_command() -> None:
    svc = AiChatAppServiceV2()
    with patch.object(svc._bus, "publish") as mock_pub:
        out = await svc.execute_command("chat", {"message": "hi"})
    assert out["success"] is True
    assert "event_id" in out
    mock_pub.assert_called_once()


@pytest.mark.asyncio
async def test_ai_chat_v2_execute_command_failure() -> None:
    svc = AiChatAppServiceV2()
    with patch.object(svc._bus, "publish", side_effect=RuntimeError("bus down")):
        out = await svc.execute_command("chat", {})
    assert out["success"] is False


def test_get_ai_chat_app_service_v2_singleton() -> None:
    a = get_ai_chat_app_service_v2()
    b = get_ai_chat_app_service_v2()
    assert a is b


# ---------------------------------------------------------------------------
# service_bridge routes
# ---------------------------------------------------------------------------


@patch("app.fastapi_routes.service_bridge.get_db")
def test_service_bridge_config_get(mock_get_db: MagicMock, bridge_client: TestClient) -> None:
    ctx, _ = _mock_db_ctx()
    mock_get_db.return_value = ctx
    r = bridge_client.get("/api/service-bridge/config")
    assert r.status_code == 200
    assert r.json()["success"] is True
    assert "instance_id" in r.json()["data"]


@patch("app.fastapi_routes.service_bridge._set_config_value")
def test_service_bridge_config_put(mock_set: MagicMock, bridge_client: TestClient) -> None:
    r = bridge_client.put(
        "/api/service-bridge/config",
        json={"main_server_url": "http://main", "instance_name": "测试实例"},
    )
    assert r.status_code == 200
    assert r.json()["success"] is True
    assert mock_set.call_count >= 1


@patch("app.fastapi_routes.service_bridge.get_db")
def test_service_bridge_list_requests(mock_get_db: MagicMock, bridge_client: TestClient) -> None:
    ctx, db = _mock_db_ctx(query_result=[])
    mock_get_db.return_value = ctx
    r = bridge_client.get("/api/service-bridge/requests")
    assert r.status_code == 200
    assert r.json()["total"] == 0


@patch("app.fastapi_routes.service_bridge.get_db")
def test_service_bridge_get_request_not_found(
    mock_get_db: MagicMock, bridge_client: TestClient
) -> None:
    ctx, _ = _mock_db_ctx(query_result=None)
    mock_get_db.return_value = ctx
    r = bridge_client.get("/api/service-bridge/requests/999")
    assert r.status_code == 404


@patch("app.fastapi_routes.service_bridge.get_db")
def test_service_bridge_receive_request(mock_get_db: MagicMock, bridge_client: TestClient) -> None:
    ctx, db = _mock_db_ctx()
    mock_get_db.return_value = ctx
    req_obj = MagicMock()
    req_obj.to_dict.return_value = {"id": 1, "title": "help"}
    with patch("app.db.models.service_request.ServiceRequest", return_value=req_obj):
        r = bridge_client.post(
            "/api/service-bridge/requests",
            json={
                "source_instance_id": "inst-1",
                "source_instance_name": "A",
                "title": "help",
            },
        )
    assert r.status_code == 200
    assert r.json()["success"] is True


def test_service_bridge_ping_main(bridge_client: TestClient) -> None:
    with patch("app.fastapi_routes.service_bridge._get_config_value", return_value=""):
        r = bridge_client.get("/api/service-bridge/ping-main")
    assert r.status_code in (200, 502, 503)


# ---------------------------------------------------------------------------
# excel_extract routes
# ---------------------------------------------------------------------------


def test_excel_extract_missing_file_path(excel_client: TestClient) -> None:
    r = excel_client.post("/api/excel/data/extract", json={})
    assert r.status_code == 400


def test_excel_extract_from_path(tmp_path, excel_client: TestClient) -> None:
    p = _write_xlsx(tmp_path / "e.xlsx", [["产品", "单价"], ["漆", 99]])
    r = excel_client.post(
        "/api/excel/data/extract",
        json={"file_path": str(p), "header_row": 1},
    )
    assert r.status_code == 200
    assert r.json()["success"] is True
    assert r.json()["total_rows"] == 1


def test_excel_extract_test_endpoint(excel_client: TestClient) -> None:
    r = excel_client.get("/api/excel/data/extract/test")
    assert r.status_code == 200


def test_excel_generate_missing_data(excel_client: TestClient) -> None:
    r = excel_client.post("/api/excel/data/generate", json={})
    assert r.status_code == 400


def test_excel_generate_and_download(excel_client: TestClient, tmp_path) -> None:
    with patch("app.fastapi_routes.excel_extract.TEMP_EXCEL_DIR", str(tmp_path)):
        rows = [{"产品": "漆", "单价": 88}]
        r = excel_client.post(
            "/api/excel/data/generate",
            json={"data": rows, "filename": "out.xlsx"},
        )
        assert r.status_code == 200
        assert r.json()["success"] is True
        r2 = excel_client.post(
            "/api/excel/data/generate/download",
            json={"data": rows, "filename": "out.xlsx"},
        )
        assert r2.status_code == 200
        assert "spreadsheetml" in r2.headers.get("content-type", "")


def test_excel_extract_upload_invalid_ext(excel_client: TestClient) -> None:
    r = excel_client.post(
        "/api/excel/data/extract/upload",
        files={"excel_file": ("bad.txt", b"hello", "text/plain")},
    )
    assert r.status_code == 400


def test_excel_extract_upload_ok(excel_client: TestClient, tmp_path) -> None:
    buf = BytesIO()
    _write_xlsx(Path(buf.name) if False else tmp_path / "up.xlsx", [["a"], [1]])
    p = tmp_path / "up.xlsx"
    with patch("app.fastapi_routes.excel_extract.TEMP_EXCEL_DIR", str(tmp_path)):
        with open(p, "rb") as f:
            r = excel_client.post(
                "/api/excel/data/extract/upload",
                files={"excel_file": ("up.xlsx", f.read(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
    assert r.status_code == 200
    assert r.json()["success"] is True


# ---------------------------------------------------------------------------
# compat_db writes (mocked engine)
# ---------------------------------------------------------------------------


def _parse_price(v):
    return float(v or 0)


def _parse_qty(v):
    return int(v or 0)


def _parse_active(v):
    return bool(v) if v is not None else True


def test_products_pg_update_row_no_engine() -> None:
    with patch(
        "app.infrastructure.persistence.compat_db.writes.get_sync_engine",
        return_value=None,
    ):
        with pytest.raises(Exception):
            products_pg_update_row(
                1,
                {"name": "x"},
                parse_price=_parse_price,
                parse_quantity=_parse_qty,
                parse_is_active=_parse_active,
            )


@patch("app.infrastructure.persistence.compat_db.writes.get_sync_engine")
def test_products_pg_delete_row_mocked(mock_eng: MagicMock) -> None:
    eng = MagicMock()
    conn = MagicMock()
    eng.begin.return_value.__enter__ = MagicMock(return_value=conn)
    eng.begin.return_value.__exit__ = MagicMock(return_value=False)
    mock_eng.return_value = eng
    products_pg_delete_row(1)


@patch("app.infrastructure.persistence.compat_db.writes._customer_pg_fetch_by_id")
@patch("app.infrastructure.persistence.compat_db.writes._customer_pg_engine_insp")
def test_customer_pg_insert_mocked(mock_ei: MagicMock, mock_fetch: MagicMock) -> None:
    eng = MagicMock()
    conn = MagicMock()
    eng.connect.return_value.__enter__ = MagicMock(return_value=conn)
    eng.connect.return_value.__exit__ = MagicMock(return_value=False)
    insp = MagicMock()
    insp.get_columns.return_value = [
        {"name": "unit_name", "type": "VARCHAR"},
        {"name": "is_active", "type": "BOOLEAN"},
        {"name": "created_at"},
        {"name": "updated_at"},
    ]
    mock_ei.return_value = (eng, insp)
    conn.execute.return_value.first.return_value = None
    conn.execute.return_value.scalar_one.return_value = 7
    mock_fetch.return_value = {"id": 7, "unit_name": "甲公司"}
    out = _customer_pg_insert("甲公司", "张三", "138", "成都")
    assert out["id"] == 7
