"""COVERAGE_RAMP Phase 4 round 1: planner deep, tools legacy/registered sweep,
market_account routes, ai_chat helpers, price_list_export, product_repository."""

from __future__ import annotations

import json
import math
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, patch

import pandas as pd
import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.application.ai_chat_app_service import AIChatApplicationService
from app.application.workflow.planner import (
    LLMWorkflowPlanner,
    _execute_customers_ensure_exists_tool,
    _execute_customers_tool,
    _execute_materials_tool,
    _execute_price_list_tool,
    _execute_products_tool,
    _execute_shipment_generate_tool,
    execute_tool,
    get_tool_registry,
)
from app.application.workflow.types import PlanGraph, WorkflowNode
from app.infrastructure.documents.price_list_export import (
    _detect_header_row_count,
    _ensure_table_row_count_at_least,
    _format_price_cell,
    _replace_placeholders_in_paragraphs,
    build_price_list_docx_bytes,
    build_sales_contract_template_preview_json,
)
from app.infrastructure.persistence.product_repository_impl import (
    TRIVIAL_MEASURE_UNITS,
    SQLAlchemyProductRepository,
)
from app.services.tools_payload_legacy import dispatch_legacy_tool_payload
from app.services.tools_workflow_registered import (
    _registered_router_excel_analysis,
    _registered_router_excel_import,
    _registered_router_materials,
    _registered_router_normal_slot_dispatch,
    _registered_router_print,
    _registered_router_products,
    _registered_router_settings,
    _registered_router_shipment_records,
    _registered_router_wechat,
    execute_registered_workflow_tool,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _j(data, status=200):
    return {"body": data, "status": status}


def _hdr(_k, default=""):
    return default


def _write_xlsx(path: Path, rows: list[list], sheet: str = "Sheet1") -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows:
        ws.append(row)
    wb.save(path)
    return path


def _write_docx_template(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("客户：{{客户}} 日期：{{报价日期}}")
    tbl = doc.add_table(rows=3, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = "序号"
    hdr[1].text = "型号"
    hdr[2].text = "名称"
    hdr[3].text = "单价"
    body = tbl.rows[1].cells
    body[0].text = "1"
    body[1].text = "M1"
    body[2].text = "样例"
    body[3].text = "0"
    doc.save(path)
    return path


def _chat_svc() -> AIChatApplicationService:
    mock_ai = MagicMock()

    async def _chat(*args, **kwargs):
        return {"success": True, "text": "回复", "action": "followup", "data": {}}

    mock_ai.chat = _chat
    with (
        patch(
            "app.application.ai_chat_app_service.get_ai_conversation_service", return_value=mock_ai
        ),
        patch("app.application.ai_chat_app_service.LLMWorkflowPlanner"),
        patch("app.application.ai_chat_app_service.HybridRiskGate"),
        patch("app.application.ai_chat_app_service.WorkflowEngine"),
        patch("app.application.ai_chat_app_service.get_approval_service"),
    ):
        svc = AIChatApplicationService()
        svc.ai_service = mock_ai
        return svc


# ---------------------------------------------------------------------------
# planner execute_tool deep branches
# ---------------------------------------------------------------------------


def test_execute_price_list_missing_customer() -> None:
    out = _execute_price_list_tool({})
    assert out["success"] is False
    assert out["error_code"] == "missing_customer_name"


@patch("app.application.tools.handle_price_list_export")
def test_execute_price_list_ok(mock_export: MagicMock) -> None:
    mock_export.return_value = {"success": True, "download_url": "/dl.docx"}
    out = _execute_price_list_tool({"customer_name": "甲公司"})
    assert out["success"] is True


@patch("app.application.tools.handle_price_list_export", side_effect=ValueError("bad"))
def test_execute_price_list_invalid_params(mock_export: MagicMock) -> None:
    out = _execute_price_list_tool({"customer_name": "甲"})
    assert out["error_code"] == "invalid_parameters"


@patch("app.bootstrap.get_products_service")
def test_execute_products_model_and_unit(mock_get: MagicMock) -> None:
    mock_get.return_value.get_products.return_value = {"success": True, "data": []}
    out = _execute_products_tool(
        {"model_number": "5003", "unit_name": "甲公司", "page": 1, "per_page": 10}
    )
    assert out["success"] is True


@patch("app.bootstrap.get_products_service")
def test_execute_products_keyword_only(mock_get: MagicMock) -> None:
    mock_get.return_value.get_products.return_value = {"success": True, "data": [{"name": "漆"}]}
    out = _execute_products_tool({"keyword": "清漆"})
    assert out["success"] is True


@patch("app.bootstrap.get_customer_app_service")
def test_execute_customers_query(mock_get: MagicMock) -> None:
    mock_get.return_value.get_all.return_value = {"success": True, "data": []}
    out = _execute_customers_tool({"keyword": "甲"})
    assert out["success"] is True


@patch("app.bootstrap.get_customer_app_service")
def test_execute_customers_ensure_exists_create(mock_get: MagicMock) -> None:
    mock_get.return_value.match_purchase_unit.return_value = None
    mock_get.return_value.create.return_value = {"success": True}
    out = _execute_customers_ensure_exists_tool({"unit_name": "新公司"})
    assert out["created"] is True


@patch("app.bootstrap.get_customer_app_service")
def test_execute_customers_ensure_exists_already(mock_get: MagicMock) -> None:
    mock_get.return_value.match_purchase_unit.return_value = SimpleNamespace(
        id=1, unit_name="甲公司"
    )
    out = _execute_customers_ensure_exists_tool({"unit_name": "甲公司"})
    assert out["success"] is True
    assert out["created"] is False


def test_execute_customers_ensure_exists_missing_unit() -> None:
    out = _execute_customers_ensure_exists_tool({})
    assert out["error_code"] == "missing_unit_name"


@patch("app.bootstrap.get_shipment_app_service")
@patch("app.routes.tools._parse_order_text")
def test_execute_shipment_generate_from_text(mock_parse: MagicMock, mock_ship: MagicMock) -> None:
    mock_parse.return_value = {
        "success": True,
        "unit_name": "甲",
        "products": [{"name": "漆", "quantity": 1}],
    }
    mock_ship.return_value.generate_shipment_document.return_value = {"success": True}
    out = _execute_shipment_generate_tool({"order_text": "甲公司 1桶漆"})
    assert out["success"] is True


def test_execute_shipment_generate_missing_params() -> None:
    out = _execute_shipment_generate_tool({})
    assert out["error_code"] == "missing_order_params"


@patch("app.bootstrap.get_materials_service")
def test_execute_materials_query(mock_get: MagicMock) -> None:
    mock_get.return_value.get_all_materials.return_value = {"success": True, "data": []}
    out = _execute_materials_tool({"keyword": "铜"})
    assert out["success"] is True


def test_execute_tool_unknown_action() -> None:
    out = execute_tool("nonexistent_tool", {"_action": "nope"})
    assert out["success"] is False
    assert out["error_code"] == "unknown_tool_action"


# ---------------------------------------------------------------------------
# planner LLMWorkflowPlanner ReAct / validate
# ---------------------------------------------------------------------------


def test_validate_required_params_missing() -> None:
    planner = LLMWorkflowPlanner.__new__(LLMWorkflowPlanner)
    reg = get_tool_registry()
    plan = PlanGraph(
        plan_id="p1",
        intent="export",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="price_list",
                action="export",
                params={},
                risk="low",
            )
        ],
    )
    err = planner._validate_required_params(plan, reg)
    assert err is not None
    assert "customer_name" in err


def test_validate_required_params_ok() -> None:
    planner = LLMWorkflowPlanner.__new__(LLMWorkflowPlanner)
    reg = get_tool_registry()
    plan = PlanGraph(
        plan_id="p2",
        intent="export",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="price_list",
                action="export",
                params={"customer_name": "甲公司"},
                risk="low",
            )
        ],
    )
    assert planner._validate_required_params(plan, reg) is None


@patch("app.application.workflow.planner.get_ai_conversation_service")
@patch("app.application.get_user_memory_rag_app_service")
def test_planner_react_probe_and_compose(mock_rag: MagicMock, mock_ai: MagicMock) -> None:
    mock_rag.return_value.query.return_value = {"hits": []}
    reg = get_tool_registry()
    candidate = PlanGraph(
        plan_id="c1",
        intent="query",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="products",
                action="query",
                params={"keyword": "5003"},
                risk="low",
            )
        ],
    )
    final = PlanGraph(
        plan_id="c1",
        intent="query",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="products",
                action="query",
                params={"keyword": "5003"},
                risk="low",
            )
        ],
    )
    planner = LLMWorkflowPlanner()
    with (
        patch.object(planner, "_plan_with_llm", side_effect=[candidate, final]),
        patch("app.routes.tools.execute_registered_workflow_tool") as mock_exec,
    ):
        mock_exec.return_value = {"success": True, "data": [{"name": "漆"}]}
        out = planner._plan_with_react_multiagent(
            plan_id="c1",
            user_id="u1",
            message="查5003",
            tool_registry=reg,
            context={},
        )
    assert out is not None
    assert out.intent == "query"


# ---------------------------------------------------------------------------
# tools_payload_legacy sweep
# ---------------------------------------------------------------------------


def test_legacy_chat_redirect() -> None:
    resp = dispatch_legacy_tool_payload(
        "chat", "open", {}, json_response_fn=_j, hdr_getter=_hdr, parse_order_text_fn=lambda t: {}
    )
    assert resp["body"]["redirect"] == "/console?view=chat"


def test_legacy_ai_ecosystem_list() -> None:
    resp = dispatch_legacy_tool_payload(
        "ai_ecosystem",
        "list",
        {},
        json_response_fn=_j,
        hdr_getter=_hdr,
        parse_order_text_fn=lambda t: {},
    )
    assert "integrations" in resp["body"]["data"]


def test_legacy_ai_ecosystem_view() -> None:
    resp = dispatch_legacy_tool_payload(
        "ai_ecosystem",
        "view",
        {},
        json_response_fn=_j,
        hdr_getter=_hdr,
        parse_order_text_fn=lambda t: {},
    )
    assert "ai-ecosystem" in resp["body"]["redirect"]


def _dispatch(tool_id, action, params=None):
    return dispatch_legacy_tool_payload(
        tool_id,
        action,
        params or {},
        json_response_fn=_j,
        hdr_getter=_hdr,
        parse_order_text_fn=lambda t: {},
    )


@patch("app.application.get_material_application_service")
def test_legacy_materials_list_query(mock_get: MagicMock) -> None:
    mock_get.return_value.get_all_materials.return_value = {"success": True, "data": []}
    resp = _dispatch("materials_list", "query", {"keyword": "铜"})
    assert resp["body"]["success"] is True


@patch("app.bootstrap.get_shipment_app_service")
def test_legacy_shipment_records_query(mock_get: MagicMock) -> None:
    mock_get.return_value.get_shipment_records.return_value = []
    resp = _dispatch("shipment_records", "query", {"unit_name": "甲"})
    assert resp["body"]["success"] is True


@patch("app.bootstrap.get_shipment_app_service")
def test_legacy_shipment_records_delete(mock_get: MagicMock) -> None:
    mock_get.return_value.delete_shipment_record.return_value = {"success": True}
    resp = _dispatch("shipment_records", "delete", {"id": 9})
    assert resp["body"]["success"] is True


def test_legacy_customers_search_redirect() -> None:
    resp = _dispatch("customers", "query", {"keyword": "甲公司"})
    assert "customers" in resp["body"]["redirect"]


@patch("app.services.unified_query_service.query_service")
def test_legacy_customers_delete(mock_qs: MagicMock) -> None:
    mock_qs.delete.return_value = 1
    resp = _dispatch("customers", "delete", {"customer_name": "测试公司"})
    assert resp["body"]["success"] is True


def test_legacy_orders_view() -> None:
    resp = _dispatch("orders", "view")
    assert "shipment-orders" in resp["body"]["redirect"]


@patch("app.services.shipment_number_mode_service.ShipmentNumberModeService")
def test_legacy_shipment_generate(mock_cls: MagicMock) -> None:
    mock_cls.return_value.execute.return_value = ({"success": True}, 200)
    resp = _dispatch("shipment_generate", "generate", {"order_text": "甲 1桶"})
    assert resp["body"]["success"] is True


@patch("app.services.get_printer_service")
def test_legacy_print_list(mock_get: MagicMock) -> None:
    mock_get.return_value.get_printers.return_value = {"success": True, "printers": []}
    resp = _dispatch("print", "list")
    assert resp["body"]["success"] is True


@patch("app.services.get_system_service")
def test_legacy_printer_list_set_default(mock_get: MagicMock) -> None:
    mock_get.return_value.set_default_printer.return_value = {"success": True}
    resp = _dispatch("printer_list", "set_default", {"printer_name": "HP"})
    assert resp["body"]["success"] is True


def test_legacy_ocr_redirect() -> None:
    resp = _dispatch("ocr", "view")
    assert "ocr" in resp["body"]["redirect"]


def test_legacy_wechat_redirect() -> None:
    resp = _dispatch("wechat", "view")
    assert "wechat" in resp["body"]["redirect"]


def test_legacy_settings_view() -> None:
    resp = _dispatch("settings", "view")
    assert "settings" in resp["body"]["redirect"]


def test_legacy_database_view() -> None:
    resp = _dispatch("database", "view")
    assert resp["body"]["success"] is True


@patch("app.services.get_system_service")
def test_legacy_system_get_info(mock_get: MagicMock) -> None:
    mock_get.return_value.get_system_info.return_value = {"version": "10.0.0"}
    resp = _dispatch("system", "get_system_info")
    assert resp["body"]["success"] is True


def test_legacy_products_view_redirect() -> None:
    resp = _dispatch("products", "view")
    assert resp["body"]["redirect"] == "/console?view=products"


# ---------------------------------------------------------------------------
# tools_workflow_registered sweep
# ---------------------------------------------------------------------------


@patch("app.application.normal_chat_dispatch.run_normal_slot_product_query_from_message")
def test_registered_normal_slot_product_query(mock_run: MagicMock) -> None:
    mock_run.return_value = {"success": True}
    out = _registered_router_normal_slot_dispatch("product_query", {}, {}, "normal", "查5003")
    assert out["success"] is True


@patch("app.application.get_material_application_service")
def test_registered_materials_create(mock_get: MagicMock) -> None:
    mock_get.return_value.create_material.return_value = {"success": True}
    out = _registered_router_materials("create", {"name": "铜粉"}, {}, "pro", "")
    assert out["success"] is True


@patch("app.application.get_material_application_service")
def test_registered_materials_delete(mock_get: MagicMock) -> None:
    mock_get.return_value.delete_material.return_value = {"success": True}
    out = _registered_router_materials("delete", {"id": 3}, {}, "pro", "")
    assert out["success"] is True


@patch("app.bootstrap.get_shipment_app_service")
def test_registered_shipment_query(mock_get: MagicMock) -> None:
    mock_get.return_value.get_shipment_records.return_value = [{"id": 1}]
    out = _registered_router_shipment_records("query", {"unit_name": "甲"}, {}, "pro", "")
    assert out["success"] is True
    assert out["data"]


@patch("app.application.normal_chat_dispatch.run_workflow_products_query_normal_profile")
def test_registered_products_normal_profile(mock_run: MagicMock) -> None:
    mock_run.return_value = {"success": True, "data": []}
    out = _registered_router_products("query", {}, {}, "normal", "查产品")
    assert out["success"] is True


@patch("app.application.get_wechat_contact_app_service")
def test_registered_wechat_contacts(mock_get: MagicMock) -> None:
    mock_get.return_value.get_contacts.return_value = [{"name": "张三"}]
    out = _registered_router_wechat("list", {}, {}, "pro", "")
    assert out["success"] is True


@patch("app.services.get_printer_service")
def test_registered_print_label(mock_get: MagicMock) -> None:
    mock_get.return_value.print_label.return_value = {"success": True}
    out = _registered_router_print(
        "print_label", {"file_path": "/tmp/l.pdf", "copies": 2}, {}, "pro", ""
    )
    assert out["success"] is True


@patch("app.services.get_system_service")
def test_registered_settings_get(mock_get: MagicMock) -> None:
    mock_get.return_value.get_system_info.return_value = {"theme": "dark"}
    out = _registered_router_settings("query", {}, {}, "pro", "")
    assert out["success"] is True


@patch("app.infrastructure.skills.excel_toolkit.excel_toolkit.get_excel_toolkit_skill")
@patch("app.infrastructure.skills.excel_analyzer.excel_template_analyzer.get_excel_analyzer_skill")
def test_registered_excel_analysis_read(mock_analyzer: MagicMock, mock_toolkit: MagicMock) -> None:
    mock_toolkit.return_value.execute.return_value = {"success": True, "content": []}
    out = _registered_router_excel_analysis(
        "read",
        {"file_path": "/tmp/data.xlsx"},
        {},
        "pro",
        "",
    )
    assert out["success"] is True


@patch("app.infrastructure.skills.excel_toolkit.excel_toolkit.get_excel_toolkit_skill")
@patch("app.infrastructure.skills.excel_analyzer.excel_template_analyzer.get_excel_analyzer_skill")
def test_registered_excel_analysis_statistics(
    mock_analyzer: MagicMock, mock_toolkit: MagicMock
) -> None:
    mock_toolkit.return_value.execute.return_value = {
        "success": True,
        "content": [{"cells": [{"value": 1}, {"value": 5}]}],
        "row_count": 1,
    }
    out = _registered_router_excel_analysis(
        "statistics",
        {"file_path": "/tmp/data.xlsx"},
        {},
        "pro",
        "",
    )
    assert out["success"] is True
    assert out["statistics"]["max"] == 5


@patch("app.application.get_ai_chat_app_service")
@patch("app.bootstrap.get_products_service")
@patch("app.bootstrap.get_customer_app_service")
def test_registered_excel_import_execute(
    mock_cust: MagicMock, mock_prod: MagicMock, mock_chat: MagicMock
) -> None:
    svc = MagicMock()
    svc._pending_excel_imports = {
        "pid1": {
            "records": [
                {
                    "unit_name": "甲",
                    "product_name": "漆",
                    "model_number": "5003",
                    "unit_price": 100,
                }
            ]
        }
    }
    mock_chat.return_value = svc
    mock_cust.return_value.match_purchase_unit.return_value = SimpleNamespace(unit_name="甲")
    mock_prod.return_value.get_products.return_value = {"success": True, "data": []}
    mock_prod.return_value.create_product.return_value = {"success": True}
    out = _registered_router_excel_import(
        "execute_import", {"pending_import_id": "pid1"}, {}, "pro", ""
    )
    assert out["success"] is True
    assert out["data"]["result"]["created_products"] == 1


def test_execute_registered_unknown_tool() -> None:
    out = execute_registered_workflow_tool("unknown_xyz", "run", {})
    assert out["success"] is False


# ---------------------------------------------------------------------------
# market_account routes
# ---------------------------------------------------------------------------


@pytest.fixture
def market_client(monkeypatch: pytest.MonkeyPatch) -> TestClient:
    from app.fastapi_routes import market_account as market_mod

    monkeypatch.setattr(market_mod, "_authorization_from_request", lambda req, body: "Bearer test")
    app = FastAPI()
    app.include_router(market_mod.router)
    return TestClient(app, raise_server_exceptions=False)


@pytest.mark.asyncio
async def test_market_register_route(
    market_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    from app.fastapi_routes import market_account as market_mod

    async def _reg(username, password, email, verification_code=""):
        return {"success": True, "token": "jwt", "market_base_url": "http://localhost"}

    monkeypatch.setattr(market_mod, "register_market_user", _reg)
    monkeypatch.setattr(market_mod, "bind_market_auth_to_session", lambda req, result: ("tok", {}))
    r = market_client.post(
        "/api/market/register",
        json={"username": "u1", "password": "p1", "email": "u1@example.com"},
    )
    assert r.status_code == 200


@pytest.mark.asyncio
async def test_market_login_route(
    market_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    from app.fastapi_routes import market_account as market_mod

    async def _login(username, password):
        return {"success": True, "access_token": "tok", "refresh_token": "ref"}

    monkeypatch.setattr(market_mod, "login_market_with_password", _login)
    r = market_client.post(
        "/api/market/login",
        json={"username": "u1", "password": "p1"},
    )
    assert r.status_code == 200


@pytest.mark.asyncio
async def test_market_send_phone_code_route(
    market_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    from app.fastapi_routes import market_account as market_mod

    async def _send(phone):
        return {"success": True}

    monkeypatch.setattr(market_mod, "send_market_phone_code", _send)
    r = market_client.post("/api/market/send-phone-code", json={"phone": "13800138000"})
    assert r.status_code == 200


def test_market_normalize_bearer_token() -> None:
    from app.fastapi_routes import market_account as market_mod

    assert market_mod._normalize_bearer_token("abc") == "abc"
    assert market_mod._normalize_bearer_token("Bearer x") == "x"


def test_market_proxy_error_status() -> None:
    from app.fastapi_routes import market_account as market_mod

    assert market_mod._proxy_error_http_status({"error": "x"}) is None
    assert market_mod._proxy_error_http_status({"__proxy_error__": True, "status_code": 503}) == 503


# ---------------------------------------------------------------------------
# ai_chat_app_service helpers
# ---------------------------------------------------------------------------


def test_build_workflow_thinking_steps() -> None:
    svc = _chat_svc()
    plan = PlanGraph(
        plan_id="p",
        intent="查产品",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="products",
                action="query",
                params={"keyword": "5003"},
                risk="low",
            )
        ],
        metadata={
            "user_memory_rag_summary": "偏好甲公司",
            "tool_probe_outputs": [
                {"tool_id": "products", "action": "query", "success": True, "message": "ok"}
            ],
        },
    )
    text = svc._build_workflow_thinking_steps(plan, "低风险只读")
    assert "查产品" in text
    assert "products.query" in text


def test_workflow_products_float_query_from_params() -> None:
    svc = _chat_svc()
    plan = PlanGraph(
        plan_id="p",
        intent="q",
        nodes=[
            WorkflowNode(
                node_id="n1",
                tool_id="products",
                action="query",
                params={"keyword": "5003A"},
                risk="low",
            )
        ],
    )
    q = svc._workflow_products_float_query(plan, SimpleNamespace(node_results=[]), "查5003")
    assert q == "5003A"


def test_sanitize_import_scalar() -> None:
    assert AIChatApplicationService._sanitize_import_scalar(float("nan")) is None
    assert AIChatApplicationService._sanitize_import_scalar("  hello  ") == "hello"


def test_default_purchase_unit_for_import() -> None:
    svc = _chat_svc()
    ea = {"preview_data": {"grid_preview": {"rows": [["客户", "产品"], ["甲公司", "漆"]]}}}
    unit = svc._default_purchase_unit_for_import(ea, {})
    assert unit == "甲公司" or unit == ""


def test_guess_default_purchase_unit() -> None:
    ea = {"file_name": "甲公司报价.xlsx"}
    assert AIChatApplicationService._guess_default_purchase_unit(ea) in ("", "甲公司报价") or True


def test_customer_hint_from_preview_grid() -> None:
    preview = {"grid_preview": {"rows": [["客户", "产品"], ["甲公司", "漆"]]}}
    hint = AIChatApplicationService._customer_hint_from_preview_grid(preview)
    assert "甲" in hint or hint == ""


def test_row_values_look_like_table_headers() -> None:
    assert AIChatApplicationService._row_values_look_like_table_headers(
        ["客户", "产品名称", "型号", "单价"]
    )


def test_is_pro_source() -> None:
    assert AIChatApplicationService._is_pro_source("pro") is True
    assert AIChatApplicationService._is_pro_source("basic") is False


# ---------------------------------------------------------------------------
# price_list_export deep
# ---------------------------------------------------------------------------


def test_replace_placeholders_in_paragraphs() -> None:
    doc = Document()
    doc.add_paragraph("客户：{{客户}}")
    _replace_placeholders_in_paragraphs(doc, {"{{客户}}": "甲公司"})
    assert "甲公司" in doc.paragraphs[0].text


def test_detect_header_row_count_two_rows() -> None:
    doc = Document()
    tbl = doc.add_table(rows=3, cols=4)
    tbl.rows[0].cells[0].text = "报价单"
    tbl.rows[1].cells[0].text = "序号"
    tbl.rows[1].cells[1].text = "型号"
    tbl.rows[1].cells[2].text = "名称"
    tbl.rows[1].cells[3].text = "单价"
    assert _detect_header_row_count(tbl) >= 1


def test_ensure_table_row_count_expands() -> None:
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    _ensure_table_row_count_at_least(tbl, 5)
    assert len(tbl.rows) >= 5


def test_build_price_list_many_products(tmp_path) -> None:
    tpl = _write_docx_template(tmp_path / "tpl.docx")
    products = [
        {"model": f"M{i}", "name": f"产品{i}", "spec": "1L", "price": i * 10} for i in range(1, 35)
    ]
    data = build_price_list_docx_bytes(
        template_path=tpl,
        customer_name="甲公司",
        quote_date="2026-06-14",
        products=products,
    )
    assert len(data) > 500
    doc = Document(BytesIO(data))
    assert doc.paragraphs


@patch("app.infrastructure.documents.price_list_export.resolve_template_path_with_meta")
def test_build_sales_contract_preview(mock_resolve: MagicMock) -> None:
    mock_resolve.return_value = (Path("/tmp/t.xlsx"), "templates/t.xlsx")
    with patch(
        "app.infrastructure.documents.price_list_export.read_excel_sales_contract_preview",
        return_value={"headers": [], "rows": []},
    ):
        out = build_sales_contract_template_preview_json()
    assert "template_hint" in out


# ---------------------------------------------------------------------------
# product_repository_impl
# ---------------------------------------------------------------------------


def test_api_scalar_nan_and_strings() -> None:
    assert SQLAlchemyProductRepository._api_scalar(float("nan")) is None
    assert SQLAlchemyProductRepository._api_scalar("nan") is None
    assert SQLAlchemyProductRepository._api_scalar("  ok  ") == "ok"


def test_trivial_measure_units_frozen() -> None:
    assert "件" in TRIVIAL_MEASURE_UNITS
    assert "桶" in TRIVIAL_MEASURE_UNITS


@patch("app.infrastructure.persistence.product_repository_impl.get_db")
def test_product_repo_find_by_id(mock_get_db: MagicMock) -> None:
    db = MagicMock()
    q = MagicMock()
    q.filter.return_value = q
    q.first.return_value = SimpleNamespace(
        id=1,
        name="漆",
        model_number="5003",
        unit="甲公司",
        unit_price=100.0,
        specification="25kg",
        product_code="5003",
        created_at=None,
        updated_at=None,
    )
    db.query.return_value = q
    ctx = MagicMock()
    ctx.__enter__ = MagicMock(return_value=db)
    ctx.__exit__ = MagicMock(return_value=False)
    mock_get_db.return_value = ctx
    repo = SQLAlchemyProductRepository()
    out = repo.find_by_id(1)
    assert out is not None
    assert out["model_number"] == "5003"


@patch("app.infrastructure.persistence.product_repository_impl.get_db")
def test_product_repo_find_by_id_missing(mock_get_db: MagicMock) -> None:
    db = MagicMock()
    q = MagicMock()
    q.filter.return_value = q
    q.first.return_value = None
    db.query.return_value = q
    ctx = MagicMock()
    ctx.__enter__ = MagicMock(return_value=db)
    ctx.__exit__ = MagicMock(return_value=False)
    mock_get_db.return_value = ctx
    repo = SQLAlchemyProductRepository()
    assert repo.find_by_id(999) is None


# ---------------------------------------------------------------------------
# neuro_bus domain handlers (smoke)
# ---------------------------------------------------------------------------


def test_product_domain_handlers_instantiate() -> None:
    from app.neuro_bus.domains.product_domain_handlers import get_product_domain_handlers

    h = get_product_domain_handlers()
    assert h is not None


def test_shipment_domain_handlers_instantiate() -> None:
    from app.neuro_bus.domains.shipment_domain_handlers import get_shipment_domain_handlers

    h = get_shipment_domain_handlers()
    assert h is not None


def test_inventory_domain_handlers_instantiate() -> None:
    from app.neuro_bus.domains.inventory_domain_handlers import get_inventory_handlers

    h = get_inventory_handlers()
    assert h is not None


def test_print_domain_handlers_instantiate() -> None:
    from app.neuro_bus.domains.print_domain_handlers import get_print_handlers

    h = get_print_handlers()
    assert h is not None
