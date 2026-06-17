"""Tests for app.infrastructure.documents.price_list_export."""
from __future__ import annotations

import os
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from app.infrastructure.documents.price_list_export import (
    _border_el_effective,
    _detect_header_row_count,
    _fill_first_table_with_products,
    _format_price_cell,
    _header_text,
    _parse_header_serial_and_column_map,
    _product_row_cell_values,
    _remove_table_row,
    _replace_placeholders_in_paragraphs,
    _row_keyword_score,
    _snapshot_tbl_borders,
    _tbl_row_count,
    build_price_list_docx_bytes,
    build_price_list_template_preview_json,
    build_sales_contract_template_preview_json,
    resolve_price_list_docx_template,
)


# ---------------------------------------------------------------------------
# _format_price_cell
# ---------------------------------------------------------------------------
class TestFormatPriceCell:
    def test_integer_value(self):
        assert _format_price_cell(100) == "100"

    def test_float_value(self):
        assert _format_price_cell(99.99) == "99.99"

    def test_none_returns_empty(self):
        assert _format_price_cell(None) == ""

    def test_empty_string_returns_empty(self):
        assert _format_price_cell("") == ""

    def test_string_number(self):
        assert _format_price_cell("50.5") == "50.50"

    def test_non_numeric_string(self):
        assert _format_price_cell("abc") == "abc"

    def test_exact_integer_float(self):
        assert _format_price_cell(100.0) == "100"

    def test_zero(self):
        assert _format_price_cell(0) == "0"

    def test_negative(self):
        assert _format_price_cell(-10.5) == "-10.50"


# ---------------------------------------------------------------------------
# _product_row_cell_values
# ---------------------------------------------------------------------------
class TestProductRowCellValues:
    def test_standard_keys(self):
        prod = {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99.9}
        result = _product_row_cell_values(prod)
        assert result == ["M1", "Widget", "10x20", "99.90"]

    def test_chinese_keys(self):
        prod = {"型号": "M2", "产品名称": "Gadget", "规格": "5x10", "单价": 50}
        result = _product_row_cell_values(prod)
        assert result == ["M2", "Gadget", "5x10", "50"]

    def test_mixed_keys(self):
        prod = {"model_number": "M3", "名称": "Thing", "spec": "1x1", "unit_price": 25}
        result = _product_row_cell_values(prod)
        assert result == ["M3", "Thing", "1x1", "25"]

    def test_missing_keys(self):
        prod = {}
        result = _product_row_cell_values(prod)
        assert result == ["", "", "", ""]


# ---------------------------------------------------------------------------
# _detect_header_row_count
# ---------------------------------------------------------------------------
class TestDetectHeaderRowCount:
    def test_single_row_header(self):
        table = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(text="型号"), MagicMock(text="名称")]
        row1 = MagicMock()
        row1.cells = [MagicMock(text="M1"), MagicMock(text="Widget")]
        table.rows = [row0, row1]
        assert _detect_header_row_count(table) == 1

    def test_two_row_header(self):
        table = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(text="产品信息")]
        row1 = MagicMock()
        row1.cells = [MagicMock(text="型号"), MagicMock(text="名称"), MagicMock(text="规格"), MagicMock(text="单价")]
        table.rows = [row0, row1]
        assert _detect_header_row_count(table) == 2

    def test_single_row_table(self):
        table = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(text="型号")]
        table.rows = [row0]
        assert _detect_header_row_count(table) == 1


# ---------------------------------------------------------------------------
# _parse_header_serial_and_column_map
# ---------------------------------------------------------------------------
class TestParseHeaderSerialAndColumnMap:
    def test_with_serial_column(self):
        cells = [MagicMock(text="序号"), MagicMock(text="型号"), MagicMock(text="名称"), MagicMock(text="规格"), MagicMock(text="单价")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert "model" in col_map
        assert "price" in col_map

    def test_without_serial_column(self):
        cells = [MagicMock(text="型号"), MagicMock(text="名称"), MagicMock(text="规格"), MagicMock(text="单价")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is False
        assert col_map["model"] == 0
        assert col_map["price"] == 3

    def test_empty_cells(self):
        with_serial, col_map = _parse_header_serial_and_column_map([])
        assert with_serial is False
        assert "model" in col_map

    def test_fallback_with_serial_and_5_columns(self):
        cells = [MagicMock(text="#"), MagicMock(text=""), MagicMock(text=""), MagicMock(text=""), MagicMock(text="")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert col_map["model"] == 1


# ---------------------------------------------------------------------------
# _border_el_effective
# ---------------------------------------------------------------------------
class TestBorderElEffective:
    def test_none_returns_false(self):
        assert _border_el_effective(None) is False

    def test_nil_val_returns_false(self):
        el = MagicMock()
        el.get.return_value = "nil"
        assert _border_el_effective(el) is False

    def test_none_val_returns_false(self):
        el = MagicMock()
        el.get.return_value = None
        assert _border_el_effective(el) is False

    def test_effective_border(self):
        el = MagicMock()
        el.get.return_value = "single"
        assert _border_el_effective(el) is True


# ---------------------------------------------------------------------------
# _header_text
# ---------------------------------------------------------------------------
class TestHeaderText:
    def test_returns_text(self):
        cell = MagicMock()
        cell.text = "  型号  "
        assert _header_text(cell) == "型号"

    def test_none_text(self):
        cell = MagicMock()
        cell.text = None
        assert _header_text(cell) == ""


# ---------------------------------------------------------------------------
# build_price_list_template_preview_json
# ---------------------------------------------------------------------------
class TestBuildPriceListTemplatePreviewJson:
    def test_returns_preview(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=(Path("/tmp/template.docx"), "template.docx"),
        ):
            result = build_price_list_template_preview_json()
        assert result["success"] is True
        assert "headers" in result


# ---------------------------------------------------------------------------
# resolve_price_list_docx_template
# ---------------------------------------------------------------------------
class TestResolvePriceListDocxTemplate:
    def test_returns_path_and_rel(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=(Path("/tmp/t.docx"), "t.docx"),
        ):
            path, rel = resolve_price_list_docx_template()
        assert str(path) == "/tmp/t.docx"
        assert rel == "t.docx"


# ---------------------------------------------------------------------------
# build_sales_contract_template_preview_json
# ---------------------------------------------------------------------------
class TestBuildSalesContractTemplatePreviewJson:
    def test_returns_preview(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=(Path("/tmp/contract.docx"), "contract.docx"),
        ), patch(
            "app.infrastructure.documents.price_list_export.read_excel_sales_contract_preview",
            return_value={"success": True, "headers": []},
        ):
            result = build_sales_contract_template_preview_json()
        assert "template_hint" in result


# ---------------------------------------------------------------------------
# build_price_list_docx_bytes
# ---------------------------------------------------------------------------
class TestBuildPriceListDocxBytes:
    def test_raises_on_no_template_path(self):
        with pytest.raises(ValueError, match="template_path"):
            build_price_list_docx_bytes()

    def test_raises_on_missing_file(self):
        with pytest.raises(FileNotFoundError):
            build_price_list_docx_bytes(template_path="/nonexistent/file.docx")

    def test_generates_docx_bytes(self, tmp_path):
        from docx import Document

        # Create a minimal docx with a table
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        data_row = table.rows[1].cells
        for c in data_row:
            c.text = ""

        template_path = tmp_path / "test_template.docx"
        doc.save(str(template_path))

        products = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99.9},
        ]
        result = build_price_list_docx_bytes(
            template_path=template_path,
            customer_name="TestCo",
            quote_date="2026-01-01",
            products=products,
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generates_docx_with_no_table(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("No table here")
        template_path = tmp_path / "no_table.docx"
        doc.save(str(template_path))

        products = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99.9},
        ]
        result = build_price_list_docx_bytes(
            template_path=template_path,
            products=products,
        )
        assert isinstance(result, bytes)

    def test_generates_docx_with_empty_products(self, tmp_path):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"

        template_path = tmp_path / "empty.docx"
        doc.save(str(template_path))

        result = build_price_list_docx_bytes(
            template_path=template_path,
            products=[],
        )
        assert isinstance(result, bytes)

    def test_uses_rows_parameter(self, tmp_path):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        for c in table.rows[1].cells:
            c.text = ""

        template_path = tmp_path / "rows.docx"
        doc.save(str(template_path))

        rows = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 50},
        ]
        result = build_price_list_docx_bytes(
            template_path=template_path,
            rows=rows,
        )
        assert isinstance(result, bytes)

    def test_template_path_arg_parameter(self, tmp_path):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"

        template_path = tmp_path / "arg.docx"
        doc.save(str(template_path))

        result = build_price_list_docx_bytes(
            template_path_arg=template_path,
            products=[],
        )
        assert isinstance(result, bytes)
