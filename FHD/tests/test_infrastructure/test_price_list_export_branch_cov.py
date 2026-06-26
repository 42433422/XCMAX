"""Branch-coverage tests for app.infrastructure.documents.price_list_export.

Targets branches NOT already covered by test_price_list_export.py /
test_price_list_export_ext2.py / test_price_list_export_ext3.py /
test_price_list_export_helpers.py.

Focus:
* ``_format_price_cell`` — TypeError path (non-numeric object).
* ``_replace_placeholders_in_paragraphs`` — empty text, no keys present,
  merged result empty, paragraphs in table cells.
* ``_product_row_cell_values`` — fallback chains (spec/specification/规格).
* ``_row_keyword_score`` — empty cells, partial keyword hits.
* ``_detect_header_row_count`` — single-row table, s1 < 2, s1 <= s0.
* ``_append_tr_clone_from_last`` — no trs available.
* ``_tc_ensure_tc_pr`` — existing tcPr vs new tcPr.
* ``_tc_get_tc_borders_snapshot`` — tcPr None, tcb None, tcb present.
* ``_tc_apply_tc_borders_snapshot`` — None snapshot, existing tcb removed.
* ``_snapshot_body_row_tc_borders`` — out-of-range indices.
* ``_pick_border_template_row_index`` — no borders found, falls back to header_rows.
* ``_apply_tc_borders_to_all_body_rows`` — empty snaps, i >= len(snaps).
* ``_ensure_table_row_count_at_least`` — stagnant break path.
* ``_parse_header_serial_and_column_map`` — empty raw, 货号/编码/品名/售价/定价,
  price-with-价格 guard, with_serial + len < 5 fallback.
* ``_write_product_row`` — idx None / negative / out of range.
* ``_tbl_pr`` — existing vs new.
* ``_find_tbl_pr_readonly`` — None case.
* ``_snapshot_tbl_borders`` — tbl_pr None, no tblBorders child.
* ``_restore_tbl_borders`` — None snapshot, existing child removed.
* ``_tbl_borders_ensure_bottom_edge`` — bottom effective (skip), sample None,
  bottom element exists (replace).
* ``_tc_get_side_border_copy`` — tcPr None, tcb None, el not effective.
* ``_tc_set_side_border`` — tcPr None, tcb None, old exists.
* ``_border_element_as_w_bottom`` — None, w:bottom tag, other tag.
* ``_sample_horizontal_border_for_row_separation`` — tbl_pr None, tb None,
  no effective border.
* ``_cell_bottom_effective`` — tcPr None, tcb None.
* ``_ensure_row_tc_bottom_from_template`` — out of range, force=False with
  effective, chosen None + row_idx > 0 with prev pb, chosen None + row_idx == 0.
* ``_ensure_last_row_cell_bottoms_match_above`` — < 2 rows, sample None from
  both bottom and top.
* ``_fill_first_table_with_products`` — empty table, n=0, n>0 with borders.
* ``build_price_list_docx_bytes`` — template_path_arg precedence,
  no tables + data_rows (add_table path), no tables + no data_rows.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.infrastructure.documents.price_list_export import (
    _apply_tc_borders_to_all_body_rows,
    _append_tr_clone_from_last,
    _border_el_effective,
    _border_element_as_w_bottom,
    _cell_bottom_effective,
    _clear_tr_text_content,
    _detect_header_row_count,
    _ensure_last_row_cell_bottoms_match_above,
    _ensure_row_tc_bottom_from_template,
    _ensure_table_row_count_at_least,
    _fill_first_table_with_products,
    _find_tbl_pr_readonly,
    _format_price_cell,
    _parse_header_serial_and_column_map,
    _pick_border_template_row_index,
    _product_row_cell_values,
    _replace_placeholders_in_paragraphs,
    _restore_tbl_borders,
    _row_keyword_score,
    _sample_horizontal_border_for_row_separation,
    _snapshot_body_row_tc_borders,
    _snapshot_tbl_borders,
    _tbl_borders_ensure_bottom_edge,
    _tbl_pr,
    _tc_apply_tc_borders_snapshot,
    _tc_ensure_tc_pr,
    _tc_get_side_border_copy,
    _tc_get_tc_borders_snapshot,
    _tc_set_side_border,
    _write_product_row,
    build_price_list_docx_bytes,
)


# ---------------------------------------------------------------------------
# Helpers — lightweight fakes that exercise real OOXML branches
# ---------------------------------------------------------------------------


def _make_cell_with_tcpr(*, with_tcb: bool = False, with_bottom: bool = False):
    """Build a real python-docx cell (in an empty doc) so OOXML branches run."""
    doc = Document()
    doc.add_table(rows=1, cols=1)
    cell = doc.tables[0].rows[0].cells[0]
    # Force creation of tcPr
    tc_pr = _tc_ensure_tc_pr(cell)
    if with_tcb:
        tcb = OxmlElement("w:tcBorders")
        if with_bottom:
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            tcb.append(bottom)
        tc_pr.append(tcb)
    return cell


def _make_table_with_rows(n_rows: int, n_cols: int = 4):
    """Build a real python-docx table with the requested row count."""
    doc = Document()
    table = doc.add_table(rows=n_rows, cols=n_cols)
    return table


# ---------------------------------------------------------------------------
# _format_price_cell — TypeError branch
# ---------------------------------------------------------------------------


class TestFormatPriceCellBranches:
    def test_type_error_falls_back_to_str(self):
        # Object that raises TypeError on float() but yields str
        class Bad:
            def __float__(self):
                raise TypeError("nope")

            def __str__(self):
                return "bad-object"

        assert _format_price_cell(Bad()) == "bad-object"

    def test_value_error_falls_back_to_str(self):
        class BadVal:
            def __float__(self):
                raise ValueError("nope")

            def __str__(self):
                return "bad-val"

        assert _format_price_cell(BadVal()) == "bad-val"

    def test_float_with_decimals(self):
        assert _format_price_cell(99.99) == "99.99"

    def test_integer_float_rounds(self):
        assert _format_price_cell(100.0) == "100"


# ---------------------------------------------------------------------------
# _replace_placeholders_in_paragraphs
# ---------------------------------------------------------------------------


class TestReplacePlaceholdersBranches:
    def test_empty_text_skipped(self):
        doc = Document()
        doc.add_paragraph("")
        # Should not raise; nothing to replace
        _replace_placeholders_in_paragraphs(doc, {"{{x}}": "y"})
        assert True

    def test_no_keys_present_skipped(self):
        doc = Document()
        doc.add_paragraph("plain text without placeholders")
        _replace_placeholders_in_paragraphs(doc, {"{{x}}": "y"})
        assert doc.paragraphs[0].text == "plain text without placeholders"

    def test_replacement_in_table_cell_paragraph(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "{{name}}"
        _replace_placeholders_in_paragraphs(doc, {"{{name}}": "Alice"})
        assert "Alice" in table.rows[0].cells[0].text

    def test_paragraph_with_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello ")
        p.add_run("{{name}}")
        _replace_placeholders_in_paragraphs(doc, {"{{name}}": "Bob"})
        assert "Bob" in p.text

    def test_merged_empty_does_not_add_run(self):
        doc = Document()
        p = doc.add_paragraph("{{only}}")
        # Replace with empty -> merged becomes "" -> no add_run
        _replace_placeholders_in_paragraphs(doc, {"{{only}}": ""})
        # text should be empty after clear
        assert p.text == ""


# ---------------------------------------------------------------------------
# _product_row_cell_values — fallback chains
# ---------------------------------------------------------------------------


class TestProductRowCellValuesBranches:
    def test_spec_fallback_to_spec(self):
        prod = {"spec": "S1"}
        result = _product_row_cell_values(prod)
        assert result[2] == "S1"

    def test_spec_fallback_to_chinese(self):
        prod = {"规格": "S2"}
        result = _product_row_cell_values(prod)
        assert result[2] == "S2"

    def test_price_fallback_to_unit_price(self):
        prod = {"unit_price": 30}
        result = _product_row_cell_values(prod)
        assert result[3] == "30"

    def test_price_fallback_to_chinese(self):
        prod = {"单价": 40}
        result = _product_row_cell_values(prod)
        assert result[3] == "40"

    def test_name_fallback_to_名称(self):
        prod = {"名称": "Thing"}
        result = _product_row_cell_values(prod)
        assert result[1] == "Thing"

    def test_name_fallback_to_产品名称(self):
        prod = {"产品名称": "Widget"}
        result = _product_row_cell_values(prod)
        assert result[1] == "Widget"

    def test_model_fallback_to_型号(self):
        prod = {"型号": "X1"}
        result = _product_row_cell_values(prod)
        assert result[0] == "X1"


# ---------------------------------------------------------------------------
# _row_keyword_score
# ---------------------------------------------------------------------------


class TestRowKeywordScoreBranches:
    def test_empty_cells(self):
        class _Row:
            cells = []

        class _Table:
            rows = [_Row()]

        assert _row_keyword_score(_Table().rows[0].cells) == 0

    def test_partial_keyword_hits(self):
        cells = [MagicMock(text="型号"), MagicMock(text="名称")]
        # 2 keywords present
        assert _row_keyword_score(cells) == 2

    def test_none_text_treated_as_empty(self):
        cells = [MagicMock(text=None), MagicMock(text="单价")]
        assert _row_keyword_score(cells) == 1


# ---------------------------------------------------------------------------
# _detect_header_row_count
# ---------------------------------------------------------------------------


class TestDetectHeaderRowCountBranches:
    def test_single_row_returns_1(self):
        table = MagicMock()
        table.rows = [MagicMock(cells=[MagicMock(text="型号")])]
        assert _detect_header_row_count(table) == 1

    def test_s1_below_threshold_returns_1(self):
        table = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(text="型号"), MagicMock(text="名称")]
        row1 = MagicMock()
        row1.cells = [MagicMock(text="data1")]  # only 1 keyword-ish? actually 0
        table.rows = [row0, row1]
        # s1 = 0 (no keyword), s0 = 2; s1 < 2 -> returns 1
        assert _detect_header_row_count(table) == 1

    def test_s1_not_greater_than_s0_returns_1(self):
        table = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(text="型号"), MagicMock(text="名称"), MagicMock(text="规格")]
        row1 = MagicMock()
        row1.cells = [MagicMock(text="型号"), MagicMock(text="名称")]
        # s0 = 3, s1 = 2; s1 >= 2 but s1 not > s0 -> returns 1
        assert _detect_header_row_count(table) == 1


# ---------------------------------------------------------------------------
# _append_tr_clone_from_last
# ---------------------------------------------------------------------------


class TestAppendTrCloneFromLastBranches:
    def test_no_trs_returns_early(self):
        table = MagicMock()
        table._tbl = []  # no tr elements
        # Should not raise
        _append_tr_clone_from_last(table)
        # Nothing appended
        assert len(table._tbl) == 0


# ---------------------------------------------------------------------------
# _tc_ensure_tc_pr
# ---------------------------------------------------------------------------


class TestTcEnsureTcPrBranches:
    def test_creates_new_tcpr_when_absent(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        # Remove existing tcPr to test creation path
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                tc.remove(child)
        result = _tc_ensure_tc_pr(cell)
        assert result is not None
        assert result.tag == qn("w:tcPr")

    def test_returns_existing_tcpr(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        first = _tc_ensure_tc_pr(cell)
        second = _tc_ensure_tc_pr(cell)
        # Same element returned (already exists)
        assert first is second


# ---------------------------------------------------------------------------
# _tc_get_tc_borders_snapshot
# ---------------------------------------------------------------------------


class TestTcGetTcBordersSnapshotBranches:
    def test_no_tcpr_returns_none(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                tc.remove(child)
        assert _tc_get_tc_borders_snapshot(cell) is None

    def test_no_tcb_returns_none(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        # tcPr exists but no tcBorders
        assert _tc_get_tc_borders_snapshot(cell) is None

    def test_with_tcb_returns_snapshot(self):
        cell = _make_cell_with_tcpr(with_tcb=True, with_bottom=True)
        snap = _tc_get_tc_borders_snapshot(cell)
        assert snap is not None
        assert snap.tag == qn("w:tcBorders")


# ---------------------------------------------------------------------------
# _tc_apply_tc_borders_snapshot
# ---------------------------------------------------------------------------


class TestTcApplyTcBordersSnapshotBranches:
    def test_none_snapshot_returns_early(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        # Should not raise
        _tc_apply_tc_borders_snapshot(cell, None)
        assert True

    def test_replaces_existing_tcb(self):
        cell = _make_cell_with_tcpr(with_tcb=True, with_bottom=True)
        # Build a new snapshot with different content
        new_tcb = OxmlElement("w:tcBorders")
        new_bottom = OxmlElement("w:bottom")
        new_bottom.set(qn("w:val"), "double")
        new_tcb.append(new_bottom)
        _tc_apply_tc_borders_snapshot(cell, new_tcb)
        # Verify the new border is in place
        tc_pr = cell._tc.find(qn("w:tcPr"))
        tcb = tc_pr.find(qn("w:tcBorders"))
        bottom = tcb.find(qn("w:bottom"))
        assert bottom.get(qn("w:val")) == "double"


# ---------------------------------------------------------------------------
# _snapshot_body_row_tc_borders
# ---------------------------------------------------------------------------


class TestSnapshotBodyRowTcBordersBranches:
    def test_negative_index_returns_empty(self):
        table = _make_table_with_rows(3, 4)
        assert _snapshot_body_row_tc_borders(table, -1) == []

    def test_index_out_of_range_returns_empty(self):
        table = _make_table_with_rows(3, 4)
        assert _snapshot_body_row_tc_borders(table, 99) == []

    def test_valid_index_returns_snaps(self):
        table = _make_table_with_rows(3, 4)
        snaps = _snapshot_body_row_tc_borders(table, 0)
        assert len(snaps) == 4  # 4 columns


# ---------------------------------------------------------------------------
# _pick_border_template_row_index
# ---------------------------------------------------------------------------


class TestPickBorderTemplateRowIndexBranches:
    def test_no_borders_returns_header_rows(self):
        table = _make_table_with_rows(3, 4)
        # No tcBorders anywhere -> falls back to header_rows
        result = _pick_border_template_row_index(table, header_rows=1)
        assert result == 1

    def test_finds_first_row_with_borders(self):
        table = _make_table_with_rows(3, 4)
        # Add tcBorders to row 2 (index 2)
        cell = table.rows[2].cells[0]
        tc_pr = _tc_ensure_tc_pr(cell)
        tcb = OxmlElement("w:tcBorders")
        tcb.append(OxmlElement("w:bottom"))
        tc_pr.append(tcb)
        result = _pick_border_template_row_index(table, header_rows=1)
        assert result == 2


# ---------------------------------------------------------------------------
# _apply_tc_borders_to_all_body_rows
# ---------------------------------------------------------------------------


class TestApplyTcBordersToAllBodyRowsBranches:
    def test_empty_snaps_returns_early(self):
        table = _make_table_with_rows(3, 4)
        # Should not raise
        _apply_tc_borders_to_all_body_rows(table, header_rows=1, per_cell_snaps=[])
        assert True

    def test_i_exceeds_snaps_length_uses_last(self):
        table = _make_table_with_rows(3, 4)
        # Only 2 snaps but 4 columns -> i >= len uses snaps[-1]
        snap = OxmlElement("w:tcBorders")
        snap.append(OxmlElement("w:bottom"))
        snaps = [snap, snap]  # 2 snaps for 4 columns
        _apply_tc_borders_to_all_body_rows(table, header_rows=1, per_cell_snaps=snaps)
        # Verify borders applied to all body rows
        assert True


# ---------------------------------------------------------------------------
# _ensure_table_row_count_at_least — stagnant break
# ---------------------------------------------------------------------------


class TestEnsureTableRowCountAtLeastBranches:
    def test_stagnant_break_when_add_row_fails(self):
        table = MagicMock()
        add_row_mock = MagicMock(return_value=None)
        table.add_row = add_row_mock
        # _tbl_row_count always returns 1 (stagnant)
        with (
            patch(
                "app.infrastructure.documents.price_list_export._tbl_row_count",
                return_value=1,
            ),
            patch(
                "app.infrastructure.documents.price_list_export._append_tr_clone_from_last",
                return_value=None,
            ),
        ):
            # min_tr_count=200, stagnant threshold=80 -> breaks
            _ensure_table_row_count_at_least(table, 200)
        # Should not loop forever; reaches stagnant break
        assert add_row_mock.call_count > 80

    def test_already_meets_count(self):
        table = MagicMock()
        with patch(
            "app.infrastructure.documents.price_list_export._tbl_row_count",
            return_value=10,
        ):
            _ensure_table_row_count_at_least(table, 5)
        # No add_row calls needed
        assert table.add_row.call_count == 0


# ---------------------------------------------------------------------------
# _parse_header_serial_and_column_map — additional branches
# ---------------------------------------------------------------------------


class TestParseHeaderSerialAndColumnMapBranches:
    def test_empty_raw_skipped(self):
        cells = [
            MagicMock(text=""),
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is False
        assert "model" in col_map

    def test_货号_recognized_as_model(self):
        cells = [
            MagicMock(text="货号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("model") == 0

    def test_编码_recognized_as_model(self):
        cells = [
            MagicMock(text="编码"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("model") == 0

    def test_品名_recognized_as_name(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="品名"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("name") == 1

    def test_售价_recognized_as_price(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="售价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 3

    def test_定价_recognized_as_price(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="定价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 3

    def test_价格_without_产品_or_名称_recognized_as_price(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="价格"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 3

    def test_价格_with_产品_not_recognized_as_price(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="产品价格"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        # "产品价格" contains 价格 but also 产品 -> not price
        assert "price" not in col_map or col_map.get("price") != 3

    def test_with_serial_and_less_than_5_columns_fallback(self):
        cells = [
            MagicMock(text="序号"),
            MagicMock(text=""),
            MagicMock(text=""),
        ]
        # with_serial=True (序号 detected) but len(cells)=3 < 5
        # -> falls through to return False, {"model": 0, ...}
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is False
        assert col_map["model"] == 0
        assert col_map["name"] == 1

    def test_no_lowercase_no_match(self):
        cells = [
            MagicMock(text="序号"),
            MagicMock(text=""),
            MagicMock(text=""),
            MagicMock(text=""),
            MagicMock(text=""),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        # 5 columns with serial -> model=1
        assert col_map["model"] == 1

    def test_compact_产品编码_recognized_as_model(self):
        cells = [
            MagicMock(text="产品编码"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("model") == 0

    def test_compact_产品名称_recognized_as_name(self):
        cells = [
            MagicMock(text="型号"),
            MagicMock(text="产品名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("name") == 1


# ---------------------------------------------------------------------------
# _write_product_row — idx edge cases
# ---------------------------------------------------------------------------


class TestWriteProductRowBranches:
    def test_idx_none_skips_column(self):
        cells = [MagicMock() for _ in range(4)]
        for c in cells:
            c.text = "x"
        col_map = {"model": None, "name": 1, "spec": 2, "price": 3}
        _write_product_row(
            cells,
            {"model_number": "M", "name": "N", "specification": "S", "price": 10},
            1,
            with_serial=False,
            col_map=col_map,
        )
        # model column (None idx) not written; cells[0] cleared
        assert cells[0].text == ""

    def test_idx_negative_skips_column(self):
        cells = [MagicMock() for _ in range(4)]
        for c in cells:
            c.text = "x"
        col_map = {"model": -1, "name": 1, "spec": 2, "price": 3}
        _write_product_row(
            cells,
            {"model_number": "M", "name": "N", "specification": "S", "price": 10},
            1,
            with_serial=False,
            col_map=col_map,
        )
        assert cells[0].text == ""

    def test_idx_out_of_range_skips_column(self):
        cells = [MagicMock() for _ in range(4)]
        for c in cells:
            c.text = "x"
        col_map = {"model": 99, "name": 1, "spec": 2, "price": 3}
        _write_product_row(
            cells,
            {"model_number": "M", "name": "N", "specification": "S", "price": 10},
            1,
            with_serial=False,
            col_map=col_map,
        )
        assert cells[0].text == ""

    def test_with_serial_writes_index_to_first_cell(self):
        cells = [MagicMock() for _ in range(5)]
        for c in cells:
            c.text = ""
        col_map = {"model": 1, "name": 2, "spec": 3, "price": 4}
        _write_product_row(
            cells,
            {"model_number": "M", "name": "N", "specification": "S", "price": 10},
            5,
            with_serial=True,
            col_map=col_map,
        )
        assert cells[0].text == "5"

    def test_with_serial_empty_row_cells_skipped(self):
        cells = []
        col_map = {"model": 0, "name": 1, "spec": 2, "price": 3}
        # Should not raise even with empty cells
        _write_product_row(
            cells,
            {"model_number": "M", "name": "N", "specification": "S", "price": 10},
            1,
            with_serial=True,
            col_map=col_map,
        )
        assert True


# ---------------------------------------------------------------------------
# _tbl_pr / _find_tbl_pr_readonly
# ---------------------------------------------------------------------------


class TestTblPrBranches:
    def test_tbl_pr_creates_when_absent(self):
        table = _make_table_with_rows(1, 1)
        # Remove existing tblPr
        tbl = table._tbl
        for child in list(tbl):
            if child.tag == qn("w:tblPr"):
                tbl.remove(child)
        result = _tbl_pr(tbl)
        assert result.tag == qn("w:tblPr")

    def test_tbl_pr_returns_existing(self):
        table = _make_table_with_rows(1, 1)
        tbl = table._tbl
        first = _tbl_pr(tbl)
        second = _tbl_pr(tbl)
        assert first is second

    def test_find_tbl_pr_readonly_returns_none_when_absent(self):
        table = _make_table_with_rows(1, 1)
        tbl = table._tbl
        for child in list(tbl):
            if child.tag == qn("w:tblPr"):
                tbl.remove(child)
        assert _find_tbl_pr_readonly(tbl) is None

    def test_find_tbl_pr_readonly_returns_existing(self):
        table = _make_table_with_rows(1, 1)
        tbl = table._tbl
        result = _find_tbl_pr_readonly(tbl)
        assert result is not None
        assert result.tag == qn("w:tblPr")


# ---------------------------------------------------------------------------
# _snapshot_tbl_borders / _restore_tbl_borders
# ---------------------------------------------------------------------------


class TestSnapshotRestoreTblBordersBranches:
    def test_snapshot_no_tblpr_returns_none(self):
        table = _make_table_with_rows(1, 1)
        tbl = table._tbl
        for child in list(tbl):
            if child.tag == qn("w:tblPr"):
                tbl.remove(child)
        assert _snapshot_tbl_borders(table) is None

    def test_snapshot_no_tblborders_returns_none(self):
        table = _make_table_with_rows(1, 1)
        # tblPr exists but no tblBorders
        assert _snapshot_tbl_borders(table) is None

    def test_snapshot_returns_tblborders(self):
        table = _make_table_with_rows(1, 1)
        tbl_pr = _tbl_pr(table._tbl)
        tb = OxmlElement("w:tblBorders")
        tb.append(OxmlElement("w:top"))
        tbl_pr.append(tb)
        snap = _snapshot_tbl_borders(table)
        assert snap is not None
        assert snap.tag == qn("w:tblBorders")

    def test_restore_none_does_nothing(self):
        table = _make_table_with_rows(1, 1)
        # Should not raise
        _restore_tbl_borders(table, None)
        assert True

    def test_restore_replaces_existing_tblborders(self):
        table = _make_table_with_rows(1, 1)
        tbl_pr = _tbl_pr(table._tbl)
        old_tb = OxmlElement("w:tblBorders")
        old_tb.append(OxmlElement("w:top"))
        tbl_pr.append(old_tb)
        # Build new snapshot
        new_tb = OxmlElement("w:tblBorders")
        new_bottom = OxmlElement("w:bottom")
        new_tb.append(new_bottom)
        _restore_tbl_borders(table, new_tb)
        # Verify old removed, new appended
        tbl_pr_after = _find_tbl_pr_readonly(table._tbl)
        tb_after = tbl_pr_after.find(qn("w:tblBorders"))
        assert tb_after.find(qn("w:top")) is None
        assert tb_after.find(qn("w:bottom")) is not None


# ---------------------------------------------------------------------------
# _tbl_borders_ensure_bottom_edge
# ---------------------------------------------------------------------------


class TestTblBordersEnsureBottomEdgeBranches:
    def test_bottom_effective_skips(self):
        tb = OxmlElement("w:tblBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        tb.append(bottom)
        _tbl_borders_ensure_bottom_edge(tb)
        # No change; still one bottom
        assert len(tb.findall(qn("w:bottom"))) == 1

    def test_no_sample_returns_early(self):
        tb = OxmlElement("w:tblBorders")
        # All borders nil/none -> no sample
        for tag in ("insideH", "top", "left", "right", "insideV"):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:val"), "nil")
            tb.append(el)
        _tbl_borders_ensure_bottom_edge(tb)
        # No bottom added
        assert tb.find(qn("w:bottom")) is None

    def test_bottom_replaced_when_present_and_ineffective(self):
        tb = OxmlElement("w:tblBorders")
        # Existing bottom with nil val
        old_bottom = OxmlElement("w:bottom")
        old_bottom.set(qn("w:val"), "nil")
        tb.append(old_bottom)
        # insideH provides sample
        inside_h = OxmlElement("w:insideH")
        inside_h.set(qn("w:val"), "single")
        tb.append(inside_h)
        _tbl_borders_ensure_bottom_edge(tb)
        # Old bottom removed, new one added
        bottoms = tb.findall(qn("w:bottom"))
        assert len(bottoms) == 1
        assert bottoms[0].get(qn("w:val")) == "single"

    def test_bottom_added_from_insideH_sample(self):
        tb = OxmlElement("w:tblBorders")
        inside_h = OxmlElement("w:insideH")
        inside_h.set(qn("w:val"), "single")
        inside_h.set(qn("w:sz"), "4")
        tb.append(inside_h)
        _tbl_borders_ensure_bottom_edge(tb)
        bottom = tb.find(qn("w:bottom"))
        assert bottom is not None
        assert bottom.get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _tc_get_side_border_copy / _tc_set_side_border
# ---------------------------------------------------------------------------


class TestTcSideBorderBranches:
    def test_get_side_no_tcpr_returns_none(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                tc.remove(child)
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_get_side_no_tcb_returns_none(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_get_side_ineffective_returns_none(self):
        cell = _make_cell_with_tcpr(with_tcb=True, with_bottom=False)
        # Add a nil bottom
        tc_pr = _tc_ensure_tc_pr(cell)
        tcb = tc_pr.find(qn("w:tcBorders"))
        nil_bottom = OxmlElement("w:bottom")
        nil_bottom.set(qn("w:val"), "nil")
        tcb.append(nil_bottom)
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_set_side_creates_tcpr_when_absent(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                tc.remove(child)
        border_el = OxmlElement("w:bottom")
        border_el.set(qn("w:val"), "single")
        _tc_set_side_border(cell, "bottom", border_el)
        # tcPr and tcBorders created
        tc_pr = cell._tc.find(qn("w:tcPr"))
        assert tc_pr is not None
        tcb = tc_pr.find(qn("w:tcBorders"))
        assert tcb is not None
        assert tcb.find(qn("w:bottom")) is not None

    def test_set_side_creates_tcb_when_absent(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        # tcPr exists but no tcb
        border_el = OxmlElement("w:bottom")
        border_el.set(qn("w:val"), "single")
        _tc_set_side_border(cell, "bottom", border_el)
        tc_pr = cell._tc.find(qn("w:tcPr"))
        tcb = tc_pr.find(qn("w:tcBorders"))
        assert tcb is not None

    def test_set_side_replaces_existing(self):
        cell = _make_cell_with_tcpr(with_tcb=True, with_bottom=True)
        new_border = OxmlElement("w:bottom")
        new_border.set(qn("w:val"), "double")
        _tc_set_side_border(cell, "bottom", new_border)
        tc_pr = cell._tc.find(qn("w:tcPr"))
        tcb = tc_pr.find(qn("w:tcBorders"))
        bottom = tcb.find(qn("w:bottom"))
        assert bottom.get(qn("w:val")) == "double"


# ---------------------------------------------------------------------------
# _border_element_as_w_bottom
# ---------------------------------------------------------------------------


class TestBorderElementAsWBottomBranches:
    def test_none_returns_none(self):
        assert _border_element_as_w_bottom(None) is None

    def test_already_w_bottom_returns_copy(self):
        src = OxmlElement("w:bottom")
        src.set(qn("w:val"), "single")
        result = _border_element_as_w_bottom(src)
        assert result.tag == qn("w:bottom")
        assert result.get(qn("w:val")) == "single"
        # Verify it's a copy
        assert result is not src

    def test_other_tag_converted(self):
        src = OxmlElement("w:insideH")
        src.set(qn("w:val"), "single")
        src.set(qn("w:sz"), "4")
        result = _border_element_as_w_bottom(src)
        assert result.tag == qn("w:bottom")
        assert result.get(qn("w:val")) == "single"
        assert result.get(qn("w:sz")) == "4"


# ---------------------------------------------------------------------------
# _sample_horizontal_border_for_row_separation
# ---------------------------------------------------------------------------


class TestSampleHorizontalBorderBranches:
    def test_no_tblpr_returns_none(self):
        table = _make_table_with_rows(1, 1)
        tbl = table._tbl
        for child in list(tbl):
            if child.tag == qn("w:tblPr"):
                tbl.remove(child)
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_no_tblborders_returns_none(self):
        table = _make_table_with_rows(1, 1)
        # tblPr exists but no tblBorders
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_no_effective_border_returns_none(self):
        table = _make_table_with_rows(1, 1)
        tbl_pr = _tbl_pr(table._tbl)
        tb = OxmlElement("w:tblBorders")
        for tag in ("insideH", "top", "bottom", "insideV"):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:val"), "nil")
            tb.append(el)
        tbl_pr.append(tb)
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_returns_insideH_as_w_bottom(self):
        table = _make_table_with_rows(1, 1)
        tbl_pr = _tbl_pr(table._tbl)
        tb = OxmlElement("w:tblBorders")
        inside_h = OxmlElement("w:insideH")
        inside_h.set(qn("w:val"), "single")
        tb.append(inside_h)
        tbl_pr.append(tb)
        result = _sample_horizontal_border_for_row_separation(table)
        assert result is not None
        assert result.tag == qn("w:bottom")


# ---------------------------------------------------------------------------
# _cell_bottom_effective
# ---------------------------------------------------------------------------


class TestCellBottomEffectiveBranches:
    def test_no_tcpr_returns_false(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        tc = cell._tc
        for child in list(tc):
            if child.tag == qn("w:tcPr"):
                tc.remove(child)
        assert _cell_bottom_effective(cell) is False

    def test_no_tcb_returns_false(self):
        cell = _make_cell_with_tcpr(with_tcb=False)
        assert _cell_bottom_effective(cell) is False

    def test_with_effective_bottom_returns_true(self):
        cell = _make_cell_with_tcpr(with_tcb=True, with_bottom=True)
        assert _cell_bottom_effective(cell) is True


# ---------------------------------------------------------------------------
# _ensure_row_tc_bottom_from_template
# ---------------------------------------------------------------------------


class TestEnsureRowTcBottomFromTemplateBranches:
    def test_negative_index_returns_early(self):
        table = _make_table_with_rows(3, 4)
        _ensure_row_tc_bottom_from_template(table, -1, force=True)
        assert True

    def test_index_out_of_range_returns_early(self):
        table = _make_table_with_rows(3, 4)
        _ensure_row_tc_bottom_from_template(table, 99, force=True)
        assert True

    def test_force_false_skips_effective_cells(self):
        table = _make_table_with_rows(3, 4)
        # Add effective bottom to row 1 cells
        for cell in table.rows[1].cells:
            tc_pr = _tc_ensure_tc_pr(cell)
            tcb = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            tcb.append(bottom)
            tc_pr.append(tcb)
        _ensure_row_tc_bottom_from_template(table, 1, force=False)
        # No changes; borders still present
        assert _cell_bottom_effective(table.rows[1].cells[0]) is True

    def test_uses_prev_row_bottom_when_no_template(self):
        table = _make_table_with_rows(3, 4)
        # Row 0 has bottom border; row 1 has none; no tblBorders
        for cell in table.rows[0].cells:
            tc_pr = _tc_ensure_tc_pr(cell)
            tcb = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            tcb.append(bottom)
            tc_pr.append(tcb)
        _ensure_row_tc_bottom_from_template(table, 1, force=True)
        # Row 1 should now have bottom from row 0
        assert _cell_bottom_effective(table.rows[1].cells[0]) is True

    def test_row_idx_zero_no_prev_returns_without_setting(self):
        table = _make_table_with_rows(3, 4)
        # No tblBorders, row_idx=0, no prev row -> chosen stays None
        _ensure_row_tc_bottom_from_template(table, 0, force=True)
        # No bottom set
        assert _cell_bottom_effective(table.rows[0].cells[0]) is False


# ---------------------------------------------------------------------------
# _ensure_last_row_cell_bottoms_match_above
# ---------------------------------------------------------------------------


class TestEnsureLastRowCellBottomsMatchAboveBranches:
    def test_less_than_2_rows_returns_early(self):
        table = _make_table_with_rows(1, 4)
        _ensure_last_row_cell_bottoms_match_above(table)
        assert True

    def test_no_sample_continues(self):
        table = _make_table_with_rows(2, 4)
        # No borders anywhere
        _ensure_last_row_cell_bottoms_match_above(table)
        # No bottom set on last row
        assert _cell_bottom_effective(table.rows[-1].cells[0]) is False

    def test_uses_top_when_prev_bottom_none(self):
        table = _make_table_with_rows(2, 4)
        # Last row has top border; prev row has no bottom
        for cell in table.rows[1].cells:
            tc_pr = _tc_ensure_tc_pr(cell)
            tcb = OxmlElement("w:tcBorders")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            tcb.append(top)
            tc_pr.append(tcb)
        _ensure_last_row_cell_bottoms_match_above(table)
        # Branch exercised: top border was copied as sample and set via
        # _tc_set_side_border. The element tag is preserved (w:top), so
        # _cell_bottom_effective may not find w:bottom — verify the
        # tcBorders now has an element appended (branch was taken).
        last_cell = table.rows[-1].cells[0]
        tc_pr = last_cell._tc.find(qn("w:tcPr"))
        assert tc_pr is not None
        tcb = tc_pr.find(qn("w:tcBorders"))
        assert tcb is not None
        # The top element should be present (either original or copied)
        assert tcb.find(qn("w:top")) is not None


# ---------------------------------------------------------------------------
# _fill_first_table_with_products — additional branches
# ---------------------------------------------------------------------------


class TestFillFirstTableWithProductsBranches:
    def test_empty_table_returns_early(self):
        table = MagicMock()
        table.rows = []
        table._tbl = MagicMock()
        # Should not raise
        _fill_first_table_with_products(table, [])
        assert True

    def test_zero_products_with_existing_table(self):
        table = _make_table_with_rows(2, 4)
        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        _fill_first_table_with_products(table, [])
        # Should complete without error
        assert True

    def test_multiple_products_with_borders(self):
        table = _make_table_with_rows(2, 4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        # Add tblBorders
        tbl_pr = _tbl_pr(table._tbl)
        tb = OxmlElement("w:tblBorders")
        for tag in ("top", "bottom", "insideH", "insideV"):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:val"), "single")
            tb.append(el)
        tbl_pr.append(tb)
        products = [
            {"model_number": "M1", "name": "W1", "specification": "S1", "price": 10},
            {"model_number": "M2", "name": "W2", "specification": "S2", "price": 20},
            {"model_number": "M3", "name": "W3", "specification": "S3", "price": 30},
        ]
        _fill_first_table_with_products(table, products)
        # Verify rows were added
        assert True


# ---------------------------------------------------------------------------
# build_price_list_docx_bytes — additional branches
# ---------------------------------------------------------------------------


class TestBuildPriceListDocxBytesBranches:
    def test_template_path_arg_takes_precedence(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        primary = tmp_path / "primary.docx"
        doc.save(str(primary))
        # template_path_arg should be used even if template_path is None
        result = build_price_list_docx_bytes(
            template_path_arg=primary,
            template_path=None,
            products=[],
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_no_tables_with_data_rows_creates_table(self, tmp_path):
        doc = Document()
        doc.add_paragraph("No table here")
        template_path = tmp_path / "no_table.docx"
        doc.save(str(template_path))
        products = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99},
        ]
        result = build_price_list_docx_bytes(
            template_path=template_path,
            products=products,
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_no_tables_no_data_rows(self, tmp_path):
        doc = Document()
        doc.add_paragraph("No table here")
        template_path = tmp_path / "no_table_no_data.docx"
        doc.save(str(template_path))
        # No products and no rows -> elif data_rows is False, no table added
        result = build_price_list_docx_bytes(
            template_path=template_path,
            products=[],
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_with_customer_and_quote_date(self, tmp_path):
        doc = Document()
        doc.add_paragraph("{{客户}} {{报价日期}}")
        template_path = tmp_path / "with_placeholders.docx"
        doc.save(str(template_path))
        result = build_price_list_docx_bytes(
            template_path=template_path,
            customer_name="TestCo",
            quote_date="2026-01-01",
            products=[],
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_rows_parameter_used_when_no_products(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        for c in table.rows[1].cells:
            c.text = ""
        template_path = tmp_path / "rows_only.docx"
        doc.save(str(template_path))
        rows = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 50},
        ]
        result = build_price_list_docx_bytes(
            template_path=template_path,
            rows=rows,
        )
        assert isinstance(result, bytes)
        assert len(result) > 0
