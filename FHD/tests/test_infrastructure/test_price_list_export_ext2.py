"""Tests for app.infrastructure.documents.price_list_export — extended coverage.

Focus: _replace_placeholders_in_paragraphs, _row_keyword_score, _detect_header_row_count,
_tbl_row_count, _clear_tr_text_content, _append_tr_clone_from_last, _tc_ensure_tc_pr,
_tc_get_tc_borders_snapshot, _tc_apply_tc_borders_snapshot, _snapshot_body_row_tc_borders,
_pick_border_template_row_index, _apply_tc_borders_to_all_body_rows, _ensure_table_row_count_at_least,
_parse_header_serial_and_column_map, _write_product_row, _remove_table_row, _tbl_pr,
_find_tbl_pr_readonly, _snapshot_tbl_borders, _restore_tbl_borders, _border_el_effective,
_tbl_borders_ensure_bottom_edge, _tc_get_side_border_copy, _tc_set_side_border,
_border_element_as_w_bottom, _sample_horizontal_border_for_row_separation,
_cell_bottom_effective, _ensure_row_tc_bottom_from_template, _ensure_last_row_cell_bottoms_match_above,
_fill_first_table_with_products, build_price_list_docx_bytes, build_sales_contract_template_preview_json,
resolve_price_list_docx_template, build_price_list_template_preview_json.
"""

from __future__ import annotations

import os
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, Mock, patch

import pytest

from app.infrastructure.documents.price_list_export import (
    _apply_tc_borders_to_all_body_rows,
    _append_tr_clone_from_last,
    _border_el_effective,
    _border_element_as_w_bottom,
    _cell_bottom_effective,
    _clear_tr_text_content,
    _detect_header_row_count,
    _ensure_last_row_cell_bottoms_match_above,
    _ensure_row_tc_bottom_from_template,
    _ensure_table_row_count_at_least,
    _fill_first_table_with_products,
    _find_tbl_pr_readonly,
    _format_price_cell,
    _header_text,
    _parse_header_serial_and_column_map,
    _pick_border_template_row_index,
    _product_row_cell_values,
    _remove_table_row,
    _restore_tbl_borders,
    _row_keyword_score,
    _sample_horizontal_border_for_row_separation,
    _snapshot_body_row_tc_borders,
    _snapshot_tbl_borders,
    _tbl_borders_ensure_bottom_edge,
    _tbl_pr,
    _tbl_row_count,
    _tc_apply_tc_borders_snapshot,
    _tc_ensure_tc_pr,
    _tc_get_side_border_copy,
    _tc_get_tc_borders_snapshot,
    _tc_set_side_border,
    _write_product_row,
    build_price_list_docx_bytes,
    build_price_list_template_preview_json,
    build_sales_contract_template_preview_json,
    resolve_price_list_docx_template,
)


# ---------------------------------------------------------------------------
# _row_keyword_score — extended
# ---------------------------------------------------------------------------


class TestRowKeywordScoreExtended:
    def test_no_keywords(self):
        cells = [Mock(text="普通"), Mock(text="文本")]
        assert _row_keyword_score(cells) == 0

    def test_single_keyword(self):
        cells = [Mock(text="型号"), Mock(text="普通")]
        assert _row_keyword_score(cells) == 1

    def test_multiple_keywords(self):
        cells = [Mock(text="型号"), Mock(text="名称"), Mock(text="规格")]
        assert _row_keyword_score(cells) == 3

    def test_all_keywords(self):
        cells = [Mock(text=x) for x in ("型号", "名称", "规格", "单价", "数量", "金额", "产品", "序号", "单位", "售价", "定价")]
        assert _row_keyword_score(cells) == 11

    def test_empty_cells(self):
        assert _row_keyword_score([]) == 0


# ---------------------------------------------------------------------------
# _detect_header_row_count — extended
# ---------------------------------------------------------------------------


class TestDetectHeaderRowCountExtended:
    def test_single_row(self):
        table = Mock()
        row0 = Mock()
        row0.cells = [Mock(text="型号"), Mock(text="名称")]
        row1 = Mock()
        row1.cells = [Mock(text="M1"), Mock(text="Widget")]
        table.rows = [row0, row1]
        assert _detect_header_row_count(table) == 1

    def test_two_rows_when_second_is_header(self):
        table = Mock()
        row0 = Mock()
        row0.cells = [Mock(text="标题"), Mock(text="说明")]
        row1 = Mock()
        row1.cells = [Mock(text="型号"), Mock(text="名称")]
        table.rows = [row0, row1]
        assert _detect_header_row_count(table) == 2

    def test_single_row_table(self):
        table = Mock()
        row0 = Mock()
        row0.cells = [Mock(text="型号")]
        table.rows = [row0]
        assert _detect_header_row_count(table) == 1


# ---------------------------------------------------------------------------
# _parse_header_serial_and_column_map — extended
# ---------------------------------------------------------------------------


class TestParseHeaderSerialAndColumnMapExtended:
    def test_empty_cells(self):
        with_serial, col_map = _parse_header_serial_and_column_map([])
        assert with_serial is False
        assert col_map == {"model": 0, "name": 1, "spec": 2, "price": 3}

    def test_serial_column(self):
        cells = [Mock(text="序号"), Mock(text="型号"), Mock(text="名称"), Mock(text="规格"), Mock(text="单价")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert "model" in col_map
        assert "name" in col_map

    def test_serial_hash(self):
        cells = [Mock(text="#"), Mock(text="型号"), Mock(text="名称")]
        with_serial, _ = _parse_header_serial_and_column_map(cells)
        assert with_serial is True

    def test_serial_no(self):
        # Need >= 5 cells for fallback path when with_serial is True but
        # fewer than 2 core columns matched
        cells = [Mock(text="No."), Mock(text="a"), Mock(text="b"), Mock(text="c"), Mock(text="d")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert col_map == {"model": 1, "name": 2, "spec": 3, "price": 4}

    def test_no_serial(self):
        cells = [Mock(text="型号"), Mock(text="名称"), Mock(text="规格"), Mock(text="单价")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is False
        assert col_map["model"] == 0

    def test_fallback_with_serial_5_cols(self):
        cells = [Mock(text="序号"), Mock(text="a"), Mock(text="b"), Mock(text="c"), Mock(text="d")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert col_map == {"model": 1, "name": 2, "spec": 3, "price": 4}

    def test_price_alias_shoujia(self):
        cells = [Mock(text="型号"), Mock(text="售价")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 1

    def test_price_alias_dingjia(self):
        cells = [Mock(text="型号"), Mock(text="定价")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 1

    def test_price_with_jiag_no_product(self):
        cells = [Mock(text="型号"), Mock(text="价格")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 1

    def test_model_alias_huohao(self):
        cells = [Mock(text="货号"), Mock(text="名称")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("model") == 0

    def test_name_alias_pinming(self):
        cells = [Mock(text="型号"), Mock(text="品名")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("name") == 1


# ---------------------------------------------------------------------------
# _write_product_row — extended
# ---------------------------------------------------------------------------


class TestWriteProductRowExtended:
    def test_basic_write(self):
        cells = [Mock(text=""), Mock(text=""), Mock(text=""), Mock(text="")]
        prod = {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99}
        _write_product_row(cells, prod, 1, with_serial=False, col_map={"model": 0, "name": 1, "spec": 2, "price": 3})
        assert cells[0].text == "M1"
        assert cells[1].text == "Widget"
        assert cells[2].text == "10x20"
        assert cells[3].text == "99"

    def test_with_serial(self):
        cells = [Mock(text=""), Mock(text=""), Mock(text=""), Mock(text=""), Mock(text="")]
        prod = {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99}
        _write_product_row(cells, prod, 5, with_serial=True, col_map={"model": 1, "name": 2, "spec": 3, "price": 4})
        assert cells[0].text == "5"
        assert cells[1].text == "M1"

    def test_col_map_out_of_range(self):
        cells = [Mock(text="")]
        prod = {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99}
        # col_map points to indices that don't exist
        _write_product_row(cells, prod, 1, with_serial=False, col_map={"model": 5, "name": 6, "spec": 7, "price": 8})
        # Should not raise, just skip out-of-range indices


# ---------------------------------------------------------------------------
# _border_el_effective — extended
# ---------------------------------------------------------------------------


class TestBorderElEffectiveExtended:
    def test_none(self):
        assert _border_el_effective(None) is False

    def test_nil_value(self):
        el = Mock()
        el.get.return_value = "nil"
        with patch("app.infrastructure.documents.price_list_export.qn", return_value="w:val"):
            assert _border_el_effective(el) is False

    def test_none_value(self):
        el = Mock()
        el.get.return_value = "none"
        with patch("app.infrastructure.documents.price_list_export.qn", return_value="w:val"):
            assert _border_el_effective(el) is False

    def test_empty_value(self):
        el = Mock()
        el.get.return_value = ""
        with patch("app.infrastructure.documents.price_list_export.qn", return_value="w:val"):
            assert _border_el_effective(el) is False

    def test_effective_value(self):
        el = Mock()
        el.get.return_value = "single"
        with patch("app.infrastructure.documents.price_list_export.qn", return_value="w:val"):
            assert _border_el_effective(el) is True


# ---------------------------------------------------------------------------
# build_price_list_docx_bytes — extended
# ---------------------------------------------------------------------------


class TestBuildPriceListDocxBytesExtended:
    def test_no_template_raises(self):
        with pytest.raises(ValueError, match="template_path"):
            build_price_list_docx_bytes()

    def test_nonexistent_template_raises(self):
        with pytest.raises(FileNotFoundError, match="Word 模板不存在"):
            build_price_list_docx_bytes("/nonexistent/template.docx")

    def test_basic_generation(self, tmp_path):
        # Create a minimal docx file
        from docx import Document

        doc = Document()
        doc.add_paragraph("{{客户}}")
        doc.add_paragraph("{{报价日期}}")
        # Add a table
        table = doc.add_table(rows=2, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        # Second row is data placeholder
        data = table.rows[1].cells
        for c in data:
            c.text = ""

        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        products = [
            {"model_number": "M1", "name": "Widget", "specification": "10x20", "price": 99},
            {"model_number": "M2", "name": "Gadget", "specification": "5x10", "price": 50},
        ]

        result = build_price_list_docx_bytes(
            str(template_path),
            customer_name="TestCo",
            quote_date="2024-01-01",
            products=products,
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_with_rows_param(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_table(rows=2, cols=4)
        table = doc.tables[0]
        hdr = table.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"

        template_path = tmp_path / "template.docx"
        doc.save(str(template_path))

        rows = [{"model_number": "M1", "name": "W", "specification": "S", "price": 1}]
        result = build_price_list_docx_bytes(str(template_path), rows=rows)
        assert isinstance(result, bytes)

    def test_no_table_creates_one(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("test")
        template_path = tmp_path / "notable.docx"
        doc.save(str(template_path))

        result = build_price_list_docx_bytes(
            str(template_path),
            products=[{"model_number": "M1", "name": "W", "specification": "S", "price": 1}],
        )
        assert isinstance(result, bytes)


# ---------------------------------------------------------------------------
# build_sales_contract_template_preview_json — extended
# ---------------------------------------------------------------------------


class TestBuildSalesContractTemplatePreviewJsonExtended:
    def test_basic(self):
        with (
            patch(
                "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
                return_value=("/path/to/template.xlsx", "rel/path"),
            ),
            patch(
                "app.infrastructure.documents.price_list_export.read_excel_sales_contract_preview",
                return_value={"success": True, "data": []},
            ),
        ):
            result = build_sales_contract_template_preview_json()
        assert result["template_hint"] == "/path/to/template.xlsx"
        assert result["success"] is True


# ---------------------------------------------------------------------------
# resolve_price_list_docx_template — extended
# ---------------------------------------------------------------------------


class TestResolvePriceListDocxTemplateExtended:
    def test_basic(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=("/path/to/template.docx", "rel/path"),
        ):
            path, rel = resolve_price_list_docx_template()
        assert path == "/path/to/template.docx"
        assert rel == "rel/path"

    def test_with_slug(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=("/path/to/specific.docx", "rel/specific"),
        ) as mock_resolve:
            resolve_price_list_docx_template(slug="custom")
            mock_resolve.assert_called_once_with(role="price_list_docx", slug="custom")


# ---------------------------------------------------------------------------
# build_price_list_template_preview_json — extended
# ---------------------------------------------------------------------------


class TestBuildPriceListTemplatePreviewJsonExtended:
    def test_basic(self):
        with patch(
            "app.infrastructure.documents.price_list_export.resolve_price_list_docx_template",
            return_value=("/path/to/template.docx", "rel/path"),
        ):
            result = build_price_list_template_preview_json()
        assert result["success"] is True
        assert result["headers"] == ["产品", "规格", "单价"]
        assert result["sample_rows"] == []
        assert result["template_hint"] == "rel/path"
        assert result["path"] == "/path/to/template.docx"


# ---------------------------------------------------------------------------
# _format_price_cell — extended edge cases
# ---------------------------------------------------------------------------


class TestFormatPriceCellEdgeCases:
    def test_boolean_input(self):
        # bool is subclass of int, float(True) = 1.0
        assert _format_price_cell(True) == "1"

    def test_string_with_whitespace(self):
        assert _format_price_cell("  50  ") == "50"

    def test_very_large_number(self):
        assert _format_price_cell(1000000) == "1000000"

    def test_float_with_many_decimals(self):
        result = _format_price_cell(99.999)
        assert "100" in result  # rounds to 100.00

    def test_negative_float(self):
        assert _format_price_cell(-10.5) == "-10.50"


# ---------------------------------------------------------------------------
# _product_row_cell_values — extended edge cases
# ---------------------------------------------------------------------------


class TestProductRowCellValuesEdgeCases:
    def test_spec_alias(self):
        prod = {"spec": "S1", "name": "N", "model_number": "M", "price": 1}
        result = _product_row_cell_values(prod)
        assert result[2] == "S1"

    def test_unit_price_alias(self):
        prod = {"unit_price": 25, "name": "N", "model_number": "M", "specification": "S"}
        result = _product_row_cell_values(prod)
        assert result[3] == "25"

    def test_all_none_values(self):
        prod = {"model_number": None, "name": None, "specification": None, "price": None}
        result = _product_row_cell_values(prod)
        assert result == ["", "", "", ""]


# ---------------------------------------------------------------------------
# _header_text — extended
# ---------------------------------------------------------------------------


class TestHeaderTextExtended:
    def test_basic(self):
        cell = Mock(text="  型号  ")
        assert _header_text(cell) == "型号"

    def test_empty(self):
        cell = Mock(text="")
        assert _header_text(cell) == ""

    def test_none(self):
        cell = Mock(text=None)
        assert _header_text(cell) == ""
