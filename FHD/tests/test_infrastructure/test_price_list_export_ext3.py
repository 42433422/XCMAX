"""Tests for app.infrastructure.documents.price_list_export — extended coverage (ext3).

Focus: uncovered branches in _format_price_cell, _replace_placeholders_in_paragraphs,
_product_row_cell_values, _row_keyword_score, _detect_header_row_count,
_tbl_row_count, _clear_tr_text_content, _append_tr_clone_from_last,
_tc_ensure_tc_pr, _tc_get_tc_borders_snapshot, _tc_apply_tc_borders_snapshot,
_snapshot_body_row_tc_borders, _pick_border_template_row_index,
_apply_tc_borders_to_all_body_rows, _ensure_table_row_count_at_least,
_header_text, _parse_header_serial_and_column_map, _write_product_row,
_remove_table_row, _tbl_pr, _find_tbl_pr_readonly, _snapshot_tbl_borders,
_restore_tbl_borders, _border_el_effective, _tbl_borders_ensure_bottom_edge,
_tc_get_side_border_copy, _tc_set_side_border, _border_element_as_w_bottom,
_sample_horizontal_border_for_row_separation, _cell_bottom_effective,
_ensure_row_tc_bottom_from_template, _ensure_last_row_cell_bottoms_match_above,
_fill_first_table_with_products, build_price_list_docx_bytes.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# _format_price_cell — additional branches
# ---------------------------------------------------------------------------


class TestFormatPriceCellAdditional:
    def test_none_returns_empty(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell(None) == ""

    def test_empty_string_returns_empty(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell("") == ""

    def test_integer_value(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell(100) == "100"

    def test_float_integer_value(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell(100.0) == "100"

    def test_float_decimal_value(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell(99.99) == "99.99"

    def test_string_numeric_value(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell("50") == "50"

    def test_string_decimal_value(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        assert _format_price_cell("50.50") == "50.50"

    def test_type_error_returns_str(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        result = _format_price_cell([1, 2, 3])
        assert "[1, 2, 3]" in result

    def test_value_error_returns_str(self):
        from app.infrastructure.documents.price_list_export import _format_price_cell

        result = _format_price_cell("not_a_number")
        assert result == "not_a_number"


# ---------------------------------------------------------------------------
# _product_row_cell_values — additional branches
# ---------------------------------------------------------------------------


class TestProductRowCellValuesAdditional:
    def test_all_fields_present(self):
        from app.infrastructure.documents.price_list_export import (
            _product_row_cell_values,
        )

        prod = {
            "model_number": "M1",
            "name": "Product 1",
            "specification": "Spec 1",
            "price": 100,
        }
        result = _product_row_cell_values(prod)
        assert result == ["M1", "Product 1", "Spec 1", "100"]

    def test_chinese_keys(self):
        from app.infrastructure.documents.price_list_export import (
            _product_row_cell_values,
        )

        prod = {
            "型号": "M1",
            "产品名称": "Product 1",
            "规格": "Spec 1",
            "单价": 100,
        }
        result = _product_row_cell_values(prod)
        assert result == ["M1", "Product 1", "Spec 1", "100"]

    def test_empty_dict(self):
        from app.infrastructure.documents.price_list_export import (
            _product_row_cell_values,
        )

        result = _product_row_cell_values({})
        assert result == ["", "", "", ""]

    def test_alternative_keys(self):
        from app.infrastructure.documents.price_list_export import (
            _product_row_cell_values,
        )

        prod = {
            "name": "Product",
            "spec": "Spec",
            "unit_price": 50.5,
        }
        result = _product_row_cell_values(prod)
        assert result[1] == "Product"
        assert result[2] == "Spec"
        assert result[3] == "50.50"


# ---------------------------------------------------------------------------
# _row_keyword_score — additional
# ---------------------------------------------------------------------------


class TestRowKeywordScoreAdditional:
    def test_no_keywords(self):
        from app.infrastructure.documents.price_list_export import _row_keyword_score

        cells = [MagicMock(text="hello"), MagicMock(text="world")]
        assert _row_keyword_score(cells) == 0

    def test_multiple_keywords(self):
        from app.infrastructure.documents.price_list_export import _row_keyword_score

        cells = [MagicMock(text="型号"), MagicMock(text="名称"), MagicMock(text="规格")]
        assert _row_keyword_score(cells) == 3

    def test_none_text(self):
        from app.infrastructure.documents.price_list_export import _row_keyword_score

        cells = [MagicMock(text=None)]
        assert _row_keyword_score(cells) == 0


# ---------------------------------------------------------------------------
# _detect_header_row_count — additional
# ---------------------------------------------------------------------------


class TestDetectHeaderRowCountAdditional:
    def test_single_row(self):
        from app.infrastructure.documents.price_list_export import (
            _detect_header_row_count,
        )

        table = MagicMock()
        table.rows = [MagicMock(cells=[MagicMock(text="型号")])]
        result = _detect_header_row_count(table)
        assert result == 1

    def test_two_rows_second_better(self):
        from app.infrastructure.documents.price_list_export import (
            _detect_header_row_count,
        )

        table = MagicMock()
        row0 = MagicMock(cells=[MagicMock(text="title")])
        row1 = MagicMock(cells=[MagicMock(text="型号"), MagicMock(text="名称")])
        table.rows = [row0, row1]
        result = _detect_header_row_count(table)
        assert result == 2

    def test_two_rows_first_better(self):
        from app.infrastructure.documents.price_list_export import (
            _detect_header_row_count,
        )

        table = MagicMock()
        row0 = MagicMock(cells=[MagicMock(text="型号"), MagicMock(text="名称")])
        row1 = MagicMock(cells=[MagicMock(text="data")])
        table.rows = [row0, row1]
        result = _detect_header_row_count(table)
        assert result == 1


# ---------------------------------------------------------------------------
# _clear_tr_text_content
# ---------------------------------------------------------------------------


class TestClearTrTextContent:
    def test_clears_text_elements(self):
        from app.infrastructure.documents.price_list_export import (
            _clear_tr_text_content,
        )

        tr = OxmlElement("w:tr")
        t = OxmlElement("w:t")
        t.text = "hello"
        tr.append(t)
        _clear_tr_text_content(tr)
        assert t.text == ""


# ---------------------------------------------------------------------------
# _append_tr_clone_from_last
# ---------------------------------------------------------------------------


class TestAppendTrCloneFromLast:
    def test_no_rows_returns(self):
        from app.infrastructure.documents.price_list_export import (
            _append_tr_clone_from_last,
        )

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        table._tbl = tbl
        _append_tr_clone_from_last(table)
        # No source row to clone -> nothing appended
        trs = [c for c in tbl if c.tag == qn("w:tr")]
        assert trs == []

    def test_clones_last_row_with_text_cleared(self):
        from app.infrastructure.documents.price_list_export import (
            _append_tr_clone_from_last,
        )

        tbl = OxmlElement("w:tbl")
        tr = OxmlElement("w:tr")
        t = OxmlElement("w:t")
        t.text = "data"
        tr.append(t)
        tbl.append(tr)

        table = MagicMock()
        table._tbl = tbl
        _append_tr_clone_from_last(table)

        trs = [c for c in tbl if c.tag == qn("w:tr")]
        assert len(trs) == 2
        # The original row keeps its text; the appended clone is structurally a
        # copy (has a <w:t>) but with the text content cleared.
        original_texts = [el.text for el in trs[0].iter() if el.tag == qn("w:t")]
        clone_texts = [el.text for el in trs[1].iter() if el.tag == qn("w:t")]
        assert original_texts == ["data"]
        assert clone_texts == [""]


# ---------------------------------------------------------------------------
# _tc_ensure_tc_pr
# ---------------------------------------------------------------------------


class TestTcEnsureTcPr:
    def test_creates_tc_pr_when_missing(self):
        from app.infrastructure.documents.price_list_export import _tc_ensure_tc_pr

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        # A pre-existing child so we can assert tcPr is inserted *before* it
        # (OOXML requires tcPr to be the first child of tc).
        existing_child = OxmlElement("w:p")
        tc.append(existing_child)
        cell._tc = tc

        result = _tc_ensure_tc_pr(cell)

        assert result.tag == qn("w:tcPr")
        # Exactly one tcPr now exists and it is the first child.
        assert len(tc.findall(qn("w:tcPr"))) == 1
        assert list(tc)[0] is result
        assert tc.find(qn("w:tcPr")) is result

    def test_returns_existing_tc_pr(self):
        from app.infrastructure.documents.price_list_export import _tc_ensure_tc_pr

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        existing_pr = OxmlElement("w:tcPr")
        tc.append(existing_pr)
        cell._tc = tc
        result = _tc_ensure_tc_pr(cell)
        assert result is existing_pr


# ---------------------------------------------------------------------------
# _tc_get_tc_borders_snapshot
# ---------------------------------------------------------------------------


class TestTcGetTcBordersSnapshot:
    def test_no_tc_pr_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _tc_get_tc_borders_snapshot,
        )

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        cell._tc = tc
        assert _tc_get_tc_borders_snapshot(cell) is None

    def test_no_tc_borders_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _tc_get_tc_borders_snapshot,
        )

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc.append(tc_pr)
        cell._tc = tc
        assert _tc_get_tc_borders_snapshot(cell) is None

    def test_returns_independent_deep_copy_preserving_children(self):
        from app.infrastructure.documents.price_list_export import (
            _tc_get_tc_borders_snapshot,
        )

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tcb = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        tcb.append(bottom)
        tc_pr.append(tcb)
        tc.append(tc_pr)
        cell._tc = tc

        result = _tc_get_tc_borders_snapshot(cell)

        # Snapshot is a detached deep copy (not the live element).
        assert result is not tcb
        assert result.tag == qn("w:tcBorders")
        # Children/attributes are copied faithfully.
        snap_bottom = result.find(qn("w:bottom"))
        assert snap_bottom is not None
        assert snap_bottom.get(qn("w:val")) == "single"
        # Mutating the original afterwards does not bleed into the snapshot.
        bottom.set(qn("w:val"), "double")
        assert result.find(qn("w:bottom")).get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _tc_apply_tc_borders_snapshot
# ---------------------------------------------------------------------------


class TestTcApplyTcBordersSnapshot:
    def test_none_snapshot_returns(self):
        from app.infrastructure.documents.price_list_export import (
            _tc_apply_tc_borders_snapshot,
        )

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        cell._tc = tc
        _tc_apply_tc_borders_snapshot(cell, None)
        # None snapshot -> returns before ensuring tcPr; nothing created
        assert tc.find(qn("w:tcPr")) is None

    def test_replaces_existing_borders(self):
        from app.infrastructure.documents.price_list_export import (
            _tc_apply_tc_borders_snapshot,
        )

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        old_tcb = OxmlElement("w:tcBorders")
        old_bottom = OxmlElement("w:bottom")
        old_tcb.append(old_bottom)
        tc_pr.append(old_tcb)
        tc.append(tc_pr)
        cell._tc = tc

        new_snapshot = OxmlElement("w:tcBorders")
        new_bottom = OxmlElement("w:bottom")
        new_bottom.set(qn("w:val"), "single")
        new_snapshot.append(new_bottom)

        _tc_apply_tc_borders_snapshot(cell, new_snapshot)

        # Old (empty) tcBorders is replaced, not appended: exactly one remains
        # and it carries the new bottom edge value.
        tcbs = [c for c in tc_pr if c.tag == qn("w:tcBorders")]
        assert len(tcbs) == 1
        assert tcbs[0] is not new_snapshot  # stored as a deep copy
        assert tcbs[0].find(qn("w:bottom")).get(qn("w:val")) == "single"
        # The original old_tcb element is no longer attached.
        assert old_tcb not in list(tc_pr)


# ---------------------------------------------------------------------------
# _snapshot_body_row_tc_borders
# ---------------------------------------------------------------------------


class TestSnapshotBodyRowTcBorders:
    def test_invalid_index_returns_empty(self):
        from app.infrastructure.documents.price_list_export import (
            _snapshot_body_row_tc_borders,
        )

        # Real 2-row table; out-of-range indices yield [] without touching cells.
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        assert _snapshot_body_row_tc_borders(table, -1) == []
        assert _snapshot_body_row_tc_borders(table, 2) == []

    def test_returns_per_cell_snapshots_for_real_row(self):
        from app.infrastructure.documents.price_list_export import (
            _snapshot_body_row_tc_borders,
            _tc_set_side_border,
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        # Cell 0 of the data row has a bottom border; cell 1 has none.
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        _tc_set_side_border(table.rows[1].cells[0], "bottom", b)

        result = _snapshot_body_row_tc_borders(table, 1)

        # One entry per cell: a real tcBorders snapshot for cell 0, None for cell 1.
        assert len(result) == 2
        assert result[0] is not None
        assert result[0].tag == qn("w:tcBorders")
        assert result[0].find(qn("w:bottom")).get(qn("w:val")) == "single"
        assert result[1] is None


# ---------------------------------------------------------------------------
# _pick_border_template_row_index
# ---------------------------------------------------------------------------


class TestPickBorderTemplateRowIndex:
    def test_falls_back_to_header_rows_when_no_body_borders(self):
        from app.infrastructure.documents.price_list_export import (
            _pick_border_template_row_index,
        )

        # Real 3-row table, none of the body rows carry tcBorders.
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # No bordered data row found -> falls back to the first body row index.
        assert _pick_border_template_row_index(table, header_rows=1) == 1

    def test_skips_to_first_bordered_body_row(self):
        from app.infrastructure.documents.price_list_export import (
            _pick_border_template_row_index,
            _tc_set_side_border,
        )

        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        # header_rows=1 -> body rows are indices 1,2,3.
        # Give only row index 2 a real border; rows 1 and 3 stay borderless.
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        _tc_set_side_border(table.rows[2].cells[0], "bottom", b)

        # The first body row that actually has tcBorders is index 2.
        assert _pick_border_template_row_index(table, header_rows=1) == 2


# ---------------------------------------------------------------------------
# _apply_tc_borders_to_all_body_rows
# ---------------------------------------------------------------------------


class TestApplyTcBordersToAllBodyRows:
    def test_empty_snaps_leaves_cells_borderless(self):
        from app.infrastructure.documents.price_list_export import (
            _apply_tc_borders_to_all_body_rows,
            _cell_bottom_effective,
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        _apply_tc_borders_to_all_body_rows(table, 1, [])

        # Empty snapshot list -> early return: no borders fabricated anywhere.
        assert _cell_bottom_effective(table.rows[1].cells[0]) is False
        assert _cell_bottom_effective(table.rows[1].cells[1]) is False

    def test_applies_snapshot_to_every_body_row(self):
        from app.infrastructure.documents.price_list_export import (
            _apply_tc_borders_to_all_body_rows,
            _cell_bottom_effective,
        )

        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        # Build a per-cell snapshot list (one tcBorders with a bottom edge per col).
        snaps = []
        for _ in range(2):
            tcb = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            tcb.append(bottom)
            snaps.append(tcb)

        # Body rows are indices 1 and 2 (header_rows=1).
        assert _cell_bottom_effective(table.rows[1].cells[0]) is False
        assert _cell_bottom_effective(table.rows[2].cells[1]) is False

        _apply_tc_borders_to_all_body_rows(table, 1, snaps)

        # Both body rows, all columns, now carry the bottom edge from the snapshot.
        for r in (1, 2):
            for c in (0, 1):
                assert _cell_bottom_effective(table.rows[r].cells[c]) is True
        # The header row (index 0) is untouched.
        assert _cell_bottom_effective(table.rows[0].cells[0]) is False


# ---------------------------------------------------------------------------
# _ensure_table_row_count_at_least
# ---------------------------------------------------------------------------


class TestEnsureTableRowCountAtLeast:
    def test_already_has_enough_rows_is_noop(self):
        from app.infrastructure.documents.price_list_export import (
            _ensure_table_row_count_at_least,
            _tbl_row_count,
        )

        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        _ensure_table_row_count_at_least(table, 3)
        # Already >= target -> no rows added.
        assert _tbl_row_count(table) == 5

    def test_grows_table_to_exact_minimum(self):
        from app.infrastructure.documents.price_list_export import (
            _ensure_table_row_count_at_least,
            _tbl_row_count,
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        assert _tbl_row_count(table) == 2

        _ensure_table_row_count_at_least(table, 6)

        # Real <w:tr> count reaches the requested minimum, and new rows keep the
        # column count of the template.
        assert _tbl_row_count(table) == 6
        assert len(table.rows[5].cells) == 3


# ---------------------------------------------------------------------------
# _header_text
# ---------------------------------------------------------------------------


class TestHeaderText:
    def test_returns_stripped_text(self):
        from app.infrastructure.documents.price_list_export import _header_text

        cell = MagicMock(text="  hello  ")
        assert _header_text(cell) == "hello"

    def test_none_text(self):
        from app.infrastructure.documents.price_list_export import _header_text

        cell = MagicMock(text=None)
        assert _header_text(cell) == ""


# ---------------------------------------------------------------------------
# _parse_header_serial_and_column_map — additional branches
# ---------------------------------------------------------------------------


class TestParseHeaderSerialAndColumnMapAdditional:
    def test_empty_cells(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        with_serial, col_map = _parse_header_serial_and_column_map([])
        assert with_serial is False
        assert col_map == {"model": 0, "name": 1, "spec": 2, "price": 3}

    def test_serial_column_with_hash(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [MagicMock(text="#"), MagicMock(text="型号"), MagicMock(text="名称")]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True

    def test_serial_column_with_no(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        # Need 5 cells and at least 2 core keywords for with_serial to be returned
        cells = [
            MagicMock(text="No."),
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, _ = _parse_header_serial_and_column_map(cells)
        assert with_serial is True

    def test_serial_column_with_bianhao(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [
            MagicMock(text="编号"),
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="单价"),
        ]
        with_serial, _ = _parse_header_serial_and_column_map(cells)
        assert with_serial is True

    def test_fallback_with_serial_5_cells(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [
            MagicMock(text="序号"),
            MagicMock(text=""),
            MagicMock(text=""),
            MagicMock(text=""),
            MagicMock(text=""),
        ]
        with_serial, col_map = _parse_header_serial_and_column_map(cells)
        assert with_serial is True
        assert col_map == {"model": 1, "name": 2, "spec": 3, "price": 4}

    def test_price_keyword_match(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="售价"),
        ]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 3

    def test_dingjia_keyword_match(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [
            MagicMock(text="型号"),
            MagicMock(text="名称"),
            MagicMock(text="规格"),
            MagicMock(text="定价"),
        ]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("price") == 3

    def test_pinming_keyword_match(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        # Need at least 2 core keywords for col_map to be returned
        cells = [MagicMock(text="品名"), MagicMock(text="型号")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("name") == 0
        assert col_map.get("model") == 1

    def test_huohao_keyword_match(self):
        from app.infrastructure.documents.price_list_export import (
            _parse_header_serial_and_column_map,
        )

        cells = [MagicMock(text="货号")]
        _, col_map = _parse_header_serial_and_column_map(cells)
        assert col_map.get("model") == 0


# ---------------------------------------------------------------------------
# _write_product_row — additional branches
# ---------------------------------------------------------------------------


class TestWriteProductRowAdditional:
    def test_with_serial(self):
        from app.infrastructure.documents.price_list_export import _write_product_row

        cells = [MagicMock(), MagicMock(), MagicMock(), MagicMock(), MagicMock()]
        prod = {"model_number": "M1", "name": "P1", "specification": "S1", "price": 100}
        _write_product_row(
            cells,
            prod,
            1,
            with_serial=True,
            col_map={"model": 1, "name": 2, "spec": 3, "price": 4},
        )
        assert cells[0].text == "1"
        assert cells[1].text == "M1"
        assert cells[2].text == "P1"
        assert cells[3].text == "S1"
        assert cells[4].text == "100"

    def test_without_serial(self):
        from app.infrastructure.documents.price_list_export import _write_product_row

        cells = [MagicMock(), MagicMock(), MagicMock(), MagicMock()]
        prod = {"model_number": "M1", "name": "P1", "specification": "S1", "price": 100}
        _write_product_row(
            cells,
            prod,
            1,
            with_serial=False,
            col_map={"model": 0, "name": 1, "spec": 2, "price": 3},
        )
        assert cells[0].text == "M1"

    def test_index_out_of_range_skipped(self):
        from app.infrastructure.documents.price_list_export import _write_product_row

        cells = [MagicMock()]
        prod = {"model_number": "M1"}
        _write_product_row(
            cells,
            prod,
            1,
            with_serial=False,
            col_map={"model": 5, "name": 6, "spec": 7, "price": 8},
        )
        # All mapped indices are out of range and serial is off: the only
        # mutation is the clear loop, so the lone cell stays empty (M1 was NOT
        # written). If the out-of-range guard were missing this would IndexError.
        assert cells[0].text == ""


# ---------------------------------------------------------------------------
# _remove_table_row
# ---------------------------------------------------------------------------


class TestRemoveTableRow:
    def test_removes_specified_row_keeping_others(self):
        from app.infrastructure.documents.price_list_export import (
            _remove_table_row,
            _tbl_row_count,
        )

        # Real 3-row table with identifiable content per row.
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        table.rows[0].cells[0].text = "ROW0"
        table.rows[1].cells[0].text = "ROW1"
        table.rows[2].cells[0].text = "ROW2"

        _remove_table_row(table, 1)

        # The middle row is gone; the other two survive in order.
        assert _tbl_row_count(table) == 2
        assert [r.cells[0].text for r in table.rows] == ["ROW0", "ROW2"]


# ---------------------------------------------------------------------------
# _tbl_pr
# ---------------------------------------------------------------------------


class TestTblPr:
    def test_returns_existing(self):
        from app.infrastructure.documents.price_list_export import _tbl_pr

        tbl = OxmlElement("w:tbl")
        existing_pr = OxmlElement("w:tblPr")
        tbl.append(existing_pr)
        result = _tbl_pr(tbl)
        assert result is existing_pr

    def test_creates_new_as_first_child(self):
        from app.infrastructure.documents.price_list_export import _tbl_pr

        tbl = OxmlElement("w:tbl")
        # Existing row child so we can confirm tblPr is inserted before it
        # (OOXML requires tblPr to precede the rows).
        tr = OxmlElement("w:tr")
        tbl.append(tr)

        result = _tbl_pr(tbl)

        assert result.tag == qn("w:tblPr")
        assert list(tbl)[0] is result
        assert len(tbl.findall(qn("w:tblPr"))) == 1
        # Idempotent: a second call returns the same element, not a duplicate.
        assert _tbl_pr(tbl) is result
        assert len(tbl.findall(qn("w:tblPr"))) == 1


# ---------------------------------------------------------------------------
# _find_tbl_pr_readonly
# ---------------------------------------------------------------------------


class TestFindTblPrReadonly:
    def test_returns_none_when_missing(self):
        from app.infrastructure.documents.price_list_export import _find_tbl_pr_readonly

        tbl = OxmlElement("w:tbl")
        assert _find_tbl_pr_readonly(tbl) is None

    def test_returns_existing(self):
        from app.infrastructure.documents.price_list_export import _find_tbl_pr_readonly

        tbl = OxmlElement("w:tbl")
        existing_pr = OxmlElement("w:tblPr")
        tbl.append(existing_pr)
        assert _find_tbl_pr_readonly(tbl) is existing_pr


# ---------------------------------------------------------------------------
# _snapshot_tbl_borders / _restore_tbl_borders
# ---------------------------------------------------------------------------


class TestSnapshotRestoreTblBorders:
    def test_snapshot_no_tbl_pr(self):
        from app.infrastructure.documents.price_list_export import _snapshot_tbl_borders

        table = MagicMock()
        table._tbl = OxmlElement("w:tbl")
        assert _snapshot_tbl_borders(table) is None

    def test_snapshot_no_borders(self):
        from app.infrastructure.documents.price_list_export import _snapshot_tbl_borders

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)
        table._tbl = tbl
        assert _snapshot_tbl_borders(table) is None

    def test_snapshot_returns_independent_deep_copy(self):
        from app.infrastructure.documents.price_list_export import _snapshot_tbl_borders

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "8")
        borders.append(top)
        tbl_pr.append(borders)
        tbl.append(tbl_pr)
        table._tbl = tbl

        result = _snapshot_tbl_borders(table)

        assert result is not borders
        assert result.tag == qn("w:tblBorders")
        snap_top = result.find(qn("w:top"))
        assert snap_top.get(qn("w:val")) == "single"
        assert snap_top.get(qn("w:sz")) == "8"
        # Detached copy: editing the original leaves the snapshot untouched.
        top.set(qn("w:sz"), "16")
        assert result.find(qn("w:top")).get(qn("w:sz")) == "8"

    def test_restore_writes_snapshot_back_into_tbl_pr(self):
        from app.infrastructure.documents.price_list_export import (
            _restore_tbl_borders,
            _snapshot_tbl_borders,
        )

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        borders.append(bottom)
        tbl_pr.append(borders)
        tbl.append(tbl_pr)
        table._tbl = tbl

        snap = _snapshot_tbl_borders(table)
        # Remove the live borders, then restore from snapshot.
        tbl_pr.remove(borders)
        assert tbl_pr.find(qn("w:tblBorders")) is None

        _restore_tbl_borders(table, snap)

        restored = tbl_pr.find(qn("w:tblBorders"))
        assert restored is not None
        # Exactly one tblBorders is present (no duplication) with the val intact.
        assert len(tbl_pr.findall(qn("w:tblBorders"))) == 1
        assert restored.find(qn("w:bottom")).get(qn("w:val")) == "single"

    def test_restore_none_leaves_table_untouched(self):
        from app.infrastructure.documents.price_list_export import _restore_tbl_borders

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        # A bare <w:tbl> with no tblPr at all.
        table._tbl = tbl

        _restore_tbl_borders(table, None)

        # None snapshot -> early return: no tblPr fabricated, tree unchanged.
        assert tbl.find(qn("w:tblPr")) is None
        assert list(tbl) == []


# ---------------------------------------------------------------------------
# _border_el_effective
# ---------------------------------------------------------------------------


class TestBorderElEffective:
    def test_none_returns_false(self):
        from app.infrastructure.documents.price_list_export import _border_el_effective

        assert _border_el_effective(None) is False

    def test_no_val_returns_false(self):
        from app.infrastructure.documents.price_list_export import _border_el_effective

        el = OxmlElement("w:bottom")
        assert _border_el_effective(el) is False

    def test_nil_val_returns_false(self):
        from app.infrastructure.documents.price_list_export import _border_el_effective

        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "nil")
        assert _border_el_effective(el) is False

    def test_none_val_returns_false(self):
        from app.infrastructure.documents.price_list_export import _border_el_effective

        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "none")
        assert _border_el_effective(el) is False

    def test_single_val_returns_true(self):
        from app.infrastructure.documents.price_list_export import _border_el_effective

        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        assert _border_el_effective(el) is True


# ---------------------------------------------------------------------------
# _tbl_borders_ensure_bottom_edge
# ---------------------------------------------------------------------------


class TestTblBordersEnsureBottomEdge:
    def test_already_has_bottom_returns(self):
        from app.infrastructure.documents.price_list_export import (
            _tbl_borders_ensure_bottom_edge,
        )

        tb = OxmlElement("w:tblBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        tb.append(bottom)
        _tbl_borders_ensure_bottom_edge(tb)
        # Already has an effective bottom edge -> unchanged (no duplicate added,
        # original value preserved)
        bottoms = tb.findall(qn("w:bottom"))
        assert len(bottoms) == 1
        assert bottoms[0].get(qn("w:val")) == "single"

    def test_no_sample_returns(self):
        from app.infrastructure.documents.price_list_export import (
            _tbl_borders_ensure_bottom_edge,
        )

        tb = OxmlElement("w:tblBorders")
        _tbl_borders_ensure_bottom_edge(tb)
        # No sample border to copy from -> no bottom edge fabricated
        assert tb.find(qn("w:bottom")) is None

    def test_copies_from_insideH(self):
        from app.infrastructure.documents.price_list_export import (
            _tbl_borders_ensure_bottom_edge,
        )

        tb = OxmlElement("w:tblBorders")
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "single")
        tb.append(insideH)
        _tbl_borders_ensure_bottom_edge(tb)
        bottom = tb.find(qn("w:bottom"))
        assert bottom is not None
        assert bottom.get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _tc_get_side_border_copy
# ---------------------------------------------------------------------------


class TestTcGetSideBorderCopy:
    def test_no_tc_pr_returns_none(self):
        from app.infrastructure.documents.price_list_export import _tc_get_side_border_copy

        cell = MagicMock()
        cell._tc = OxmlElement("w:tc")
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_no_tcb_returns_none(self):
        from app.infrastructure.documents.price_list_export import _tc_get_side_border_copy

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc.append(tc_pr)
        cell._tc = tc
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_ineffective_border_returns_none(self):
        from app.infrastructure.documents.price_list_export import _tc_get_side_border_copy

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tcb = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "nil")
        tcb.append(bottom)
        tc_pr.append(tcb)
        tc.append(tc_pr)
        cell._tc = tc
        assert _tc_get_side_border_copy(cell, "bottom") is None

    def test_returns_independent_copy_of_effective_side(self):
        from app.infrastructure.documents.price_list_export import _tc_get_side_border_copy

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tcb = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:color"), "FF0000")
        tcb.append(bottom)
        tc_pr.append(tcb)
        tc.append(tc_pr)
        cell._tc = tc

        result = _tc_get_side_border_copy(cell, "bottom")

        assert result is not bottom
        assert result.tag == qn("w:bottom")
        # Style attributes are carried over on the copy.
        assert result.get(qn("w:val")) == "single"
        assert result.get(qn("w:color")) == "FF0000"
        # Detached: editing the live border does not change the returned copy.
        bottom.set(qn("w:color"), "00FF00")
        assert result.get(qn("w:color")) == "FF0000"


# ---------------------------------------------------------------------------
# _tc_set_side_border
# ---------------------------------------------------------------------------


class TestTcSetSideBorder:
    def test_creates_tc_pr_when_missing(self):
        from app.infrastructure.documents.price_list_export import _tc_set_side_border

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        cell._tc = tc
        border = OxmlElement("w:bottom")
        border.set(qn("w:val"), "single")
        _tc_set_side_border(cell, "bottom", border)
        tc_pr = tc.find(qn("w:tcPr"))
        assert tc_pr is not None
        tcb = tc_pr.find(qn("w:tcBorders"))
        assert tcb is not None
        # The requested side border is written with the supplied value, stored
        # as a deep copy (not the same element instance).
        stored = tcb.find(qn("w:bottom"))
        assert stored is not None
        assert stored is not border
        assert stored.get(qn("w:val")) == "single"

    def test_replaces_existing(self):
        from app.infrastructure.documents.price_list_export import _tc_set_side_border

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tcb = OxmlElement("w:tcBorders")
        old = OxmlElement("w:bottom")
        old.set(qn("w:val"), "nil")
        tcb.append(old)
        tc_pr.append(tcb)
        tc.append(tc_pr)
        cell._tc = tc

        new_border = OxmlElement("w:bottom")
        new_border.set(qn("w:val"), "single")
        _tc_set_side_border(cell, "bottom", new_border)
        bottoms = tcb.findall(qn("w:bottom"))
        # The pre-existing 'nil' bottom is replaced in place by the new 'single'.
        assert len(bottoms) == 1
        assert bottoms[0].get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _border_element_as_w_bottom
# ---------------------------------------------------------------------------


class TestBorderElementAsWBottom:
    def test_none_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _border_element_as_w_bottom,
        )

        assert _border_element_as_w_bottom(None) is None

    def test_already_bottom_returns_independent_copy(self):
        from app.infrastructure.documents.price_list_export import (
            _border_element_as_w_bottom,
        )

        src = OxmlElement("w:bottom")
        src.set(qn("w:val"), "single")
        src.set(qn("w:sz"), "12")
        result = _border_element_as_w_bottom(src)

        assert result is not src
        assert result.tag == qn("w:bottom")
        assert result.get(qn("w:val")) == "single"
        assert result.get(qn("w:sz")) == "12"
        # Editing the source must not mutate the returned copy.
        src.set(qn("w:sz"), "24")
        assert result.get(qn("w:sz")) == "12"

    def test_converts_other_tag(self):
        from app.infrastructure.documents.price_list_export import (
            _border_element_as_w_bottom,
        )

        src = OxmlElement("w:insideH")
        src.set(qn("w:val"), "single")
        result = _border_element_as_w_bottom(src)
        assert result is not None
        assert result.tag == qn("w:bottom")
        assert result.get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _sample_horizontal_border_for_row_separation
# ---------------------------------------------------------------------------


class TestSampleHorizontalBorderForRowSeparation:
    def test_no_tbl_pr_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _sample_horizontal_border_for_row_separation,
        )

        table = MagicMock()
        table._tbl = OxmlElement("w:tbl")
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_no_tbl_borders_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _sample_horizontal_border_for_row_separation,
        )

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)
        table._tbl = tbl
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_no_effective_borders_returns_none(self):
        from app.infrastructure.documents.price_list_export import (
            _sample_horizontal_border_for_row_separation,
        )

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        tb = OxmlElement("w:tblBorders")
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "nil")
        tb.append(insideH)
        tbl_pr.append(tb)
        tbl.append(tbl_pr)
        table._tbl = tbl
        assert _sample_horizontal_border_for_row_separation(table) is None

    def test_returns_insideH_as_bottom(self):
        from app.infrastructure.documents.price_list_export import (
            _sample_horizontal_border_for_row_separation,
        )

        table = MagicMock()
        tbl = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        tb = OxmlElement("w:tblBorders")
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "single")
        tb.append(insideH)
        tbl_pr.append(tb)
        tbl.append(tbl_pr)
        table._tbl = tbl
        result = _sample_horizontal_border_for_row_separation(table)
        assert result is not None
        # insideH is converted into a w:bottom carrying the same style value.
        assert result.tag == qn("w:bottom")
        assert result.get(qn("w:val")) == "single"


# ---------------------------------------------------------------------------
# _cell_bottom_effective
# ---------------------------------------------------------------------------


class TestCellBottomEffective:
    def test_no_tc_pr_returns_false(self):
        from app.infrastructure.documents.price_list_export import _cell_bottom_effective

        cell = MagicMock()
        cell._tc = OxmlElement("w:tc")
        assert _cell_bottom_effective(cell) is False

    def test_no_tcb_returns_false(self):
        from app.infrastructure.documents.price_list_export import _cell_bottom_effective

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc.append(tc_pr)
        cell._tc = tc
        assert _cell_bottom_effective(cell) is False

    def test_effective_bottom_returns_true(self):
        from app.infrastructure.documents.price_list_export import _cell_bottom_effective

        cell = MagicMock()
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tcb = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        tcb.append(bottom)
        tc_pr.append(tcb)
        tc.append(tc_pr)
        cell._tc = tc
        assert _cell_bottom_effective(cell) is True


# ---------------------------------------------------------------------------
# _ensure_row_tc_bottom_from_template
# ---------------------------------------------------------------------------


class TestEnsureRowTcBottomFromTemplate:
    def test_invalid_index_writes_no_borders(self):
        from app.infrastructure.documents.price_list_export import (
            _cell_bottom_effective,
            _ensure_row_tc_bottom_from_template,
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        # Give the table a usable insideH sample so that, if the guard were
        # missing, a border *could* be written somewhere observable.
        tbl_pr = table._tbl.find(qn("w:tblPr"))
        tb = OxmlElement("w:tblBorders")
        ih = OxmlElement("w:insideH")
        ih.set(qn("w:val"), "single")
        tb.append(ih)
        tbl_pr.append(tb)

        _ensure_row_tc_bottom_from_template(table, -1, force=True)
        _ensure_row_tc_bottom_from_template(table, 5, force=True)

        # Out-of-range indices -> early return: no cell gained a bottom border.
        for r in range(2):
            for c in range(2):
                assert _cell_bottom_effective(table.rows[r].cells[c]) is False

    def test_force_copies_table_insideH_to_row_cells(self):
        from app.infrastructure.documents.price_list_export import (
            _cell_bottom_effective,
            _ensure_row_tc_bottom_from_template,
            _tc_get_side_border_copy,
        )

        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        tbl_pr = table._tbl.find(qn("w:tblPr"))
        tb = OxmlElement("w:tblBorders")
        ih = OxmlElement("w:insideH")
        ih.set(qn("w:val"), "single")
        ih.set(qn("w:sz"), "6")
        tb.append(ih)
        tbl_pr.append(tb)

        assert _cell_bottom_effective(table.rows[1].cells[0]) is False
        _ensure_row_tc_bottom_from_template(table, 1, force=True)

        # Each cell of row 1 now has a bottom border derived from insideH,
        # preserving its style attributes.
        for c in range(2):
            assert _cell_bottom_effective(table.rows[1].cells[c]) is True
            copied = _tc_get_side_border_copy(table.rows[1].cells[c], "bottom")
            assert copied.get(qn("w:val")) == "single"
            assert copied.get(qn("w:sz")) == "6"


# ---------------------------------------------------------------------------
# _ensure_last_row_cell_bottoms_match_above
# ---------------------------------------------------------------------------


class TestEnsureLastRowCellBottomsMatchAbove:
    def test_single_row_table_is_left_untouched(self):
        from app.infrastructure.documents.price_list_export import (
            _cell_bottom_effective,
            _ensure_last_row_cell_bottoms_match_above,
        )

        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        _ensure_last_row_cell_bottoms_match_above(table)
        # Fewer than 2 rows -> early return, no bottom fabricated.
        assert _cell_bottom_effective(table.rows[0].cells[0]) is False
        assert _cell_bottom_effective(table.rows[0].cells[1]) is False

    def test_last_row_inherits_bottom_from_row_above(self):
        from app.infrastructure.documents.price_list_export import (
            _cell_bottom_effective,
            _ensure_last_row_cell_bottoms_match_above,
            _tc_get_side_border_copy,
            _tc_set_side_border,
        )

        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        # The middle row (index 1, the one above the last) carries bottom borders.
        for c in range(2):
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:color"), "123456")
            _tc_set_side_border(table.rows[1].cells[c], "bottom", b)

        assert _cell_bottom_effective(table.rows[2].cells[0]) is False
        _ensure_last_row_cell_bottoms_match_above(table)

        # Last row now mirrors the bottom style of the row above it.
        for c in range(2):
            assert _cell_bottom_effective(table.rows[2].cells[c]) is True
            copied = _tc_get_side_border_copy(table.rows[2].cells[c], "bottom")
            assert copied.get(qn("w:val")) == "single"
            assert copied.get(qn("w:color")) == "123456"


# ---------------------------------------------------------------------------
# _fill_first_table_with_products — additional branches
# ---------------------------------------------------------------------------


class TestFillFirstTableWithProductsAdditional:
    def test_empty_table_is_left_unchanged(self):
        from app.infrastructure.documents.price_list_export import (
            _fill_first_table_with_products,
            _tbl_row_count,
        )

        # A real <w:tbl> with zero <w:tr> -> table.rows is empty.
        doc = Document()
        table = doc.add_table(rows=0, cols=0)
        assert _tbl_row_count(table) == 0

        _fill_first_table_with_products(table, [])

        # No rows -> early return: still an empty table, nothing fabricated.
        assert _tbl_row_count(table) == 0

    def test_products_fill_rows_and_keep_trailing_blank(self):
        from app.infrastructure.documents.price_list_export import (
            _fill_first_table_with_products,
            _tbl_row_count,
        )

        doc = Document()
        tbl = doc.add_table(rows=2, cols=4)
        tbl.rows[0].cells[0].text = "型号"
        tbl.rows[0].cells[1].text = "名称"
        tbl.rows[0].cells[2].text = "规格"
        tbl.rows[0].cells[3].text = "单价"
        tbl.rows[1].cells[0].text = "STALE"

        products = [
            {"model_number": "M1", "name": "P1", "specification": "S1", "price": 100},
            {"model_number": "M2", "name": "P2", "specification": "S2", "price": 12.5},
        ]
        _fill_first_table_with_products(tbl, products)

        # header(1) + 2 products + 1 trailing blank row.
        assert _tbl_row_count(tbl) == 4
        assert [c.text for c in tbl.rows[0].cells] == ["型号", "名称", "规格", "单价"]
        assert [c.text for c in tbl.rows[1].cells] == ["M1", "P1", "S1", "100"]
        assert [c.text for c in tbl.rows[2].cells] == ["M2", "P2", "S2", "12.50"]
        assert [c.text for c in tbl.rows[3].cells] == ["", "", "", ""]

    def test_no_products(self):
        from app.infrastructure.documents.price_list_export import (
            _fill_first_table_with_products,
            _tbl_row_count,
        )

        doc = Document()
        tbl = doc.add_table(rows=2, cols=4)
        tbl.rows[0].cells[0].text = "型号"
        tbl.rows[0].cells[1].text = "名称"
        tbl.rows[0].cells[2].text = "规格"
        tbl.rows[0].cells[3].text = "单价"
        tbl.rows[1].cells[0].text = "OLD_DATA"
        _fill_first_table_with_products(tbl, [])
        # No products: body target is 0 rows, so the single data row is trimmed
        # away leaving only the header, which stays intact.
        assert _tbl_row_count(tbl) == 1
        assert tbl.rows[0].cells[0].text == "型号"
        assert tbl.rows[0].cells[3].text == "单价"


# ---------------------------------------------------------------------------
# build_price_list_docx_bytes — additional branches
# ---------------------------------------------------------------------------


class TestBuildPriceListDocxBytesAdditional:
    def test_none_path_raises_value_error(self):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_docx_bytes,
        )

        with pytest.raises(ValueError, match="template_path"):
            build_price_list_docx_bytes(None)

    def test_nonexistent_file_raises(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_docx_bytes,
        )

        with pytest.raises(FileNotFoundError, match="Word 模板不存在"):
            build_price_list_docx_bytes(str(tmp_path / "nonexistent.docx"))

    def test_template_path_kwarg_round_trips_to_valid_docx(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_docx_bytes,
        )

        # Create a minimal docx file with a known paragraph.
        doc = Document()
        doc.add_paragraph("hello-template")
        path = tmp_path / "template.docx"
        doc.save(str(path))

        result = build_price_list_docx_bytes(template_path=str(path))

        # Output is a real .docx (ZIP magic) that re-opens and preserves content.
        assert result[:2] == b"PK"
        reopened = Document(BytesIO(result))
        assert reopened.paragraphs[0].text == "hello-template"
        # No products + no tables -> no table is fabricated.
        assert reopened.tables == []

    def test_with_products_and_no_tables_creates_filled_table(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_docx_bytes,
        )

        # Template has a placeholder paragraph but no tables.
        doc = Document()
        doc.add_paragraph("客户：{{客户}} 日期：{{报价日期}}")
        path = tmp_path / "template.docx"
        doc.save(str(path))

        products = [
            {"model_number": "M1", "name": "P1", "specification": "S1", "price": 100},
            {"model_number": "M2", "name": "P2", "specification": "S2", "price": 99.5},
        ]
        result = build_price_list_docx_bytes(
            str(path),
            products=products,
            customer_name="ACME",
            quote_date="2026-01-01",
        )

        reopened = Document(BytesIO(result))
        # Placeholders are substituted in the paragraph.
        assert reopened.paragraphs[0].text == "客户：ACME 日期：2026-01-01"
        # A 4-column product table is created and filled with the standard header.
        assert len(reopened.tables) == 1
        tbl = reopened.tables[0]
        assert [c.text for c in tbl.rows[0].cells] == ["型号", "名称", "规格", "单价"]
        assert [c.text for c in tbl.rows[1].cells] == ["M1", "P1", "S1", "100"]
        assert [c.text for c in tbl.rows[2].cells] == ["M2", "P2", "S2", "99.50"]
        # n>0 -> one trailing blank data row is preserved for the bottom edge.
        assert [c.text for c in tbl.rows[3].cells] == ["", "", "", ""]

    def test_with_rows_kwarg_fills_existing_template_table(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_docx_bytes,
        )

        # Template already has a 4-col table with a recognizable header.
        doc = Document()
        tbl = doc.add_table(rows=2, cols=4)
        tbl.rows[0].cells[0].text = "型号"
        tbl.rows[0].cells[1].text = "名称"
        tbl.rows[0].cells[2].text = "规格"
        tbl.rows[0].cells[3].text = "单价"
        tbl.rows[1].cells[0].text = "STALE"
        path = tmp_path / "template.docx"
        doc.save(str(path))

        # `rows` is honored as the data source (alias for `products`).
        rows = [{"model_number": "M9", "name": "P9", "specification": "S9", "price": 7}]
        result = build_price_list_docx_bytes(str(path), rows=rows)

        reopened = Document(BytesIO(result))
        out_tbl = reopened.tables[0]
        # Header is preserved, stale data is overwritten with the supplied row.
        assert [c.text for c in out_tbl.rows[0].cells] == ["型号", "名称", "规格", "单价"]
        assert [c.text for c in out_tbl.rows[1].cells] == ["M9", "P9", "S9", "7"]


# ---------------------------------------------------------------------------
# _replace_placeholders_in_paragraphs
# ---------------------------------------------------------------------------


class TestReplacePlaceholdersInParagraphs:
    def test_replaces_in_paragraphs(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            _replace_placeholders_in_paragraphs,
        )

        doc = Document()
        doc.add_paragraph("Hello {{客户}}")
        _replace_placeholders_in_paragraphs(doc, {"{{客户}}": "World"})
        assert doc.paragraphs[0].text == "Hello World"

    def test_replaces_in_table_cells(self, tmp_path):
        from app.infrastructure.documents.price_list_export import (
            _replace_placeholders_in_paragraphs,
        )

        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        tbl.rows[0].cells[0].text = "{{单位}}"
        _replace_placeholders_in_paragraphs(doc, {"{{单位}}": "Test"})
        assert tbl.rows[0].cells[0].text == "Test"

    def test_no_replacement_when_no_match(self):
        from app.infrastructure.documents.price_list_export import (
            _replace_placeholders_in_paragraphs,
        )

        doc = Document()
        doc.add_paragraph("Hello World")
        _replace_placeholders_in_paragraphs(doc, {"{{客户}}": "X"})
        assert doc.paragraphs[0].text == "Hello World"


# ---------------------------------------------------------------------------
# build_sales_contract_template_preview_json / build_price_list_template_preview_json
# ---------------------------------------------------------------------------


class TestTemplatePreviewJson:
    def test_build_price_list_template_preview_json(self):
        from app.infrastructure.documents.price_list_export import (
            build_price_list_template_preview_json,
        )

        with patch(
            "app.infrastructure.documents.price_list_export.resolve_price_list_docx_template",
            return_value=(Path("/tmp/test.docx"), "test.docx"),
        ):
            result = build_price_list_template_preview_json()
        assert result["success"] is True
        assert result["template_hint"] == "test.docx"
        assert result["path"] == "/tmp/test.docx"

    def test_build_sales_contract_template_preview_json(self):
        from app.infrastructure.documents.price_list_export import (
            build_sales_contract_template_preview_json,
        )

        with (
            patch(
                "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
                return_value=(Path("/tmp/contract.docx"), "contract.docx"),
            ),
            patch(
                "app.infrastructure.documents.price_list_export.read_excel_sales_contract_preview",
                return_value={"success": True, "data": []},
            ),
        ):
            result = build_sales_contract_template_preview_json()
        assert result["success"] is True
        assert result["template_hint"] == "/tmp/contract.docx"

    def test_resolve_price_list_docx_template(self):
        from app.infrastructure.documents.price_list_export import (
            resolve_price_list_docx_template,
        )

        with patch(
            "app.infrastructure.documents.price_list_export.resolve_template_path_with_meta",
            return_value=(Path("/tmp/test.docx"), "rel"),
        ):
            path, rel = resolve_price_list_docx_template()
        assert path == Path("/tmp/test.docx")
        assert rel == "rel"
