"""价目表 Word 导出 + 销售合同模板预览。

Phase 3B 从 ``app.legacy.price_list_docx_export`` 吸收。
"""

from __future__ import annotations

import copy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.infrastructure.documents.sales_contract_excel import (
    read_excel_sales_contract_preview,
)
from app.infrastructure.documents.template_registry import (
    resolve_template_path_with_meta,
)


def build_sales_contract_template_preview_json(slug: str | None = None) -> dict:
    path, _ = resolve_template_path_with_meta(role="sales_contract_docx", slug=slug)
    out = read_excel_sales_contract_preview(path)
    out["template_hint"] = str(path)
    return out


def resolve_price_list_docx_template(slug: str | None = None):
    path, rel = resolve_template_path_with_meta(role="price_list_docx", slug=slug)
    return path, rel


def build_price_list_template_preview_json(slug: str | None = None) -> dict:
    path, rel = resolve_price_list_docx_template(slug)
    return {
        "success": True,
        "headers": ["产品", "规格", "单价"],
        "sample_rows": [],
        "template_hint": rel,
        "path": str(path),
    }


def _format_price_cell(val: Any) -> str:
    if val is None or val == "":
        return ""
    try:
        n = float(val)
        if abs(n - round(n)) < 1e-9:
            return str(int(round(n)))
        return f"{n:.2f}"
    except (TypeError, ValueError):
        return str(val)


def _replace_placeholders_in_paragraphs(doc, mapping: dict[str, str]) -> None:
    def repl(text: str) -> str:
        if not text:
            return text
        for k, v in mapping.items():
            text = text.replace(k, v)
        return text

    keys = tuple(mapping.keys())

    def fix_paragraph(p) -> None:
        txt = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
        if not txt or not any(k in txt for k in keys):
            return
        merged = repl(txt)
        p.clear()
        if merged:
            p.add_run(merged)

    for p in doc.paragraphs:
        fix_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_paragraph(p)


def _product_row_cell_values(prod: dict[str, Any]) -> list[str]:
    model = str(prod.get("model_number") or prod.get("型号") or "")
    name = str(prod.get("name") or prod.get("产品名称") or prod.get("名称") or "")
    spec = str(
        prod.get("specification")
        or prod.get("spec")
        or prod.get("规格")
        or ""
    )
    price = _format_price_cell(prod.get("price") or prod.get("单价") or prod.get("unit_price"))
    return [model, name, spec, price]


def _row_keyword_score(row_cells) -> int:
    blob = "".join((c.text or "") for c in row_cells)
    keywords = ("型号", "名称", "规格", "单价", "数量", "金额", "产品", "序号", "单位", "售价", "定价")
    return sum(1 for k in keywords if k in blob)


def _detect_header_row_count(table) -> int:
    """表头占几行：常见 1 行；若第 2 行更像列标题且第 1 行不像，则视为「标题行 + 表头」共 2 行。"""
    if len(table.rows) < 2:
        return 1
    s0 = _row_keyword_score(table.rows[0].cells)
    s1 = _row_keyword_score(table.rows[1].cells)
    if s1 >= 2 and s1 > s0:
        return 2
    return 1


def _tbl_row_count(table) -> int:
    """以 ``<w:tr>`` 实际个数为准（避免个别版式下 ``len(table.rows)`` 与真实行数不一致）。"""
    return len([c for c in table._tbl if c.tag == qn("w:tr")])


def _clear_tr_text_content(tr_el: Any) -> None:
    for el in tr_el.iter():
        if el.tag == qn("w:t"):
            el.text = ""


def _append_tr_clone_from_last(table) -> None:
    """``add_row()`` 未增加行数时的兜底：深拷贝最后一行 ``<w:tr>`` 并追加。"""
    tbl = table._tbl
    trs = [c for c in tbl if c.tag == qn("w:tr")]
    if not trs:
        return
    new_tr = copy.deepcopy(trs[-1])
    _clear_tr_text_content(new_tr)
    tbl.append(new_tr)


def _tc_ensure_tc_pr(cell) -> Any:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    return tc_pr


def _tc_get_tc_borders_snapshot(cell) -> Any | None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    tcb = tc_pr.find(qn("w:tcBorders"))
    return copy.deepcopy(tcb) if tcb is not None else None


def _tc_apply_tc_borders_snapshot(cell, tcb_snapshot: Any | None) -> None:
    """用快照替换单元格 ``w:tcBorders``（整块复制，保证新增行与模板数据行线框一致）。"""
    if tcb_snapshot is None:
        return
    tc_pr = _tc_ensure_tc_pr(cell)
    for child in list(tc_pr):
        if child.tag == qn("w:tcBorders"):
            tc_pr.remove(child)
    tc_pr.append(copy.deepcopy(tcb_snapshot))


def _snapshot_body_row_tc_borders(table, body_row_idx: int) -> list[Any | None]:
    """在清空表体前，抓取首行数据区各格的 ``tcBorders`` 作为后续所有数据行（含追加行）的线框模板。"""
    if body_row_idx < 0 or body_row_idx >= _tbl_row_count(table):
        return []
    return [_tc_get_tc_borders_snapshot(c) for c in table.rows[body_row_idx].cells]


def _pick_border_template_row_index(table, header_rows: int) -> int:
    """若首行数据无单元格边框，则向下找第一条带 ``tcBorders`` 的数据行作线框模板。"""
    lim = min(_tbl_row_count(table), header_rows + 25)
    for r in range(header_rows, lim):
        for c in table.rows[r].cells:
            if _tc_get_tc_borders_snapshot(c) is not None:
                return r
    return header_rows


def _apply_tc_borders_to_all_body_rows(table, header_rows: int, per_cell_snaps: list[Any | None]) -> None:
    """把首行数据区线框套到表体每一行（解决第 31 行及以后追加行无线框）。"""
    if not per_cell_snaps:
        return
    for r_idx in range(header_rows, _tbl_row_count(table)):
        cells = table.rows[r_idx].cells
        for i, cell in enumerate(cells):
            snap = per_cell_snaps[i] if i < len(per_cell_snaps) else per_cell_snaps[-1]
            _tc_apply_tc_borders_snapshot(cell, snap)


def _ensure_table_row_count_at_least(table, min_tr_count: int) -> None:
    """保证表格至少有 ``min_tr_count`` 个 ``<w:tr>``（用于产品数超过模板预画行数时继续向下扩展）。"""
    stagnant = 0
    while _tbl_row_count(table) < min_tr_count:
        before = _tbl_row_count(table)
        table.add_row()
        if _tbl_row_count(table) <= before:
            _append_tr_clone_from_last(table)
        if _tbl_row_count(table) <= before:
            stagnant += 1
            if stagnant > 80:
                break
        else:
            stagnant = 0


def _header_text(cell) -> str:
    return (cell.text or "").strip()


def _parse_header_serial_and_column_map(header_row_cells) -> tuple[bool, dict[str, int]]:
    """根据表头文字匹配列：仅当首列**明确**为序号列时才写行号；否则不在首列塞数字。

    返回 ``(是否有序号列, {model,name,spec,price}->列索引)``；无法从表头推断时退回 0..3 连续四列。
    """
    cells = list(header_row_cells)
    if not cells:
        return False, {"model": 0, "name": 1, "spec": 2, "price": 3}

    t0 = _header_text(cells[0])
    with_serial = bool(
        "序号" in t0
        or t0 in ("#", "No.", "NO", "编号", "序")
        or t0.lower() in ("no.", "no")
    )

    col_map: dict[str, int] = {}
    for idx, cell in enumerate(cells):
        raw = _header_text(cell)
        if not raw:
            continue
        compact = raw.replace(" ", "")
        if "规格" in raw:
            col_map.setdefault("spec", idx)
        elif "型号" in raw or "编码" in raw or "货号" in raw or "产品编码" in compact:
            col_map.setdefault("model", idx)
        elif "名称" in raw or "品名" in raw or "产品名称" in compact:
            col_map.setdefault("name", idx)
        elif any(x in raw for x in ("单价", "售价", "定价")) or (
            "价格" in raw and "产品" not in raw and "名称" not in raw
        ):
            col_map.setdefault("price", idx)

    core = ("model", "name", "spec", "price")
    if sum(1 for k in core if k in col_map) >= 2:
        return with_serial, col_map

    # 回退：按常见四列顺序（有序号列则整体右移一列）
    if with_serial and len(cells) >= 5:
        return True, {"model": 1, "name": 2, "spec": 3, "price": 4}
    return False, {"model": 0, "name": 1, "spec": 2, "price": 3}


def _write_product_row(
    row_cells,
    prod: dict[str, Any],
    display_index: int,
    *,
    with_serial: bool,
    col_map: dict[str, int],
) -> None:
    m, n, s, p = _product_row_cell_values(prod)
    vals = {"model": m, "name": n, "spec": s, "price": p}
    for c in row_cells:
        c.text = ""
    if with_serial and len(row_cells) > 0:
        row_cells[0].text = str(display_index)
    for key in ("model", "name", "spec", "price"):
        idx = col_map.get(key)
        if idx is None or idx < 0 or idx >= len(row_cells):
            continue
        row_cells[idx].text = vals[key]


def _remove_table_row(table, row_idx: int) -> None:
    row = table.rows[row_idx]
    table._tbl.remove(row._tr)


def _tbl_pr(tbl) -> Any:
    """取或创建 ``w:tblPr``（兼容不同 python-docx / oxml 版本）。"""
    for child in tbl:
        if child.tag == qn("w:tblPr"):
            return child
    tbl_pr = OxmlElement("w:tblPr")
    tbl.insert(0, tbl_pr)
    return tbl_pr


def _find_tbl_pr_readonly(tbl) -> Any | None:
    for child in tbl:
        if child.tag == qn("w:tblPr"):
            return child
    return None


def _snapshot_tbl_borders(table) -> Any | None:
    """深拷贝模板中整张表的 ``w:tblBorders``（线型/粗细/颜色等），供填数后原样写回。"""
    tbl_pr = _find_tbl_pr_readonly(table._tbl)
    if tbl_pr is None:
        return None
    for child in tbl_pr:
        if child.tag == qn("w:tblBorders"):
            return copy.deepcopy(child)
    return None


def _restore_tbl_borders(table, borders_el: Any | None) -> None:
    """用快照还原表级边框；无快照时不写入，避免覆盖模板仅用单元格边框的版式。"""
    if borders_el is None:
        return
    tbl_pr = _tbl_pr(table._tbl)
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)
    tbl_pr.append(copy.deepcopy(borders_el))


def _border_el_effective(el: Any | None) -> bool:
    if el is None:
        return False
    v = el.get(qn("w:val"))
    if not v:
        return False
    return str(v).lower() not in ("nil", "none")


def _tbl_borders_ensure_bottom_edge(tbl_borders_el: Any) -> None:
    """表级 ``tblBorders`` 若缺底边或底边为 nil，则从 ``insideH``/``top`` 复制样式补底边（保留线型/颜色）。"""
    bottom = tbl_borders_el.find(qn("w:bottom"))
    if _border_el_effective(bottom):
        return
    sample = None
    for tag in ("insideH", "top", "left", "right", "insideV"):
        cand = tbl_borders_el.find(qn(f"w:{tag}"))
        if _border_el_effective(cand):
            sample = cand
            break
    if sample is None:
        return
    nb = OxmlElement("w:bottom")
    for k, v in sample.attrib.items():
        nb.set(k, v)
    if bottom is not None:
        tbl_borders_el.remove(bottom)
    tbl_borders_el.append(nb)


def _tc_get_side_border_copy(cell, side: str) -> Any | None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    tcb = tc_pr.find(qn("w:tcBorders"))
    if tcb is None:
        return None
    el = tcb.find(qn(f"w:{side}"))
    if not _border_el_effective(el):
        return None
    return copy.deepcopy(el)


def _tc_set_side_border(cell, side: str, border_el: Any) -> None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    tcb = tc_pr.find(qn("w:tcBorders"))
    if tcb is None:
        tcb = OxmlElement("w:tcBorders")
        tc_pr.append(tcb)
    tag = qn(f"w:{side}")
    old = tcb.find(tag)
    if old is not None:
        tcb.remove(old)
    tcb.append(copy.deepcopy(border_el))


def _border_element_as_w_bottom(src: Any | None) -> Any | None:
    """把 ``insideH``/``top`` 等边线元素转成 ``w:bottom``（属性沿用模板线型/颜色/粗细）。"""
    if src is None:
        return None
    if src.tag == qn("w:bottom"):
        return copy.deepcopy(src)
    out = OxmlElement("w:bottom")
    for k, v in src.attrib.items():
        out.set(k, v)
    return out


def _sample_horizontal_border_for_row_separation(table) -> Any | None:
    """行间横线参考：优先 ``tblBorders/insideH``，其次 ``top``/``bottom``/``insideV``。"""
    tbl_pr = _find_tbl_pr_readonly(table._tbl)
    if tbl_pr is None:
        return None
    tb = tbl_pr.find(qn("w:tblBorders"))
    if tb is None:
        return None
    for tag in ("insideH", "top", "bottom", "insideV"):
        el = tb.find(qn(f"w:{tag}"))
        if _border_el_effective(el):
            return _border_element_as_w_bottom(el)
    return None


def _cell_bottom_effective(cell) -> bool:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return False
    tcb = tc_pr.find(qn("w:tcBorders"))
    if tcb is None:
        return False
    return _border_el_effective(tcb.find(qn("w:bottom")))


def _ensure_row_tc_bottom_from_template(table, row_idx: int, *, force: bool) -> None:
    """为指定行各单元格补全 ``w:bottom``（与模板横线一致），解决「最后一行产品下面少一条线」。"""
    if row_idx < 0 or row_idx >= _tbl_row_count(table):
        return
    row_cells = table.rows[row_idx].cells
    tbl_bottom_style = _sample_horizontal_border_for_row_separation(table)
    for i, cell in enumerate(row_cells):
        if not force and _cell_bottom_effective(cell):
            continue
        chosen = tbl_bottom_style
        if chosen is None and row_idx > 0:
            prev = table.rows[row_idx - 1].cells
            if i < len(prev):
                pb = _tc_get_side_border_copy(prev[i], "bottom")
                if pb is not None:
                    chosen = _border_element_as_w_bottom(pb)
        if chosen is not None:
            _tc_set_side_border(cell, "bottom", chosen)


def _ensure_last_row_cell_bottoms_match_above(table) -> None:
    """最后一行各单元格补 ``tcBorders/w:bottom``，避免表尾少一条横线（样式来自上一行同列底边或顶边）。"""
    rows = table.rows
    if _tbl_row_count(table) < 2:
        return
    last = rows[-1].cells
    prev = rows[-2].cells
    m = min(len(last), len(prev))
    for i in range(m):
        sample = _tc_get_side_border_copy(prev[i], "bottom")
        if sample is None:
            sample = _tc_get_side_border_copy(last[i], "top")
        if sample is None:
            continue
        _tc_set_side_border(last[i], "bottom", sample)


def _fill_first_table_with_products(table, products: list[dict[str, Any]]) -> None:
    """按导出条数调整数据行行数；表级线框尽量保持模板原样。

    1. 在改动行前快照 ``w:tblBorders``（若有），保留线型/颜色/粗细等 OOXML。
    2. 自动识别表头占 1 行或 2 行（标题行 + 列标题），数据从表头下一行写入。
    3. 仅当表头首列**明确**为「序号」等时才写行号；列顺序按表头文字匹配（型号/名称/规格/单价），避免首列多出一格无表头。
    4. 从首条带线数据行快照 ``tcBorders``，填数后套到全部表体行，避免第 31 行等追加行缺格线。
    5. 若有快照则写回 ``tblBorders``；与单元格线框叠加。
    6. 有产品时在表尾多保留 **一行空白数据行**，避免裁到只剩 n 行时吃掉模板底行/底边框显得「下面少一行」。
    """
    if not table.rows:
        return

    n = len(products)
    borders_snap = _snapshot_tbl_borders(table)
    header_rows = _detect_header_row_count(table)
    header_cells = table.rows[header_rows - 1].cells
    with_serial, col_map = _parse_header_serial_and_column_map(header_cells)

    # 表体目标行数：n 条产品 +（有数据时）表尾留白一行，便于底边线闭合、与常见价目表版式一致
    body_rows_target = n + (1 if n > 0 else 0)
    need_tr = header_rows + body_rows_target

    _ensure_table_row_count_at_least(table, need_tr)

    # 在清空表体文字之前，抓取一条「有线」的数据行各格 tcBorders，套到所有表体行（含第 31 行及以后追加行）
    cell_border_snaps: list[Any | None] = []
    if _tbl_row_count(table) > header_rows:
        tpl_r = _pick_border_template_row_index(table, header_rows)
        cell_border_snaps = _snapshot_body_row_tc_borders(table, tpl_r)

    for r_idx in range(header_rows, _tbl_row_count(table)):
        for cell in table.rows[r_idx].cells:
            cell.text = ""

    while _tbl_row_count(table) > need_tr and _tbl_row_count(table) > header_rows:
        _remove_table_row(table, _tbl_row_count(table) - 1)

    _ensure_table_row_count_at_least(table, need_tr)

    for i, prod in enumerate(products):
        _write_product_row(
            table.rows[header_rows + i].cells,
            prod,
            i + 1,
            with_serial=with_serial,
            col_map=col_map,
        )

    for r_idx in range(header_rows + n, _tbl_row_count(table)):
        for cell in table.rows[r_idx].cells:
            cell.text = ""

    _apply_tc_borders_to_all_body_rows(table, header_rows, cell_border_snaps)

    _restore_tbl_borders(table, borders_snap)
    tbl_pr = _find_tbl_pr_readonly(table._tbl)
    if tbl_pr is not None:
        tb = tbl_pr.find(qn("w:tblBorders"))
        if tb is not None:
            _tbl_borders_ensure_bottom_edge(tb)
    # 最后一行「产品」底下横线：仅靠表尾留白行不够时，用 insideH/上一行样式显式写 bottom
    if n >= 1:
        last_product_row = header_rows + n - 1
        _ensure_row_tc_bottom_from_template(table, last_product_row, force=True)
    _ensure_last_row_cell_bottoms_match_above(table)


def build_price_list_docx_bytes(
    template_path_arg: Path | str | None = None,
    *,
    customer_name: str | None = None,
    template_slug: str | None = None,
    template_path: Path | str | None = None,
    quote_date: str | None = None,
    products: list[dict[str, Any]] | None = None,
    rows: list[dict[str, Any]] | None = None,
) -> bytes:
    """基于用户 .docx 模板生成价目表二进制（保留版式，首表数据区重写为当前产品行）。"""
    path = template_path_arg or template_path
    if path is None:
        raise ValueError("build_price_list_docx_bytes: template_path 不能为空")
    src = Path(path)
    if not src.is_file():
        raise FileNotFoundError(f"Word 模板不存在: {src}")

    data_rows = list(rows or products or [])
    cust = (customer_name or "").strip()
    qd = (quote_date or "").strip()

    from docx import Document

    doc = Document(str(src))

    mapping = {
        "{{客户}}": cust,
        "{{购买单位}}": cust,
        "{{单位}}": cust,
        "{{报价日期}}": qd,
        "{{日期}}": qd,
    }
    _replace_placeholders_in_paragraphs(doc, mapping)

    if doc.tables:
        _fill_first_table_with_products(doc.tables[0], data_rows)
    elif data_rows:
        tbl = doc.add_table(rows=1, cols=4)
        hdr = tbl.rows[0].cells
        hdr[0].text = "型号"
        hdr[1].text = "名称"
        hdr[2].text = "规格"
        hdr[3].text = "单价"
        _fill_first_table_with_products(tbl, data_rows)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = [
    "build_sales_contract_template_preview_json",
    "resolve_price_list_docx_template",
    "build_price_list_template_preview_json",
    "build_price_list_docx_bytes",
]
