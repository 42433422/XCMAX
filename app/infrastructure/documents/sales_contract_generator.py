from docx import Document
import os
import uuid
from datetime import datetime
from typing import List, Dict
from decimal import Decimal, ROUND_HALF_UP


class SalesContractGenerator:
    def __init__(self, template_path: str, output_dir: str = None):
        self.template_path = template_path
        if output_dir:
            self.output_dir = output_dir
        else:
            app_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            project_dir = os.path.dirname(app_dir)
            self.output_dir = os.path.join(project_dir, "generated_contracts")
        os.makedirs(self.output_dir, exist_ok=True)

    def _num_to_chinese_upper(self, num):
        if isinstance(num, Decimal):
            num = float(num)
        if num == 0:
            return '零元整'

        digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
        units = ['', '拾', '佰', '仟', '万']

        int_num = int(num)
        num_str = str(int_num)
        length = len(num_str)

        result = []
        for i, d in enumerate(num_str):
            digit = int(d)
            unit_idx = length - i - 1
            if digit != 0:
                result.append(digits[digit])
                if unit_idx > 0:
                    result.append(units[unit_idx])
            else:
                if result and result[-1] != '零' and result[-1] not in units:
                    result.append('零')

        s = ''.join(result)
        while '零零' in s:
            s = s.replace('零零', '零')
        s = s.strip('零')

        if not s.endswith('元') and not s.endswith('角') and not s.endswith('分'):
            s = s + '元'

        if '角' not in s and '分' not in s:
            s = s + '零角零分'
        elif '角' not in s:
            s = s + '零角'
        elif '分' not in s:
            s = s + '零分'

        return s

    def generate(
        self,
        customer_name: str,
        customer_phone: str = "",
        contract_date: str = None,
        products: List[Dict] = None,
        return_buckets_expected: int = 0,
        return_buckets_actual: int = 0
    ) -> Dict:
        if contract_date is None:
            now = datetime.now()
            contract_date = f"{now.year}年{now.month:02d}月{now.day:02d}日"

        if products is None:
            products = [{
                "model_number": "306B",
                "name": "PU亮光硬化剂",
                "spec": "10KG×1",
                "unit": "桶",
                "quantity": "10 KG",
                "unit_price": "39.2",
                "amount": "392"
            }]

        doc = Document(self.template_path)

        self._fill_customer(doc, customer_name, customer_phone)
        self._fill_address_date(doc, customer_name, contract_date)
        self._fill_table(doc, products, return_buckets_expected, return_buckets_actual)

        contract_id = str(uuid.uuid4())[:8]
        filename = f"{customer_name}_{contract_id}_销售合同.docx"
        filename = filename.replace(" ", "_").replace("/", "_")
        file_path = os.path.join(self.output_dir, filename)

        doc.save(file_path)

        def _to_decimal(val):
            try:
                return Decimal(str(val).replace(" KG", "").replace("kg", "").strip())
            except:
                return Decimal("0")

        total_quantity = sum(_to_decimal(p.get("quantity", "0")) for p in products if p.get("quantity"))
        total_amount = sum(Decimal(str(p.get("amount", "0"))) for p in products if p.get("amount"))

        return {
            "success": True,
            "contract_id": contract_id,
            "filename": filename,
            "file_path": file_path,
            "customer_name": customer_name,
            "contract_date": contract_date,
            "products": products,
            "total_quantity": float(total_quantity),
            "total_amount": float(total_amount)
        }

    def _fill_customer(self, doc, customer_name, customer_phone):
        for para in doc.paragraphs:
            if "惠州市宝盈家具有限公司" in para.text:
                para.text = para.text.replace("惠州市宝盈家具有限公司", customer_name)
                break

    def _fill_address_date(self, doc, customer_name, contract_date):
        for para in doc.paragraphs:
            if "ADDRESS" in para.text and "DATE" in para.text:
                para.text = para.text.replace("ADDRESS", customer_name)
                para.text = para.text.replace("DATE", contract_date)
                break

    def _fill_table(self, doc, products, return_buckets_expected, return_buckets_actual):
        table = doc.tables[0]

        total_quantity = 0
        total_amount = 0

        for idx, product in enumerate(products):
            if idx + 1 < len(table.rows):
                row = table.rows[idx + 1]

                row.cells[0].text = product.get("model_number", "")
                row.cells[1].text = product.get("name", "")
                row.cells[2].text = product.get("spec", "")
                row.cells[3].text = product.get("unit", "")
                row.cells[4].text = product.get("unit", "")
                row.cells[5].text = product.get("quantity", "")
                row.cells[6].text = product.get("quantity", "")
                row.cells[7].text = f"{product.get('unit_price', '')}元／KG"
                row.cells[8].text = f"{product.get('amount', '')}元"
                row.cells[9].text = ""

                qty = Decimal(product.get("quantity", "0").replace(" KG", "").replace("kg", ""))
                amt = Decimal(product.get("amount", "0"))
                total_quantity += qty
                total_amount += amt

        for i in range(len(products) + 1, 12):
            if i < len(table.rows):
                row = table.rows[i]
                for cell in row.cells:
                    cell.text = ""

        row14 = table.rows[14]
        row14.cells[4].text = ""
        row14.cells[5].text = f"{int(total_quantity)} KG"
        row14.cells[6].text = f"{int(total_quantity)} KG"
        row14.cells[7].text = ""
        row14.cells[8].text = ""
        row14.cells[9].text = ""

        row15 = table.rows[15]
        chinese_amount = self._num_to_chinese_upper(total_amount)
        row15.cells[0].text = "合计TOTAL\n（大写）"
        row15.cells[1].text = f"      {chinese_amount}"
        for i in range(2, 9):
            row15.cells[i].text = ""
        row15.cells[9].text = f"¥{total_amount:.2f}"