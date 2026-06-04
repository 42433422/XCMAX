#!/usr/bin/env python3
"""生成 XCAGI 技术服务合同 Word 模板（精简版，含 {{field_key}} 占位符）。

注意：仓库默认已使用「完善版」`1_完善版_字段模板.docx`，请勿随意运行本脚本覆盖。

输出：
  FHD/data/contracts/xcagi_service_v1/template.docx
  FHD/data/contracts/xcagi_service_v1/fadada/法大大上传_template.docx

用法（在 FHD 目录，确认要覆盖完善版后再执行）：
  .venv/bin/python scripts/build_contract_template_docx.py --force
"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

CONTRACT_DIR = ROOT / "data" / "contracts" / "xcagi_service_v1"
FADADA_DIR = CONTRACT_DIR / "fadada"
TEMPLATE_PATH = CONTRACT_DIR / "template.docx"
CONFIG_PATH = CONTRACT_DIR / "fill_config.json"


def _p(doc, text: str, *, bold: bool = False, size: int = 12) -> None:
    from docx import Document  # noqa: F401 — runtime import below
    from docx.shared import Pt
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True


def build_template_docx() -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("XCAGI 技术服务合同")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("合同编号：{{contract_number}}    版本：{{contract_version}}")

    _p(doc, "")
    _p(doc, "甲方（委托方）", bold=True)
    _p(doc, "名称：{{party_a_name}}")
    _p(doc, "统一社会信用代码/身份证号：{{party_a_credit_code}}")
    _p(doc, "地址：{{party_a_address}}")
    _p(doc, "法定代表人/负责人：{{party_a_legal_rep}}")
    _p(doc, "联系人：{{party_a_contact}}    电话：{{party_a_phone}}")
    _p(doc, "电子邮箱：{{party_a_email}}    即时通讯：{{party_a_im}}")

    _p(doc, "")
    _p(doc, "乙方（受托方）", bold=True)
    _p(doc, "名称：{{party_b_name}}")
    _p(doc, "统一社会信用代码：{{party_b_credit_code}}")
    _p(doc, "地址：{{party_b_address}}")
    _p(doc, "法定代表人：{{party_b_legal_rep}}")
    _p(doc, "联系人：{{party_b_contact}}    电话：{{party_b_phone}}")

    _p(doc, "")
    _p(doc, "鉴于甲乙双方就 XCAGI 企业 AI 员工平台相关软件服务事宜，经协商一致，订立本合同。", bold=True)

    _p(doc, "第一条  服务内容与期限", bold=True)
    _p(doc, "1.1 服务期限：自 {{service_start_date}} 起至 {{service_end_date}} 止。")
    _p(doc, "1.2 主要功能/模块：{{main_function_list}}")
    _p(doc, "1.3 授权：区域 {{license_area}}；账号 {{account_quantity}}；终端 {{device_quantity}}；并发 {{concurrent_quantity}}。")
    _p(doc, "1.4 授权模块：{{authorized_modules}}")

    _p(doc, "第二条  费用与支付", bold=True)
    _p(doc, "2.1 合同总金额：人民币 {{total_amount_number}} 元（大写：{{total_amount_upper}}），{{tax_included}}。")
    _p(
        doc,
        "2.2 付款：首期 {{first_payment_percent}} / {{first_payment_amount_number}} 元（{{first_payment_amount_upper}}），"
        "于签订后 {{first_payment_days}} 个工作日内支付；"
        "第二期条件 {{second_payment_condition}}，比例 {{second_payment_percent}}，金额 {{second_payment_amount_number}} 元，"
        "{{second_payment_days}} 个工作日内；"
        "尾款条件 {{final_payment_condition}}，比例 {{final_payment_percent}}，金额 {{final_payment_amount_number}} 元，"
        "{{final_payment_days}} 个工作日内。",
    )
    _p(doc, "2.3 收款账户：户名 {{bank_account_name}}；开户行 {{bank_name}}；账号 {{bank_account_number}}。")
    _p(doc, "2.4 逾期违约金：{{overdue_penalty_rate}}；逾期超 {{overdue_suspend_days}} 日乙方可暂停服务。")

    _p(doc, "第三条  交付与验收", bold=True)
    _p(doc, "3.1 首付款后 {{delivery_workdays}} 个工作日内完成交付；验收期 {{acceptance_workdays}} 个工作日。")
    _p(
        doc,
        "3.2 软件交付：App（{{app_platform}} {{app_version}}，{{app_package_name}}，{{app_delivery_method}}）；"
        "桌面端（{{desktop_platform}} {{desktop_version}}，{{desktop_package_name}}，{{desktop_delivery_method}}）。",
    )
    _p(doc, "3.3 后台地址（如有）：{{backend_service_url}}")

    _p(doc, "第四条  维护与支持", bold=True)
    _p(doc, "4.1 响应：一般问题 {{normal_response_hours}} 小时；重大问题 {{urgent_response_hours}} 小时。")
    _p(doc, "4.2 支持方式：{{support_channels}}")
    _p(doc, "4.3 额外服务（如有）：{{custom_service_scope}}")

    _p(doc, "第五条  保密、不可抗力与争议", bold=True)
    _p(doc, "5.1 保密期限 {{confidentiality_years}} 年（商业秘密除外）。")
    _p(doc, "5.2 不可抗力持续 {{force_majeure_days}} 日可协商变更或解除。")
    _p(doc, "5.3 争议提交 {{jurisdiction_court}} 解决。")
    _p(doc, "5.4 提前终止须提前 {{termination_notice_days}} 日书面通知。")

    _p(doc, "第六条  签署", bold=True)
    _p(doc, "6.1 本合同一式 {{copy_total_count}} 份，甲方执 {{party_a_copy_count}} 份，乙方执 {{party_b_copy_count}} 份。")
    _p(doc, "6.2 签署地点：{{sign_place}}    签署日期：{{sign_date}}")

    _p(doc, "")
    _p(doc, "甲方（盖章）：____________________          乙方（盖章）：____________________")
    _p(doc, "法定代表人或授权代表：________________      法定代表人或授权代表：________________")

    _p(doc, "")
    _p(doc, "【附件：软件交付与验收确认单】", bold=True)
    _p(doc, "交付日期：{{delivery_date}}    渠道：{{delivery_channel}}    校验值：{{package_checksum}}")
    _p(doc, "验收结果：{{acceptance_result}}    问题说明：{{acceptance_issues}}")

    CONTRACT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE_PATH)
    return TEMPLATE_PATH


def build_sample_filled_docx() -> Path:
    """生成一份示例填充合同，供法大大上传前预览版式。"""
    from app.services.service_contract_fill import build_merged_fields, _render_docx_bytes

    if not TEMPLATE_PATH.is_file():
        build_template_docx()
    values = build_merged_fields(
        0,
        username="示例客户科技有限公司",
        extra={
            "party_a_name": "示例客户科技有限公司",
            "party_a_credit_code": "91110000MA0000000X",
            "party_a_legal_rep": "张三",
            "party_a_contact": "李四",
            "party_a_phone": "13800000000",
            "total_amount_number": "10000.00",
            "total_amount_upper": "壹万元整",
            "service_start_date": "2026-06-01",
            "service_end_date": "2027-05-31",
            "sign_place": "成都市高新区",
        },
    )
    out = FADADA_DIR / "示例_已填充_供预览.docx"
    out.write_bytes(_render_docx_bytes(TEMPLATE_PATH, values))
    return out


def write_fadada_readme() -> Path:
    FADADA_DIR.mkdir(parents=True, exist_ok=True)
    cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8")) if CONFIG_PATH.is_file() else {}
    party_b = (cfg.get("party_b_registry_info") or {}).get("统一社会信用代码", "")
    readme = FADADA_DIR / "法大大签署模板配置指南.md"
    readme.write_text(
        """# 法大大签署模板配置指南（XCAGI 技术服务合同）

## 1. 本地生成模板文件

在 `FHD` 目录执行：

```bash
.venv/bin/python scripts/build_contract_template_docx.py
```

生成文件：

| 文件 | 用途 |
|------|------|
| `../template.docx` | 系统内 `{{字段}}` 填充母版 |
| `法大大上传_template.docx` | 复制到法大大控制台上传 |
| `示例_已填充_供预览.docx` | 上传前检查排版（可选） |

完整条款正文见：`FHD/合同/软件服务合同_模板.md`（Markdown 源稿，内容较长，未整份塞进 Word 以免单文件过大）。

## 2. 法大大控制台操作顺序

1. 登录 [法大大开放平台](https://dev.fadada.com/) → 创建应用 → 记录 **AppId**、**AppSecret**。
2. 企业认证 → 记录企业 **openCorpId**（即环境变量 `FADADA_OPEN_CORP_ID`）。
3. **模板管理** → 上传 `法大大上传_template.docx`（若平台提示过大，可只上传主合同前六条 + 签署页，或联系法大大技术支持拆册）。
4. 在签署模板中配置参与方（必须与代码一致）：
   - **发起方**：企业，对应乙方「成都修茈科技有限公司」（`FADADA_OPEN_CORP_ID`）
   - **签署方 actorId**：`甲方`（客户个人签署，对应 `FADADA_SIGN_ACTOR_ID=甲方`）
5. 拖好签章/签字控件后，发布模板 → 复制 **签署模板 ID** → 填入 `FADADA_SIGN_TEMPLATE_ID`。
6. **回调地址**：`https://你的域名/api/contract-lifecycle/esign/webhook`

## 3. FHD 环境变量（复制 `config/fadada.env.example` → `.env`）

```bash
ESIGN_PROVIDER=fadada
FADADA_APP_ID=
FADADA_APP_SECRET=
FADADA_OPEN_CORP_ID=
FADADA_SIGN_TEMPLATE_ID=
FADADA_SERVER_URL=https://api.fadada.com/api/v5
FADADA_CALLBACK_URL=https://你的域名/api/contract-lifecycle/esign/webhook
FADADA_SIGN_ACTOR_ID=甲方
```

## 4. XCAGI 侧业务流程

1. 内部客服 / 财务统计 → 电子签章 → 填写市场用户 ID → **发起电子签**
2. 系统调用法大大 `create-with-template`，`transReferenceId` = `market_user_id`
3. 将返回的 `sign_url` 发给客户（甲方）
4. 客户签署完成后，法大大回调 webhook → 合同状态 **effective**

## 5. 字段与法大大控件（可选）

若需在法大大模板里做「填写控件」，字段 key 与 `fill_config.json` 一致，例如：

- `party_a_name`、`party_a_credit_code`
- `total_amount_number`、`total_amount_upper`
- `service_start_date`、`service_end_date`

完整字段清单见同目录 `template_fields_manifest.json`。

## 6. 乙方工商信息（模板默认）

"""
        + (f"- 乙方：成都修茈科技有限公司\n- 统一社会信用代码：{party_b or '91510100MAEUUDXP3J'}\n" if party_b else "")
        + """
## 7. 常见问题

- **文件过大**：开放平台单文件限制因账号而异；优先上传本仓库生成的 `法大大上传_template.docx`（已精简条款），完整版仅作 CRM 生成 PDF/Word 归档。
- **actorId 不匹配**：接口报错时检查模板参与方 ID 是否为 `甲方`。
- **回调未生效**：确认公网 HTTPS 可达，且响应 body 为 `{"msg":"success"}`。
""",
        encoding="utf-8",
    )
    return readme


def write_fields_manifest() -> Path:
    cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    fields = cfg.get("fields") or {}
    rows = []
    for key, fdef in fields.items():
        rows.append(
            {
                "key": key,
                "label": fdef.get("label"),
                "placeholder": fdef.get("placeholder") or f"{{{{{key}}}}}",
                "required": fdef.get("required"),
                "group": fdef.get("group"),
            }
        )
    out = FADADA_DIR / "template_fields_manifest.json"
    out.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def main() -> None:
    import sys

    if "--force" not in sys.argv and TEMPLATE_PATH.is_file():
        print(
            "已存在完善版 template.docx，跳过生成。若确要覆盖请加 --force",
            file=sys.stderr,
        )
        return
    path = build_template_docx()
    size_kb = path.stat().st_size // 1024
    print(f"OK template.docx ({size_kb} KB) -> {path}")

    FADADA_DIR.mkdir(parents=True, exist_ok=True)
    upload_copy = FADADA_DIR / "法大大上传_template.docx"
    shutil.copy2(path, upload_copy)
    print(f"OK copy -> {upload_copy}")

    sample = build_sample_filled_docx()
    print(f"OK sample -> {sample}")

    write_fields_manifest()
    write_fadada_readme()
    print(f"OK fadada docs -> {FADADA_DIR}")


if __name__ == "__main__":
    main()
