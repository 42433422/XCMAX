from docx import Document

doc = Document(r'e:\FHD\424\转 Word_扫描全能王 11.04.26 13.31.docx')

print('=== 段落详细信息 ===')
for i, para in enumerate(doc.paragraphs):
    print(f'段落{i}: "{para.text}" | 对齐:{para.alignment}')
    for run in para.runs:
        print(f'  Run: "{run.text}" | Bold:{run.font.bold} | Size:{run.font.size}')

print()
print('=== 表格单元格详细信息 ===')
table = doc.tables[0]
for r_idx, row in enumerate(table.rows):
    for c_idx, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            if para.text.strip():
                print(f'[{r_idx},{c_idx}] "{para.text}" | 对齐:{para.alignment}')
                for run in para.runs:
                    print(f'  Run: "{run.text}" | Bold:{run.font.bold} | Size:{run.font.size}')