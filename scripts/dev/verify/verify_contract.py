from docx import Document

doc = Document(r'e:\FHD\销售合同送货凭证_新.docx')

print('=== 段落内容 ===')
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'段落{i}: {repr(para.text)}')

print()
print('=== 表格内容 ===')
table = doc.tables[0]
for r_idx, row in enumerate(table.rows):
    row_data = []
    for cell in row.cells:
        text = cell.text.strip()
        row_data.append(text if text else '')
    if any(row_data):
        print(f'行{r_idx}: {row_data}')