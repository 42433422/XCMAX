"""Word 全量提取员工：检测、规则、兜底 convert 与包体验证。"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

WORD_EXTRACT_KEYWORDS = (
    "word",
    "docx",
    ".doc",
    "文档",
    "文本",
    "txt",
)
WORD_EXTRACT_ACTION_KEYWORDS = (
    "提取",
    "解析",
    "保存",
    "转换",
    "全量",
    "格式",
    "信息",
    "extract",
    "parse",
)

WORD_OUTPUT_FIELDS = (
    "metadata",
    "paragraphs",
    "tables",
    "headers_footers",
    "styles",
    "images",
    "comments",
    "core_properties",
    "plain_text",
    "outline",
    "blocks",
    "sections",
)


def is_word_full_extract(brief: str) -> bool:
    from modstore_server.csv_tabular_runtime import is_csv_full_read, is_csv_generate
    from modstore_server.excel_tabular_runtime import is_excel_full_read, is_excel_generate
    from modstore_server.pdf_extract_runtime import is_pdf_full_read, is_pdf_generate
    from modstore_server.ppt_extract_runtime import is_ppt_full_read, is_ppt_generate
    from modstore_server.txt_extract_runtime import is_txt_full_read, is_txt_generate

    if is_csv_full_read(brief) or is_csv_generate(brief):
        return False
    if is_excel_full_read(brief) or is_excel_generate(brief):
        return False
    if is_txt_full_read(brief) or is_txt_generate(brief):
        return False
    if is_pdf_full_read(brief) or is_pdf_generate(brief):
        return False
    if is_ppt_full_read(brief) or is_ppt_generate(brief):
        return False
    bl = (brief or "").lower()
    has_word_doc = any(k in bl for k in ("word", "docx", ".doc"))
    has_txt_only = (
        any(k in bl for k in (".txt", "txt文件", "txt 文件", "纯文本")) and not has_word_doc
    )
    if has_txt_only:
        return False
    if (
        any(k in bl for k in ("生成", "写入", "write", "generate", "重建", "render"))
        and "提取" not in bl
        and "全量提取" not in bl
    ):
        return False
    has_doc = any(k in bl for k in WORD_EXTRACT_KEYWORDS)
    has_action = any(k in bl for k in WORD_EXTRACT_ACTION_KEYWORDS)
    if not has_doc or not has_action:
        return False
    if not has_word_doc and "txt" in bl:
        return False
    if any(
        k in bl
        for k in (
            "合同",
            "法务",
            "合规",
            "审核",
            "条款",
            "contract",
            "legal",
            "compliance",
            "review",
        )
    ):
        return False
    return True


def word_extract_structured_spec(brief: str) -> Dict[str, Any]:
    return {
        "domain": "文档处理 / Word 全量提取",
        "goal": (brief or "").strip().splitlines()[0][:200] or "全量提取 Word 文档所有格式与信息",
        "input": "用户上传的 .docx / .doc 文件",
        "output": "document_full.json（结构化全量）+ document_full.txt + images/",
        "output_schema": {
            "fields": list(WORD_OUTPUT_FIELDS),
            "json_file": "outputs/document_full.json",
            "text_file": "outputs/document_full.txt",
            "images_dir": "outputs/images/",
        },
        "constraints": [
            "必须真实解析 docx，禁止 LLM 编造内容",
            "覆盖段落、标题层级、列表、分章节、表格、字体段落样式、图片、页眉页脚、元数据、批注",
            "输出 outline、blocks、sections 与 paragraphs/tables 并存",
            "handlers 必须为 direct_python",
        ],
        "suggested_capabilities": [
            "doc.full_extract",
            "doc.tables",
            "doc.images",
            "doc.metadata",
        ],
        "suggested_handlers": ["direct_python"],
    }


def build_word_extract_rule_spec(brief: str) -> Dict[str, Any]:
    return {
        "brief": brief,
        "mode": "direct_python_file_transform",
        "accepted_extensions": [".docx", ".doc"],
        "default_action": "convert",
        "default_output_relpath": "outputs/document_full.json",
        "default_text_output_relpath": "outputs/document_full.txt",
        "default_images_dir": "outputs/images",
        "runtime_kind": "word_full_extract",
        "output_schema": list(WORD_OUTPUT_FIELDS),
        "requirements": [
            'Use direct_python only; handlers must be ["direct_python"].',
            "Extract paragraphs (headings/lists/fonts), tables, outline, blocks, sections, images, styles, headers/footers, core properties, comments.",
            "Write document_full.json and document_full.txt; export images to outputs/images/.",
            "Never claim success unless output files are actually written.",
            "Return {ok, summary, items, warnings, error, meta}.",
        ],
    }


def render_word_fallback_convert_module() -> str:
    return r"""from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from xml.etree import ElementTree as ET

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_CP_NS = "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}"
_DC_NS = "{http://purl.org/dc/elements/1.1}"
_DCT_NS = "{http://purl.org/dc/terms/}"


def _text_of(el: ET.Element) -> str:
    parts: List[str] = []
    if el.text:
        parts.append(el.text)
    for child in el.iter():
        if child is not el and child.text:
            parts.append(child.text)
        if child.tail:
            parts.append(child.tail)
    return "".join(parts).strip()


def _parse_core_props(zf: zipfile.ZipFile) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for name in ("docProps/core.xml", "docProps/app.xml"):
        if name not in zf.namelist():
            continue
        root = ET.fromstring(zf.read(name))
        for child in root:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            val = (child.text or "").strip()
            if val:
                out[tag] = val
    return out


def _parse_styles(zf: zipfile.ZipFile) -> List[Dict[str, Any]]:
    path = "word/styles.xml"
    if path not in zf.namelist():
        return []
    root = ET.fromstring(zf.read(path))
    styles: List[Dict[str, Any]] = []
    for st in root.findall(f"{_W_NS}style"):
        sid = st.get(f"{_W_NS}styleId") or st.get("styleId") or ""
        stype = st.get(f"{_W_NS}type") or st.get("type") or ""
        name_el = st.find(f"{_W_NS}name")
        name = name_el.get(f"{_W_NS}val") if name_el is not None else ""
        styles.append({"id": sid, "type": stype, "name": name})
    return styles


def _int_attr(el: Optional[ET.Element], attr: str) -> Optional[int]:
    if el is None:
        return None
    raw = el.get(f"{_W_NS}{attr}") or el.get(attr)
    if raw is None:
        return None
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def _parse_run_style(rpr: Optional[ET.Element]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    if rpr is None:
        return out
    if rpr.find(f"{_W_NS}b") is not None:
        out["bold"] = True
    if rpr.find(f"{_W_NS}i") is not None:
        out["italic"] = True
    if rpr.find(f"{_W_NS}u") is not None:
        out["underline"] = True
    fonts = rpr.find(f"{_W_NS}rFonts")
    if fonts is not None:
        for key in ("ascii", "eastAsia", "hAnsi", "cs"):
            val = fonts.get(f"{_W_NS}{key}") or fonts.get(key)
            if val:
                out.setdefault("font_names", {})[key] = val
    sz = rpr.find(f"{_W_NS}sz")
    if sz is not None:
        half = _int_attr(sz, "val")
        if half is not None:
            out["font_size_half_pt"] = half
    color = rpr.find(f"{_W_NS}color")
    if color is not None:
        cval = color.get(f"{_W_NS}val") or color.get("val")
        if cval and cval != "auto":
            out["color"] = cval
    return out


def _parse_paragraph_format(ppr: Optional[ET.Element]) -> Dict[str, Any]:
    fmt: Dict[str, Any] = {}
    if ppr is None:
        return fmt
    jc = ppr.find(f"{_W_NS}jc")
    if jc is not None:
        fmt["alignment"] = jc.get(f"{_W_NS}val") or jc.get("val") or ""
    ind = ppr.find(f"{_W_NS}ind")
    if ind is not None:
        for key in ("left", "right", "firstLine", "hanging"):
            v = _int_attr(ind, key)
            if v is not None:
                fmt.setdefault("indent", {})[key] = v
    spacing = ppr.find(f"{_W_NS}spacing")
    if spacing is not None:
        for key in ("before", "after", "line", "lineRule"):
            v = spacing.get(f"{_W_NS}{key}") or spacing.get(key)
            if v is not None:
                fmt.setdefault("spacing", {})[key] = v
    return fmt


def _heading_level_from_style(style: str, outline_lvl: Optional[int]) -> Tuple[bool, Optional[int]]:
    if outline_lvl is not None and 0 <= outline_lvl <= 8:
        return True, outline_lvl + 1
    sl = (style or "").lower()
    for n in range(1, 10):
        if f"heading{n}" in sl.replace(" ", "") or f"标题{n}" in style or f"heading {n}" in sl:
            return True, n
    return False, None


def _parse_paragraphs(body: ET.Element) -> List[Dict[str, Any]]:
    paras: List[Dict[str, Any]] = []
    for idx, p in enumerate(body.findall(f"{_W_NS}p")):
        ppr = p.find(f"{_W_NS}pPr")
        style = ""
        outline_lvl: Optional[int] = None
        list_type = ""
        list_level: Optional[int] = None
        para_fmt: Dict[str, Any] = {}
        if ppr is not None:
            ps = ppr.find(f"{_W_NS}pStyle")
            if ps is not None:
                style = ps.get(f"{_W_NS}val") or ps.get("val") or ""
            ol = ppr.find(f"{_W_NS}outlineLvl")
            outline_lvl = _int_attr(ol, "val")
            num_pr = ppr.find(f"{_W_NS}numPr")
            if num_pr is not None:
                ilvl = num_pr.find(f"{_W_NS}ilvl")
                list_level = _int_attr(ilvl, "val")
                num_id_el = num_pr.find(f"{_W_NS}numId")
                num_id = _int_attr(num_id_el, "val")
                list_type = "numbered" if (num_id or 0) > 0 else "bullet"
            para_fmt = _parse_paragraph_format(ppr)
        runs = []
        for r in p.findall(f"{_W_NS}r"):
            run_style = _parse_run_style(r.find(f"{_W_NS}rPr"))
            text = _text_of(r)
            if text:
                runs.append({"text": text, "style": run_style})
        text = _text_of(p)
        is_heading, heading_level = _heading_level_from_style(style, outline_lvl)
        if text or runs or is_heading or list_type:
            entry: Dict[str, Any] = {
                "index": idx,
                "style": style,
                "text": text,
                "runs": runs,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "paragraph_format": para_fmt,
            }
            if list_type:
                entry["list_type"] = list_type
                entry["list_level"] = list_level if list_level is not None else 0
            paras.append(entry)
    return paras


def _parse_sections(body: ET.Element, paragraph_count: int) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    sec_idx = 0
    para_cursor = 0
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            ppr = child.find(f"{_W_NS}pPr")
            if ppr is not None and ppr.find(f"{_W_NS}sectPr") is not None:
                sections.append({
                    "section_index": sec_idx,
                    "kind": "paragraph_sectPr",
                    "paragraph_index": para_cursor,
                })
                sec_idx += 1
            para_cursor += 1
        elif tag == "sectPr":
            sections.append({
                "section_index": sec_idx,
                "kind": "body_sectPr",
                "paragraph_index": para_cursor,
            })
            sec_idx += 1
    if not sections and paragraph_count:
        sections.append({"section_index": 0, "kind": "default", "paragraph_index": 0, "paragraph_end": paragraph_count - 1})
    return sections


def _build_outline(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    outline: List[Dict[str, Any]] = []
    stack: List[Dict[str, Any]] = []
    for p in paragraphs:
        if not p.get("is_heading"):
            continue
        level = int(p.get("heading_level") or 1)
        node = {
            "level": level,
            "title": p.get("text") or "",
            "paragraph_index": p.get("index"),
            "children": [],
        }
        while stack and int(stack[-1].get("level") or 1) >= level:
            stack.pop()
        if stack:
            stack[-1]["children"].append(node)
        else:
            outline.append(node)
        stack.append(node)
    return outline


def _build_blocks(body: ET.Element, paragraphs: List[Dict[str, Any]], tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    p_idx = 0
    t_idx = 0
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if p_idx < len(paragraphs):
                p = paragraphs[p_idx]
                blocks.append({
                    "type": "paragraph",
                    "index": p.get("index"),
                    "text": p.get("text"),
                    "is_heading": p.get("is_heading"),
                    "heading_level": p.get("heading_level"),
                    "list_type": p.get("list_type"),
                    "list_level": p.get("list_level"),
                })
                ppr = child.find(f"{_W_NS}pPr")
                if ppr is not None and ppr.find(f"{_W_NS}sectPr") is not None:
                    blocks.append({"type": "section_break", "paragraph_index": p.get("index")})
            p_idx += 1
        elif tag == "tbl":
            if t_idx < len(tables):
                tbl = tables[t_idx]
                blocks.append({
                    "type": "table",
                    "index": tbl.get("index"),
                    "row_count": tbl.get("row_count"),
                    "col_count": tbl.get("col_count"),
                })
            t_idx += 1
        elif tag == "sectPr":
            blocks.append({"type": "section_break", "kind": "body_sectPr"})
    return blocks


def _parse_tables(body: ET.Element) -> List[Dict[str, Any]]:
    tables: List[Dict[str, Any]] = []
    for t_idx, tbl in enumerate(body.findall(f"{_W_NS}tbl")):
        rows: List[List[str]] = []
        for tr in tbl.findall(f"{_W_NS}tr"):
            row: List[str] = []
            for tc in tr.findall(f"{_W_NS}tc"):
                row.append(_text_of(tc))
            if row:
                rows.append(row)
        tables.append({"index": t_idx, "rows": rows, "row_count": len(rows), "col_count": max((len(r) for r in rows), default=0)})
    return tables


def _parse_header_footer(zf: zipfile.ZipFile, prefix: str) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for name in sorted(zf.namelist()):
        if not name.startswith(f"word/{prefix}") or not name.endswith(".xml"):
            continue
        root = ET.fromstring(zf.read(name))
        body = root.find(f"{_W_NS}body") or root
        out.append({"part": name, "text": _text_of(body), "paragraphs": _parse_paragraphs(body)})
    return out


def _parse_comments(zf: zipfile.ZipFile) -> List[Dict[str, Any]]:
    path = "word/comments.xml"
    if path not in zf.namelist():
        return []
    root = ET.fromstring(zf.read(path))
    comments: List[Dict[str, Any]] = []
    for c in root.findall(f"{_W_NS}comment"):
        cid = c.get(f"{_W_NS}id") or c.get("id") or ""
        author = c.get(f"{_W_NS}author") or c.get("author") or ""
        comments.append({"id": cid, "author": author, "text": _text_of(c)})
    return comments


def _extract_images(zf: zipfile.ZipFile, out_dir: Path) -> List[Dict[str, Any]]:
    images: List[Dict[str, Any]] = []
    out_dir.mkdir(parents=True, exist_ok=True)
    for name in zf.namelist():
        if not name.startswith("word/media/"):
            continue
        data = zf.read(name)
        fname = Path(name).name
        dest = out_dir / fname
        dest.write_bytes(data)
        images.append({"source": name, "filename": fname, "size": len(data), "path": str(dest)})
    return images


def _extract_with_docx(src_path: Path) -> Optional[Dict[str, Any]]:
    try:
        import docx
    except ImportError:
        return None
    doc = docx.Document(str(src_path))
    paras = [{"index": i, "style": (p.style.name if p.style else ""), "text": p.text, "runs": []} for i, p in enumerate(doc.paragraphs) if p.text]
    tables = []
    for ti, tbl in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        tables.append({"index": ti, "rows": rows, "row_count": len(rows), "col_count": max((len(r) for r in rows), default=0)})
    core = {}
    try:
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "keywords", "category", "comments", "created", "modified", "last_modified_by"):
            val = getattr(cp, attr, None)
            if val is not None:
                core[attr] = str(val)
    except Exception:
        pass
    return {"paragraphs": paras, "tables": tables, "core_properties": core}


def _extract_docx_ooxml(src_path: Path, output_dir: Path) -> Dict[str, Any]:
    with zipfile.ZipFile(src_path) as zf:
        doc_xml = zf.read("word/document.xml")
        root = ET.fromstring(doc_xml)
        body = root.find(f"{_W_NS}body")
        if body is None:
            raise ValueError("无效的 docx：缺少 word/document.xml body")
        paragraphs = _parse_paragraphs(body)
        tables = _parse_tables(body)
        sections = _parse_sections(body, len(paragraphs))
        outline = _build_outline(paragraphs)
        blocks = _build_blocks(body, paragraphs, tables)
        styles = _parse_styles(zf)
        headers = _parse_header_footer(zf, "header")
        footers = _parse_header_footer(zf, "footer")
        comments = _parse_comments(zf)
        core = _parse_core_props(zf)
        images_dir = output_dir / "images"
        images = _extract_images(zf, images_dir)
    plain_parts = [p.get("text", "") for p in paragraphs if p.get("text")]
    for tbl in tables:
        for row in tbl.get("rows") or []:
            plain_parts.append(" | ".join(row))
    for hf in headers + footers:
        if hf.get("text"):
            plain_parts.append(hf["text"])
    return {
        "metadata": {"source": src_path.name, "format": "docx", "extractor": "ooxml"},
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
        "outline": outline,
        "blocks": blocks,
        "headers_footers": {"headers": headers, "footers": footers},
        "styles": styles,
        "images": images,
        "comments": comments,
        "core_properties": core,
        "plain_text": "\n".join(plain_parts),
        "stats": {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "image_count": len(images),
            "comment_count": len(comments),
            "style_count": len(styles),
            "section_count": len(sections),
            "outline_node_count": len(outline),
            "block_count": len(blocks),
        },
    }


def convert_file(
    src_path: Path,
    output_path: Path,
    *,
    template_path: Optional[Path] = None,
    payload: Dict[str, Any],
    ctx: Dict[str, Any],
    rule_spec: Dict[str, Any],
) -> Dict[str, Any]:
    from word_full_read.legacy_doc import ensure_docx_for_extract

    output_dir = output_path.parent
    legacy_work = output_dir / ".legacy_doc_work"
    extract_src, legacy_meta = ensure_docx_for_extract(Path(src_path), legacy_work)
    src_path = extract_src

    output_dir.mkdir(parents=True, exist_ok=True)
    json_path = output_path if output_path.suffix.lower() == ".json" else output_dir / "document_full.json"
    txt_path = output_dir / "document_full.txt"
    if str(rule_spec.get("default_text_output_relpath") or "").endswith(".txt"):
        txt_path = output_dir / Path(str(rule_spec.get("default_text_output_relpath"))).name

    docx_enhanced = _extract_with_docx(src_path)
    payload_data = _extract_docx_ooxml(src_path, output_dir)
    if isinstance(legacy_meta, dict) and legacy_meta:
        payload_data.setdefault("metadata", {})
        if isinstance(payload_data["metadata"], dict):
            payload_data["metadata"]["legacy_doc"] = legacy_meta
    if docx_enhanced:
        if docx_enhanced.get("core_properties"):
            payload_data["core_properties"] = {**payload_data.get("core_properties", {}), **docx_enhanced["core_properties"]}
        if len(docx_enhanced.get("paragraphs") or []) > len(payload_data.get("paragraphs") or []):
            payload_data["paragraphs"] = docx_enhanced["paragraphs"]
        if len(docx_enhanced.get("tables") or []) > len(payload_data.get("tables") or []):
            payload_data["tables"] = docx_enhanced["tables"]
        payload_data["metadata"]["extractor"] = "ooxml+python-docx"

    json_path.write_text(json.dumps(payload_data, ensure_ascii=False, indent=2), encoding="utf-8")
    txt_path.write_text(payload_data.get("plain_text") or "", encoding="utf-8")

    stats = payload_data.get("stats") or {}
    return {
        "output_path": str(json_path),
        "text_output_path": str(txt_path),
        "images_dir": str(output_dir / "images"),
        "paragraph_count": stats.get("paragraph_count", 0),
        "table_count": stats.get("table_count", 0),
        "image_count": stats.get("image_count", 0),
        "style_count": stats.get("style_count", 0),
        "comment_count": stats.get("comment_count", 0),
        "output_schema": list(rule_spec.get("output_schema") or []),
    }
"""


def validate_word_extract_backend(pack_dir: Path) -> Tuple[List[str], List[str]]:
    errors: List[str] = []
    warnings: List[str] = []
    backend = pack_dir / "backend"
    if not backend.is_dir():
        errors.append("缺少 backend 目录")
        return errors, warnings

    py_blob = ""
    has_convert = False
    for py_path in backend.rglob("*.py"):
        try:
            text = py_path.read_text(encoding="utf-8", errors="ignore")
            py_blob += text.lower()
            if "def convert_file" in text and "vendor" in str(py_path).lower():
                has_convert = True
        except OSError:
            pass

    mf_path = pack_dir / "manifest.json"
    handlers: List[str] = []
    if mf_path.is_file():
        try:
            from modstore_server.employee_asset_pipeline import manifest_actions_handlers

            mf = json.loads(mf_path.read_text(encoding="utf-8"))
            handlers = manifest_actions_handlers(mf)
        except (json.JSONDecodeError, OSError):
            warnings.append("manifest.json 无法解析")

    if handlers and "direct_python" not in handlers:
        errors.append("Word 全量提取员工 handlers 必须包含 direct_python")
    if not has_convert:
        errors.append("backend/vendor 中缺少 convert_file 实现")
    if not any(tok in py_blob for tok in ("docx", "zipfile", "wordprocessingml", "document.xml")):
        warnings.append("未发现 docx/OOXML 解析相关代码")
    rs_path = pack_dir / "rule_spec.json"
    accepted_ext: List[str] = []
    if rs_path.is_file():
        try:
            rs_data = json.loads(rs_path.read_text(encoding="utf-8"))
            if isinstance(rs_data, dict):
                accepted_ext = list(rs_data.get("accepted_extensions") or [])
        except (json.JSONDecodeError, OSError):
            pass
    if ".doc" in accepted_ext or "legacy_doc" in py_blob:
        if "legacy_doc" not in py_blob and "ensure_docx_for_extract" not in py_blob:
            warnings.append("已声明 .doc 但 convert 未集成 legacy_doc.ensure_docx_for_extract")

    for field in (
        "paragraphs",
        "tables",
        "images",
        "core_properties",
        "outline",
        "blocks",
        "sections",
    ):
        if field not in py_blob:
            warnings.append(f"convert 模块可能未覆盖输出字段：{field}")
    for tok in ("heading_level", "list_type", "outline", "numpr", "outlinelvl", "sectpr"):
        if tok not in py_blob:
            warnings.append(f"convert 模块可能未覆盖解析能力：{tok}")

    return errors, warnings


def minimal_docx_bytes() -> bytes:
    """Minimal valid OOXML docx for pipeline smoke tests."""
    import io
    import zipfile

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body><w:p><w:r><w:t>smoke</w:t></w:r></w:p></w:body></w:document>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document)
    return buf.getvalue()


def minimal_docx_fixture_b64() -> str:
    import base64

    return base64.b64encode(minimal_docx_bytes()).decode("ascii")


def word_extract_orchestration_plan(brief: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    from modstore_server.employee_brief_utils import compact_routing_brief

    checklist = payload.get("execution_checklist")
    checklist_text = (
        "\n".join(f"- {x}" for x in checklist if isinstance(x, str))
        if isinstance(checklist, list)
        else ""
    )
    clean = compact_routing_brief(brief, max_len=400) or (brief or "").strip()
    merged = "\n".join(x for x in [clean, checklist_text] if x).strip()
    short = "Word 全量提取员"
    script_brief = (
        f"{merged or clean}\n\n"
        "请生成 Python 脚本：读取 inputs/ 中的 .docx，全量提取段落、表格、图片、样式、页眉页脚、元数据、批注；"
        "写入 outputs/document_full.json、outputs/document_full.txt，图片导出到 outputs/images/；"
        "无输入时在 outputs/readme.md 说明期望输入格式。"
    )
    script_runtime = (
        "只能读 inputs/、写 outputs/；允许 python-docx、zipfile、xml.etree；"
        "禁止联网和越界文件访问；必须输出 JSON + TXT + images。"
    )
    workflow_brief = (
        f"{merged or clean}\n\n"
        "Skill 组流程：①接收 Word 上传 ②校验格式 ③全量解析（正文/表格/图片/样式/元数据）"
        "④生成 document_full.json + txt ⑤导出图片 ⑥质量校验 ⑦交付用户。"
    )
    return {
        "employee_name": short,
        "employee_brief": (
            f"{merged or clean}\n\n"
            "员工必须使用 direct_python 真实解析 docx，输出结构化 JSON（含 paragraphs/tables/images/styles/metadata）"
            "与纯文本摘要，禁止 LLM 编造文档内容。"
        ),
        "script_workflow_name": f"{short} 脚本工作流",
        "script_brief": script_brief,
        "script_runtime_notes": script_runtime,
        "workflow_name": str(payload.get("employee_workflow_name") or short).strip() or short,
        "workflow_brief": workflow_brief,
        "acceptance": [
            "员工包 handlers 为 direct_python 且含 vendor convert 模块",
            "输出 document_full.json 含 paragraphs/tables/images/core_properties",
            "脚本工作流可空跑并生成 outputs/ 说明或样例",
            "Skill 组覆盖上传→解析→校验→交付",
        ],
    }
