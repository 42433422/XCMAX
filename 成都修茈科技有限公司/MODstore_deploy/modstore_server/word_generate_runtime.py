"""Word 生成员工：JSON（+ 可选模板 docx）→ 生成 .docx。"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Tuple

WORD_GEN_DOC_KEYWORDS = (
    "word",
    "docx",
    ".doc",
    "文档",
    "word文档",
)
WORD_GEN_ACTION_KEYWORDS = (
    "生成",
    "写入",
    "写word",
    "写 docx",
    "输出",
    "重建",
    "导出",
    "write",
    "generate",
    "render",
    "build",
)
WORD_GEN_EXCLUDE = (
    "全量提取",
    "仅提取",
    "只提取",
    "仅读取",
    "只读",
    "extract only",
    "read only",
)
WORD_GEN_OUTPUT_FIELDS = (
    "metadata",
    "paragraphs",
    "tables",
    "outline",
    "blocks",
    "plain_text",
    "stats",
)


def _brief_lower(brief: str) -> str:
    return (brief or "").lower()


def _has_word_doc_signal(bl: str) -> bool:
    return any(k in bl for k in WORD_GEN_DOC_KEYWORDS)


def is_word_generate(brief: str) -> bool:
    """JSON（document_full.json 等同 schema）+ 可选模板 → 生成 .docx。"""
    bl = _brief_lower(brief)
    if not _has_word_doc_signal(bl):
        return False
    if any(k in bl for k in WORD_GEN_EXCLUDE):
        return False
    has_gen = any(k in bl for k in WORD_GEN_ACTION_KEYWORDS)
    if not has_gen:
        return False
    if "json" in bl or "document_full" in bl or "模板" in bl or "template" in bl:
        return True
    from modstore_server.word_extract_runtime import is_word_full_extract

    if is_word_full_extract(brief) and not has_gen:
        return False
    if has_gen and any(k in bl for k in ("生成", "写入", "write", "generate", "render", "build")):
        return True
    return False


def word_generate_structured_spec(brief: str) -> Dict[str, Any]:
    return {
        "domain": "文档处理 / Word 生成",
        "goal": (brief or "").strip().splitlines()[0][:200] or "由 JSON 结构化数据生成 Word 文档",
        "input": "document_full.json / user_query 纯文本 / .txt + 可选 template.docx",
        "output": "outputs/generated_document.docx",
        "output_schema": {
            "fields": list(WORD_GEN_OUTPUT_FIELDS),
            "docx_file": "outputs/generated_document.docx",
            "json_input": "inputs/document_full.json",
            "template_file": "inputs/template.docx（可选）",
        },
        "constraints": [
            "必须基于真实 JSON 输入写 docx，禁止无输入时 LLM 编造正文",
            "handlers 必须含 direct_python；润色/改写可选用 agent",
            "兼容 paragraphs/tables/outline/blocks 字段",
        ],
        "suggested_capabilities": ["doc.generate", "doc.template_merge", "doc.styles"],
        "suggested_handlers": ["direct_python", "agent"],
    }


def build_word_generate_rule_spec(brief: str) -> Dict[str, Any]:
    bl = _brief_lower(brief)
    wants_polish = any(k in bl for k in ("润色", "改写", "polish", "rewrite"))
    return {
        "brief": brief,
        "mode": "direct_python_file_transform",
        "accepted_extensions": [".json", ".txt"],
        "template_extensions": [".docx"],
        "default_action": "convert",
        "default_output_relpath": "outputs/generated_document.docx",
        "default_json_input": "inputs/document_full.json",
        "default_template_relpath": "inputs/template.docx",
        "runtime_kind": "word_generate",
        "optional_llm_polish": wants_polish,
        "output_schema": list(WORD_GEN_OUTPUT_FIELDS),
        "requirements": [
            'handlers must include "direct_python"; may include "agent" for optional polish.',
            "Read JSON (document_full.json schema); optional template.docx for styles.",
            "Write outputs/generated_document.docx with headings, lists, tables, fonts where available.",
            "Never fabricate content when JSON input is empty.",
            "Return {ok, summary, items, warnings, error, meta}.",
        ],
    }


def render_word_generate_convert_module() -> str:
    return r"""from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from modstore_server.office_plaintext_generate import resolve_word_document_spec


def _style_for_paragraph(doc: Any, p: Dict[str, Any]) -> Optional[str]:
    if p.get("is_heading") and p.get("heading_level"):
        lvl = int(p["heading_level"])
        name = f"Heading {min(max(lvl, 1), 9)}"
        try:
            _ = doc.styles[name]
            return name
        except KeyError:
            return None
    style = str(p.get("style") or "").strip()
    if style:
        try:
            _ = doc.styles[style]
            return style
        except KeyError:
            pass
    return None


def _apply_run_format(run: Any, run_data: Dict[str, Any]) -> None:
    style = run_data.get("style") or {}
    if not isinstance(style, dict):
        return
    if style.get("bold"):
        run.bold = True
    if style.get("italic"):
        run.italic = True
    if style.get("underline"):
        run.underline = True
    fn = style.get("font_names") or {}
    if isinstance(fn, dict):
        for key in ("eastAsia", "ascii", "hAnsi"):
            if fn.get(key):
                try:
                    run.font.name = fn[key]
                    break
                except Exception:
                    pass
    half = style.get("font_size_half_pt")
    if half is not None:
        try:
            from docx.shared import Pt
            run.font.size = Pt(int(half) / 2)
        except Exception:
            pass


def _add_paragraph(doc: Any, p: Dict[str, Any]) -> None:
    text = str(p.get("text") or "").strip()
    runs = p.get("runs") or []
    style_name = _style_for_paragraph(doc, p)
    para = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    list_type = p.get("list_type")
    if list_type in ("numbered", "bullet"):
        try:
            para.style = "List Number" if list_type == "numbered" else "List Bullet"
        except KeyError:
            pass
    if runs:
        for rd in runs:
            if not isinstance(rd, dict):
                continue
            rt = str(rd.get("text") or "")
            if not rt:
                continue
            run = para.add_run(rt)
            _apply_run_format(run, rd)
    elif text:
        para.add_run(text)


def _add_table(doc: Any, tbl: Dict[str, Any]) -> None:
    rows = tbl.get("rows") or []
    if not rows:
        return
    nrows = len(rows)
    ncols = max(len(r) for r in rows)
    if ncols < 1:
        return
    table = doc.add_table(rows=nrows, cols=ncols)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = str(val or "")


def _build_from_blocks(doc: Any, spec: Dict[str, Any]) -> None:
    blocks = spec.get("blocks") or []
    if blocks:
        for blk in blocks:
            if not isinstance(blk, dict):
                continue
            btype = blk.get("type")
            if btype == "paragraph":
                paras = spec.get("paragraphs") or []
                idx = blk.get("index")
                p = next((x for x in paras if x.get("index") == idx), None)
                if p:
                    _add_paragraph(doc, p)
                elif blk.get("text"):
                    _add_paragraph(doc, blk)
            elif btype == "table":
                tables = spec.get("tables") or []
                ti = blk.get("index")
                t = next((x for x in tables if x.get("index") == ti), None)
                if t:
                    _add_table(doc, t)
        return
    for p in spec.get("paragraphs") or []:
        if isinstance(p, dict):
            _add_paragraph(doc, p)
    for t in spec.get("tables") or []:
        if isinstance(t, dict):
            _add_table(doc, t)


async def convert_file(
    src_path: Path,
    output_path: Path,
    *,
    template_path: Optional[Path] = None,
    payload: Dict[str, Any],
    ctx: Dict[str, Any],
    rule_spec: Dict[str, Any],
) -> Dict[str, Any]:
    suffix = src_path.suffix.lower()
    if suffix not in (".json", ".txt"):
        raise ValueError(f"不支持的文件类型：{suffix or '(无后缀)'}，支持 .json / .txt")

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法生成 Word") from exc

    spec, _warnings = await resolve_word_document_spec(src_path, payload or {}, ctx or {}, rule_spec or {})
    if not (spec.get("paragraphs") or spec.get("blocks") or spec.get("plain_text")):
        raise ValueError("缺少 paragraphs/blocks/plain_text，无法生成 docx")

    output_dir = output_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_path
    if output_path.suffix.lower() != ".docx":
        docx_path = output_dir / "generated_document.docx"
    rel = str(rule_spec.get("default_output_relpath") or "")
    if rel.endswith(".docx"):
        docx_path = output_dir / Path(rel).name

    tpl = template_path
    if tpl is None or not tpl.is_file():
        tpl_rel = str(rule_spec.get("default_template_relpath") or "")
        if tpl_rel:
            cand = src_path.parent.parent / tpl_rel
            if cand.is_file():
                tpl = cand
        inputs_tpl = src_path.parent / "template.docx"
        if inputs_tpl.is_file():
            tpl = inputs_tpl

    if tpl and tpl.is_file():
        doc = Document(str(tpl))
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)
    else:
        doc = Document()

    _build_from_blocks(doc, spec)
    if not doc.paragraphs and spec.get("plain_text"):
        doc.add_paragraph(str(spec.get("plain_text") or ""))

    doc.save(str(docx_path))

    para_count = len(spec.get("paragraphs") or [])
    table_count = len(spec.get("tables") or [])
    return {
        "output_path": str(docx_path),
        "template_path": str(tpl or ""),
        "paragraph_count": para_count,
        "table_count": table_count,
        "output_schema": list(rule_spec.get("output_schema") or []),
        "source_json": src_path.name,
    }
"""


def validate_word_generate_backend(pack_dir: Path) -> Tuple[List[str], List[str]]:
    errors: List[str] = []
    warnings: List[str] = []
    backend = pack_dir / "backend"
    if not backend.is_dir():
        errors.append("缺少 backend 目录")
        return errors, warnings

    py_blob = ""
    has_convert = False
    for py_path in backend.rglob("*.py"):
        try:
            text = py_path.read_text(encoding="utf-8", errors="ignore")
            py_blob += text.lower()
            if "def convert_file" in text and "vendor" in str(py_path).lower():
                has_convert = True
        except OSError:
            pass

    mf_path = pack_dir / "manifest.json"
    handlers: List[str] = []
    if mf_path.is_file():
        try:
            from modstore_server.employee_asset_pipeline import manifest_actions_handlers

            mf = json.loads(mf_path.read_text(encoding="utf-8"))
            handlers = manifest_actions_handlers(mf)
        except (json.JSONDecodeError, OSError):
            warnings.append("manifest.json 无法解析")

    if handlers and "direct_python" not in handlers:
        errors.append("Word 生成员工 handlers 必须包含 direct_python")
    if not has_convert:
        errors.append("backend/vendor 中缺少 convert_file 实现")
    if not any(
        tok in py_blob for tok in ("docx", "document", "add_paragraph", "generated_document")
    ):
        warnings.append("未发现 docx 生成相关代码")
    for tok in ("template", "json", "heading"):
        if tok not in py_blob:
            warnings.append(f"convert 模块可能未覆盖：{tok}")

    return errors, warnings


def minimal_document_full_json() -> Dict[str, Any]:
    return {
        "metadata": {"source": "fixture", "format": "json"},
        "paragraphs": [
            {"index": 0, "text": "第一章 概述", "is_heading": True, "heading_level": 1, "runs": []},
            {
                "index": 1,
                "text": "这是正文段落。",
                "is_heading": False,
                "runs": [{"text": "这是正文段落。", "style": {"bold": True}}],
            },
            {"index": 2, "text": "列表项一", "list_type": "bullet", "list_level": 0, "runs": []},
        ],
        "tables": [
            {"index": 0, "rows": [["列A", "列B"], ["1", "2"]], "row_count": 2, "col_count": 2}
        ],
        "blocks": [
            {"type": "paragraph", "index": 0},
            {"type": "paragraph", "index": 1},
            {"type": "paragraph", "index": 2},
            {"type": "table", "index": 0},
        ],
        "outline": [{"level": 1, "title": "第一章 概述", "paragraph_index": 0, "children": []}],
        "plain_text": "第一章 概述\n这是正文段落。\n列表项一",
    }


def word_generate_orchestration_plan(brief: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    from modstore_server.employee_brief_utils import compact_routing_brief

    checklist = payload.get("execution_checklist")
    checklist_text = (
        "\n".join(f"- {x}" for x in checklist if isinstance(x, str))
        if isinstance(checklist, list)
        else ""
    )
    clean = compact_routing_brief(brief, max_len=400) or (brief or "").strip()
    merged = "\n".join(x for x in [clean, checklist_text] if x).strip()
    short = "Word 生成员"
    script_brief = (
        f"{merged or clean}\n\n"
        "请生成 Python：读取 inputs/document_full.json、.txt 或 payload.user_query 纯文本，"
        "可选 inputs/template.docx 作为样式模板；写出 outputs/generated_document.docx。"
    )
    return {
        "employee_name": short,
        "employee_brief": (
            f"{merged or clean}\n\n"
            "direct_python 根据 JSON 或 user_query 纯文本真实写 docx；可选模板保留样式；禁止无输入编造正文。"
        ),
        "script_workflow_name": f"{short} 脚本工作流",
        "script_brief": script_brief,
        "script_runtime_notes": "允许 python-docx；读 inputs/ 写 outputs/；禁止联网。",
        "workflow_name": str(payload.get("employee_workflow_name") or short).strip() or short,
        "workflow_brief": (
            f"{merged or clean}\n\nSkill：上传 JSON → 可选模板 → 生成 docx → 校验 → 交付。"
        ),
        "acceptance": [
            "handlers 含 direct_python",
            "输出 generated_document.docx",
            "兼容 document_full.json schema",
        ],
    }
